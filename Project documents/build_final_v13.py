"""
Samvak Project Review Document — Complete Generator v13
Implements ALL formatting rules from the prompt exactly.
Run: python build_final_v13.py
"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from lxml import etree
import copy

# ── Paths ──────────────────────────────────────────────────────────────────────
BASE   = r"E:\sam\Project documents"
LOGO   = os.path.join(BASE, "extracted_images", "image_0.png")
DIAGRAMS = os.path.join(BASE, "diagrams")
SHOTS    = os.path.join(BASE, "screenshots")
OUT      = os.path.join(BASE, "Samvak_Review_Final_V13b.docx")

# ── Document & section setup ───────────────────────────────────────────────────
doc = Document()

# US Letter, margins per spec
sec = doc.sections[0]
sec.page_width  = Inches(8.5)
sec.page_height = Inches(11)
sec.left_margin   = Inches(0.92)
sec.right_margin  = Inches(0.99)
sec.top_margin    = Inches(0.32)
sec.bottom_margin = Inches(0.19)

# Global Normal style
ns = doc.styles['Normal']
ns.font.name = 'Times New Roman'
ns.font.size = Pt(12)
ns.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
ns.paragraph_format.line_spacing = 1.5
ns.paragraph_format.space_after  = Pt(0)

# ── Page border helper ─────────────────────────────────────────────────────────
def _add_border(section):
    sectPr = section._sectPr
    # remove existing pgBorders if any
    for old in sectPr.findall(qn('w:pgBorders')):
        sectPr.remove(old)
    pgB = parse_xml(
        '<w:pgBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' w:offsetFrom="page">'
        '<w:top w:val="single" w:sz="12" w:space="24" w:color="000000"/>'
        '<w:left w:val="single" w:sz="12" w:space="24" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="12" w:space="24" w:color="000000"/>'
        '<w:right w:val="single" w:sz="12" w:space="24" w:color="000000"/>'
        '</w:pgBorders>'
    )
    sectPr.append(pgB)

_add_border(sec)

# ── Core paragraph helpers ─────────────────────────────────────────────────────
def _set_font(run, size, bold=False, name='Times New Roman'):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold

def para(text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12,
         space_before=0, name='Times New Roman'):
    """Add a paragraph with 1.5 line spacing."""
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing  = 1.5
    pf.space_before  = Pt(space_before)
    pf.space_after   = Pt(0)
    r = p.add_run(text)
    _set_font(r, size, bold, name)
    return p

def chapter(text):
    """16pt bold, left, 1.5 spacing — chapter heading."""
    return para(text, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, size=16, space_before=12)

def subhead(text):
    """14pt bold, left, 1.5 spacing, space_before=12 — sub-heading."""
    return para(text, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, size=14, space_before=12)

def body(text, space_before=12):
    """12pt, justified, 1.5 spacing."""
    return para(text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_before=space_before)

def center(text, bold=False, size=12):
    return para(text, bold=bold, align=WD_ALIGN_PARAGRAPH.CENTER, size=size)

def blank(n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after  = Pt(0)

def pb():
    doc.add_page_break()

# ── Table helper ───────────────────────────────────────────────────────────────
def make_table(headers, rows, col_widths=None):
    """Create a table with black single borders, bold headers, no shading, 12pt TNR, 1.5 spacing."""
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    def _cell(cell, text, bold=False):
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.5
        pf.space_before = Pt(0)
        pf.space_after  = Pt(0)
        r = p.add_run(text)
        _set_font(r, 12, bold)
        # cell padding (DXA: 120 left/right, 80 top/bottom)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        mar = parse_xml(
            '<w:tcMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:top w:w="80" w:type="dxa"/>'
            '<w:left w:w="120" w:type="dxa"/>'
            '<w:bottom w:w="80" w:type="dxa"/>'
            '<w:right w:w="120" w:type="dxa"/>'
            '</w:tcMar>'
        )
        # remove old tcMar
        for old in tcPr.findall(qn('w:tcMar')):
            tcPr.remove(old)
        tcPr.append(mar)

    # header row
    for j, h in enumerate(headers):
        _cell(t.rows[0].cells[j], h, bold=True)

    # data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            _cell(t.rows[i+1].cells[j], str(val))

    # set column widths if given
    if col_widths:
        for row in t.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = Inches(w)

    blank()
    return t

# ── Image helper ───────────────────────────────────────────────────────────────
def insert_image(path, width=6.0, caption=''):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width))
    if caption:
        cp = para(caption, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)

# ── Code block helper ──────────────────────────────────────────────────────────
def code_block(text):
    """Courier New 10pt, light grey background, left-aligned."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing  = 1.5
    pf.space_before  = Pt(6)
    pf.space_after   = Pt(6)
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(10)
    # grey shading on paragraph
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(
        '<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' w:val="clear" w:color="auto" w:fill="F3F4F6"/>'
    )
    pPr.append(shd)
    return p

print("Helpers loaded.")

# ═══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════
blank(2)
if os.path.exists(LOGO):
    lp = doc.add_paragraph()
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp.add_run().add_picture(LOGO, width=Inches(1.6))

blank(1)
center("KONERU LAKSHMAIAH EDUCATION FOUNDATION", bold=True, size=16)
center("(Deemed to be University)", size=11)
center("Department of Computer Science and Engineering", bold=True, size=14)
blank(1)
center("MULTI-LINGUAL SIGN LANGUAGE TO SPEECH\nAND SPEECH TO SIGN TRANSLATOR", bold=True, size=16)
blank(1)
center("(Saṁvāk)", bold=True, size=13)
blank(2)
center("A Project Report", size=12)
center("Submitted in partial fulfillment of the requirements for the award of the degree of", size=12)
blank(1)
center("Master of Computer Applications", bold=True, size=14)
center("In", size=12)
center("Department of Computer Science and Engineering", bold=True, size=13)
blank(1)
center("By", size=12)
blank(1)
center("Uppalapati Naveen Varma", bold=True, size=14)
center("(2401600155)", size=13)
blank(1)
center("Under the supervision of", size=12)
center("Mrs. Swathi Voddi", bold=True, size=13)
center("Assistant Professor", size=12)
blank(2)
center("K L University, Vaddeswaram, Guntur, Andhra Pradesh — 522502", size=11)
center("Academic Year 2024–25", bold=True, size=13)
pb()

# ═══════════════════════════════════════════════════════════════════════════════
# DECLARATION
# ═══════════════════════════════════════════════════════════════════════════════
blank(3)
center("KONERU LAKSHMAIAH EDUCATION FOUNDATION", bold=True, size=13)
center("DEPARTMENT OF COMPUTER SCIENCE AND ENGINEERING", bold=True, size=12)
blank(2)
center("DECLARATION", bold=True, size=16)
blank(2)
body(
    'The Project Report entitled "MULTI-LINGUAL SIGN LANGUAGE TO SPEECH AND SPEECH TO SIGN '
    'TRANSLATOR" is a record of bonafide work carried out by me under the guidance of '
    'Mrs. Swathi Voddi, Assistant Professor, Department of Computer Science and Engineering, '
    'Koneru Lakshmaiah Education Foundation, Vaddeswaram.'
)
blank(1)
body(
    'I hereby declare that this project work is original and has not been submitted to any '
    'other university or institution for the award of any degree or diploma. All the information '
    'furnished in this project report is genuine to the best of my knowledge and belief.'
)
blank(5)
para("Signature of the Student                                    Uppalapati Naveen Varma (2401600155)",
     bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, size=12)
pb()

# ═══════════════════════════════════════════════════════════════════════════════
# CERTIFICATE
# ═══════════════════════════════════════════════════════════════════════════════
blank(3)
center("KONERU LAKSHMAIAH EDUCATION FOUNDATION", bold=True, size=13)
center("DEPARTMENT OF COMPUTER SCIENCE AND ENGINEERING", bold=True, size=12)
blank(2)
center("CERTIFICATE", bold=True, size=16)
blank(2)
body(
    'This is to certify that the Project Report entitled "MULTI-LINGUAL SIGN LANGUAGE TO '
    'SPEECH AND SPEECH TO SIGN TRANSLATOR" is a bonafide record of the work done by '
    'Uppalapati Naveen Varma (2401600155) in partial fulfillment of the requirements for the '
    'award of the degree of Master of Computer Applications from the Department of Computer '
    'Science and Engineering, Koneru Lakshmaiah Education Foundation, Vaddeswaram, during '
    'the academic year 2024–25.'
)
blank(5)
p = para("Signature of the Supervisor", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, size=12)
p.add_run("                                        Signature of the HOD").bold = True
p2 = para("(Mrs. Swathi Voddi)", align=WD_ALIGN_PARAGRAPH.LEFT, size=12)
p2.add_run("                                              (Dr. Ch. Kiran Kumar)")
p3 = para("Assistant Professor", align=WD_ALIGN_PARAGRAPH.LEFT, size=12)
p3.add_run("                                                 Professor & HOD")
blank(4)
para("Signature of the Examiner", align=WD_ALIGN_PARAGRAPH.LEFT, size=12)
pb()

# ═══════════════════════════════════════════════════════════════════════════════
# ACKNOWLEDGEMENT
# ═══════════════════════════════════════════════════════════════════════════════
blank(1)
center("ACKNOWLEDGEMENT", bold=True, size=16)
blank(2)
body(
    'The satisfaction that accompanies the successful completion of any task would be '
    'incomplete without the mention of people who made it possible, whose constant guidance '
    'and encouragement crown all the efforts with success.'
)
body(
    'I am very thankful to my project guide Mrs. Swathi Voddi, Assistant Professor, '
    'Department of Computer Science and Engineering, K L University, for her continuous '
    'support, encouragement, and invaluable guidance in completing this project. Her technical '
    'expertise in the domain of machine learning and computer vision was instrumental in '
    'shaping this work.'
)
body(
    'I express my heartfelt gratitude to Dr. Ch. Kiran Kumar, Head of the Department of '
    'Computer Science and Engineering, for providing me with the opportunity and facilities '
    'to carry out this project successfully.'
)
body(
    'I express my sincere thanks to all the faculty members of the Department of Computer '
    'Science and Engineering for imparting the knowledge that laid the foundation for this '
    'project. Last but not the least, I thank all Teaching and Non-Teaching Staff of our '
    'department and especially my classmates and friends who directly or indirectly helped me '
    'in completing this project.'
)
blank(4)
para("Signature of the Student", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, size=12)
para("Uppalapati Naveen Varma", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, size=12)
para("(2401600155)", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, size=12)
pb()

# ═══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════════════════════
center("TABLE OF CONTENTS", bold=True, size=16)
blank(1)
toc_rows = [
    ("ABSTRACT", "i"),
    ("1.  INTRODUCTION", "1"),
    ("1.1 Motivation:", "1"),
    ("1.2 Problem Statement:", "1"),
    ("1.3 Objective:", "2"),
    ("1.4 Scope:", "2"),
    ("1.5 Project Introduction:", "2"),
    ("2.  LITERATURE SURVEY", "4"),
    ("3.  SYSTEM ANALYSIS", "8"),
    ("3.1 Existing System:", "8"),
    ("3.2 Disadvantages:", "8"),
    ("3.3 Proposed System:", "9"),
    ("3.4 Advantages:", "9"),
    ("3.5 Work Flow:", "9"),
    ("4.  REQUIREMENT ANALYSIS", "10"),
    ("4.1 Functional & Non-Functional Requirements:", "10"),
    ("4.2 Hardware Requirements:", "11"),
    ("4.3 Software Requirements:", "11"),
    ("4.4 Architecture:", "12"),
    ("5.  SYSTEM DESIGN", "13"),
    ("5.1 Input Design:", "13"),
    ("5.2 UML Diagrams:", "14"),
    ("5.3 DFD Diagrams:", "22"),
    ("6.  CODE, IMPLEMENTATION AND RESULTS", "24"),
    ("6.1 app.py — Application Factory", "24"),
    ("6.2 sign.py — LSTM Loader & Prediction", "26"),
    ("6.3 speech.py — Speech-to-Sign Pipeline", "28"),
    ("6.4 grammar_helper.py — ISL Grammar", "30"),
    ("6.5 models.py — Database Models", "32"),
    ("6.6 train_classifier.py — LSTM Training", "34"),
    ("6.7 isign_retrieval.py — Cosine Retrieval", "36"),
    ("6.8 sign.js — Browser ML Pipeline", "38"),
    ("6 .1 Module Descriptions", "58"),
    ("7.  SYSTEM STUDY AND TESTING", "64"),
    ("7.1 Feasibility Study:", "64"),
    ("7.2 Test Cases:", "65"),
    ("8.  CONCLUSION", "69"),
    ("9.  FUTURE ENHANCEMENT", "70"),
    ("10. REFERENCES", "71"),
]
toc_t = doc.add_table(rows=len(toc_rows), cols=2)
toc_t.style = 'Table Grid'
toc_t.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, (title, pg) in enumerate(toc_rows):
    is_ch = not title[0].isdigit() or (title[1] == '.' and title[2] == ' ')
    # check if chapter-level (single digit + dot + spaces)
    is_chapter = len(title) > 2 and title[1:3] in ['.  ', '. '] or title[0:2] in ['AB']
    for j, val in enumerate([title, pg]):
        c = toc_t.rows[i].cells[j]
        c.text = ''
        p2 = c.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.RIGHT
        pf = p2.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.5
        pf.space_after = Pt(0)
        r2 = p2.add_run(val)
        _set_font(r2, 12, bold=True)  # both cols bold per spec
    # set widths
    toc_t.rows[i].cells[0].width = Inches(5.5)
    toc_t.rows[i].cells[1].width = Inches(1.0)
blank()
pb()

# ═══════════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ═══════════════════════════════════════════════════════════════════════════════
center("ABSTRACT", bold=True, size=16)
blank(1)
body(
    'This project presents the design and implementation of Saṁvāk — a Multi-Lingual Sign '
    'Language to Speech and Speech to Sign Translator — a real-time, web-based application '
    'that bridges the communication gap between the hearing-impaired community and the general '
    'population. India alone has approximately 63 million hearing-impaired individuals, yet no '
    'affordable, bidirectional, hardware-free ISL translation system exists. Saṁvāk operates in '
    'two primary modes: (1) Sign-to-Speech, where a webcam captures hand gestures performed in '
    'Indian Sign Language (ISL), processes them through a pipeline using Google MediaPipe '
    'Holistic for 258-dimensional landmark extraction and a TensorFlow/Keras LSTM model '
    'trained on the iSign dataset (14,674 clips, 40 classes) achieving 94.6% top-1 accuracy '
    'on 30-frame sequences; and (2) Speech-to-Sign, where spoken or typed input is transcribed, '
    'converted to ISL gloss via the Gemini API (OSV grammar), and rendered as real-time sign '
    'animations through a Three.js 3D avatar in the browser.'
)
body(
    'The application backend is built with Python 3.10+, Flask 3.x, Flask-SocketIO, '
    'SQLAlchemy, and SQLite. The browser-side ML pipeline uses TensorFlow.js with the '
    'tfjs_lstm/ model bundle and the MediaPipe Holistic JavaScript API, extracting 258 '
    'features per frame and buffering 30 frames for in-browser classification — eliminating '
    'the need for server-side GPU inference. For speech-to-sign translation, speech.py '
    'leverages spaCy (en_core_web_sm) for tokenization and the Google Gemini API for '
    'grammatically correct ISL Object-Subject-Verb (OSV) gloss generation, with a local '
    'rule-based fallback in grammar_helper.py. The 132 MB cosine similarity retrieval index '
    '(isign_retrieval_index.npz) provides nearest-neighbor phrase matching for unknown inputs. '
    'Text-to-speech uses gTTS with a Web Speech API fallback. Authentication is handled by '
    'Flask-Login with bcrypt password hashing, and the system is deployed on Render.com using '
    'Gunicorn and Eventlet workers.'
)
body(
    'Saṁvāk supports output in eight languages (English, Hindi, Telugu, Tamil, Kannada, '
    'Malayalam, Spanish, French) and provides an interactive ISL dictionary powered by the '
    'iSign clip retrieval index, a gamified learning module with XP-based progress tracking, '
    'user authentication, personalized preference management, and full translation history '
    'logging. The core system comprises key files including app.py (factory and blueprints), '
    'sign.py (4,146 lines — LSTM inference and 258-feature extraction), sign.js (1,600+ lines '
    '— browser ML pipeline with TF.js and MediaPipe), speech.py, grammar_helper.py, '
    'models.py, dictionary.py, isign_retrieval.py, geometry_brain.py, and '
    'fingerspell_recognizer.py. Saṁvāk demonstrates that accessible, real-time, bidirectional '
    'sign language translation is achievable using only a webcam, microphone, and a web browser.'
)

pb()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════════
chapter("1.  INTRODUCTION")

subhead("1.1 Motivation:")
body(
    'According to the World Health Organization, over 430 million people worldwide suffer '
    'from disabling hearing loss, with approximately 63 million individuals in India alone '
    'classified as hearing-impaired. The primary mode of communication for the Deaf and '
    'hard-of-hearing community is Indian Sign Language (ISL), yet the vast majority of the '
    'hearing population does not understand sign language. This communication barrier leads '
    'to profound social isolation, limited access to education, restricted healthcare '
    'interactions, and reduced employment opportunities for the hearing-impaired population.'
)
body(
    'While human sign language interpreters serve as intermediaries, their availability is '
    'severely limited, particularly in rural India and during medical emergencies. Indian Sign '
    'Language has its own grammar, syntax, and vocabulary that differ significantly from '
    'spoken languages — ISL follows Object-Subject-Verb (OSV) ordering rather than the '
    'English Subject-Verb-Object (SVO) order — making direct, word-for-word translation '
    'linguistically incorrect and socially harmful. Existing assistive technologies are either '
    'prohibitively expensive (hardware gloves costing $500–$5,000), require specialized '
    'depth-sensing cameras (Microsoft Kinect), operate only in American Sign Language (ASL), '
    'or are restricted to a single output language. These challenges motivated the development '
    'of Saṁvāk — a software-based, hardware-minimal, multilingual, bidirectional solution '
    'that runs on any device equipped with a standard webcam and microphone.'
)

subhead("1.2 Problem Statement:")
body(
    'The core problem addressed by Saṁvāk is the absence of an affordable, real-time, '
    'bidirectional translation system between Indian Sign Language and multiple spoken or '
    'written human languages. Five specific limitations define this problem: (1) Existing '
    'sign language recognition systems are predominantly designed for American Sign Language '
    '(ASL), with no commercially available ISL-specific system for Indian users; (2) Current '
    'tools provide only one-way translation (sign to text), with no mechanism for hearing '
    'users to respond in sign language; (3) All known real-time sign recognition systems '
    'require specialized hardware such as depth cameras, sensor-embedded gloves, or '
    'GPU-accelerated local workstations; (4) No existing tool supports multilingual output '
    'allowing Deaf users to communicate with speakers of regional Indian languages such as '
    'Telugu, Tamil, Kannada, or Malayalam; (5) Available speech-to-sign systems perform '
    'naive English word-to-sign mapping without applying ISL grammar rules, producing sign '
    'sequences that are grammatically incorrect and incomprehensible to fluent ISL users.'
)

subhead("1.3 Objective:")
body(
    '1. Develop a real-time sign language recognition system using only webcam input, '
    'Google MediaPipe Holistic for 258-dimensional landmark extraction, and a '
    'TensorFlow/Keras LSTM model trained on the iSign dataset to recognize 40 ISL classes '
    'with 94.6% top-1 accuracy on 30-frame sequences.\n'
    '2. Implement a complete speech-to-sign translation pipeline that accepts voice or text '
    'input, converts it to grammatically correct ISL gloss sequences using the Google Gemini '
    'API with OSV reordering, and animates a Three.js 3D browser avatar to perform the '
    'corresponding ISL signs in real time.\n'
    '3. Support multilingual input and output across eight languages: English, Hindi, Telugu, '
    'Tamil, Kannada, Malayalam, Spanish, and French.\n'
    '4. Build a web-based, mobile-responsive application requiring no specialized hardware, '
    'with in-browser TF.js inference and MediaPipe JavaScript to eliminate server-round-trip '
    'latency.\n'
    '5. Incorporate an ISL retrieval dictionary using cosine similarity over the 132 MB '
    'isign_retrieval_index.npz embedding of 14,674 iSign video clips across 101 phrase '
    'folders.\n'
    '6. Provide a gamified interactive learning module with XP-based progress tracking, '
    'user authentication (Flask-Login + bcrypt), personalized preferences, and full '
    'translation history logging via SQLAlchemy/SQLite.'
)

subhead("1.4 Scope:")
body(
    'The scope of Saṁvāk encompasses the following included capabilities: real-time '
    'recognition of 40 ISL phrases via webcam using a browser-side TF.js LSTM pipeline; '
    'speech-to-sign conversion with ISL OSV grammar transformation via Gemini API and '
    '3D avatar animation; a fingerspelling recognizer (fingerspell_recognizer.py) for '
    'spelling out unrecognized words; an ISL dictionary backed by the iSign dataset '
    'retrieval index; a gamified learning module; user account management with '
    'authentication, preference storage, and history; and cloud deployment on Render.com '
    'using Gunicorn and Eventlet. The system does not currently cover continuous sign '
    'language recognition for freely signed sentences beyond the 40-class vocabulary, '
    'ISL regional dialects, or offline use without an internet connection for Gemini API '
    'calls.'
)

subhead("1.5 Project Introduction:")
body(
    'Saṁvāk (Sanskrit: सँवाक, meaning "dialogue") is a web-based, real-time, bidirectional '
    'translation application hosted at https://github.com/Naveen8803/Samvak and deployed '
    'on Render.com. The system integrates five major technology domains into a cohesive '
    'web application. In the Sign-to-Speech pipeline: the browser captures webcam frames '
    'at up to 30 FPS; MediaPipe Holistic JavaScript extracts 258-dimensional pose and hand '
    'landmark features (33 pose landmarks × 4 coordinates + 21 left-hand × 3 + 21 '
    'right-hand × 3); 30 consecutive frames are buffered into a (1, 30, 258) tensor; '
    'TF.js loads the tfjs_lstm/ model bundle and performs in-browser classification with '
    'sub-200ms latency against 40 ISL classes; and a confidence threshold filters false '
    'positives. The recognized phrase is then translated to the selected target language '
    'using deep-translator and converted to speech via gTTS. A geometry-based static '
    'classifier in geometry_brain.py provides fallback recognition for basic hand '
    'configurations (thumbs up, V sign, etc.). In the Speech-to-Sign pipeline: voice '
    'input is captured via the Web Speech API or uploaded audio file; speech.py uses '
    'the SpeechRecognition library with Google STT for transcription; spaCy '
    '(en_core_web_sm) tokenizes the English text; grammar_helper.py applies ISL grammar '
    'rules (drop articles, drop auxiliary verbs, reorder to OSV, move WH-words to end); '
    'the Gemini API (gemini-1.5-flash) refines the gloss with context-aware ISL OSV '
    'conversion; and sign.js (1,600+ lines) drives the Three.js r128 3D avatar to '
    'animate the sign sequence. The Flask backend (app.py) uses a factory pattern with '
    'registered blueprints (auth, sign, speech, dictionary, learn), SQLAlchemy ORM with '
    'five database models (User, Translation, UserProgress, UserPreference, '
    'ContactMessage), and Flask-SocketIO for real-time communication.'
)
pb()

# ═══════════════════════════════════════════════════════════════════════════════
# 2. LITERATURE SURVEY
# ═══════════════════════════════════════════════════════════════════════════════
chapter("2.  LITERATURE SURVEY")
blank(1)

lit_papers = [
    {
        "ref": "[1] Wadhawan, A. & Kumar, P., \"Sign Language Recognition Systems: A Decade Systematic Literature Review\", Archives of Computational Methods in Engineering, vol. 28, pp. 785–813, 2021.",
        "objective": "Objective: This comprehensive survey systematically reviews sign language recognition (SLR) research published between 2010 and 2020, classifying 150+ papers into sensor-based and vision-based methodological categories, and evaluating architectures including CNN, RNN, LSTM, and HMM-based approaches in terms of accuracy, dataset size, and real-time feasibility.",
        "methodology": "Methodology: The authors employ a systematic literature review protocol, extracting performance benchmarks, hardware requirements, and dataset characteristics from each work. Sensor-based methods (data gloves, accelerometers) consistently outperform on small datasets while vision-based deep learning methods scale better with larger corpora. The survey identifies that LSTM-based temporal classifiers on pose-estimated landmark sequences achieve the highest accuracy on dynamic gesture datasets.",
        "advantages": "Advantages: The survey provides a rigorous taxonomy distinguishing static gesture recognition (single frame CNN) from dynamic gesture recognition (temporal sequence LSTM), enabling direct comparison. It confirms that pose-estimation features are more robust than raw RGB features to lighting and skin-tone variation.",
        "proposed": "Proposed System Relevance: The survey's finding that LSTM classifiers on pose-landmark sequences achieve state-of-the-art SLR accuracy directly motivated Saṁvāk's choice of a TensorFlow/Keras LSTM model operating on 258-dimensional MediaPipe Holistic pose features extracted across 30-frame sliding windows, trained on the iSign ISL dataset.",
    },
    {
        "ref": "[2] Lugaresi, C. et al., \"MediaPipe: A Framework for Building Perception Pipelines\", arXiv:1906.08172, 2019.",
        "objective": "Objective: This paper presents MediaPipe, an open-source, cross-platform, production-ready framework for building real-time multimedia perception pipelines for mobile, desktop, and edge devices. It supports graph-defined pipelines for tasks including face detection, hand tracking, pose estimation, and holistic body landmark detection.",
        "methodology": "Methodology: MediaPipe defines perception as a directed acyclic graph of calculators (processing nodes). The Holistic pipeline applies BlazePose for 33 full-body pose landmarks, Palm Detection and Hand Landmark models for 21 keypoints per hand, and Face Mesh for 468 face landmarks — all running concurrently at 30+ FPS on mobile hardware without GPU acceleration. Saṁvāk uses MediaPipe Holistic to extract 33 pose × 4 + 21 left-hand × 3 + 21 right-hand × 3 = 258 features per frame.",
        "advantages": "Advantages: MediaPipe provides a JavaScript API enabling all landmark extraction to run client-side in the browser, eliminating the need for server-round-trip latency for each video frame. The framework is hardware-agnostic and runs at real-time speeds on consumer webcams.",
        "proposed": "Proposed System Relevance: Saṁvāk's sign.js uses MediaPipe Holistic v0.5 JavaScript API to extract 258-dimensional landmark vectors from every webcam frame in the browser. These vectors feed the TF.js LSTM buffer enabling fully client-side sign recognition without any server-side ML inference.",
    },
    {
        "ref": "[3] Hochreiter, S. & Schmidhuber, J., \"Long Short-Term Memory\", Neural Computation, vol. 9, no. 8, pp. 1735–1780, 1997.",
        "objective": "Objective: This foundational paper introduces the Long Short-Term Memory (LSTM) architecture to solve the vanishing gradient problem in recurrent neural networks, enabling learning of long-range temporal dependencies in sequential data.",
        "methodology": "Methodology: LSTM introduces three gating mechanisms — input gate, forget gate, and output gate — along with a cell state that carries information across many time steps. Gradient flow is stabilized through constant error carousels within memory cells. The architecture is validated on tasks requiring memory over hundreds of time steps, including speech recognition and handwriting sequencing.",
        "advantages": "Advantages: LSTM's ability to model temporal dependencies across sequences of arbitrary length makes it ideal for dynamic gesture recognition where the same sign may span varying numbers of frames depending on signer speed. The architecture naturally handles the 30-frame temporal window used in Saṁvāk.",
        "proposed": "Proposed System Relevance: Saṁvāk's sign_language.h5 LSTM model (trained in train_classifier.py) processes 30-frame × 258-feature sequences to classify 40 ISL gesture classes. The LSTM layers capture temporal hand trajectory patterns — such as the outward arc of \"Thank You\" or the wrist-rotation of \"Hello\" — that cannot be captured by single-frame CNN classifiers.",
    },
    {
        "ref": "[4] Vasudevan, A. et al. (Exploration-Lab, IIT Kharagpur), \"iSign: A Benchmark for Indian Sign Language Processing\", Proceedings of ACL 2024.",
        "objective": "Objective: iSign is the first large-scale benchmark dataset for Indian Sign Language processing, providing 14,674 video clips across 101 ISL phrase categories with associated RGB video, extracted pose landmarks, and corresponding English text annotations, released alongside baseline sign recognition and retrieval models.",
        "methodology": "Methodology: ISL phrases were recorded by multiple native ISL signers under controlled studio conditions. Landmark features are pre-extracted using MediaPipe Holistic and stored as normalized numpy arrays per clip. A phrase-level retrieval index using sentence embedding cosine similarity enables nearest-neighbor lookup for unknown phrase queries. Baseline LSTM classifiers trained on the 101-class subset achieve 78% top-1 accuracy, setting a benchmark for the community.",
        "advantages": "Advantages: iSign uniquely provides pre-extracted MediaPipe landmark arrays, enabling direct use in LSTM training without re-processing raw video. The cosine similarity retrieval index supports open-vocabulary ISL recognition beyond the fixed training classes.",
        "proposed": "Proposed System Relevance: Saṁvāk is trained on 40 of the 101 iSign phrase classes using the pre-extracted landmark arrays, achieving 94.6% top-1 accuracy (exceeding the iSign baseline). The 132 MB isign_retrieval_index.npz forms the backbone of Saṁvāk's dictionary and retrieval fallback when LSTM confidence is below threshold. The iSign dataset of 14,674 clips from 101 phrase folders is the foundation of Saṁvāk's recognition capability.",
    },
    {
        "ref": "[5] Saunders, B., Camgoz, N.C. & Bowden, R., \"Progressive Transformers for End-to-End Sign Language Production\", ECCV 2020.",
        "objective": "Objective: This paper proposes a Transformer-based neural architecture for sign language production (SLP) — the task of converting spoken language sentences into continuous sign pose sequences, enabling automated generation of sign language from text.",
        "methodology": "Methodology: Progressive Transformers use a counter-decoding mechanism that progressively generates sign pose sequences token by token, conditioning each output frame on both the accumulated sign pose history and the source text encoding. The model is trained on the PHOENIX14T dataset (German Sign Language weather broadcasts) with DTW-based evaluation metrics specifically designed for sign pose sequence quality.",
        "advantages": "Advantages: The progressive generation mechanism produces temporally smooth sign sequences without requiring explicit gloss-to-pose alignment annotations, demonstrating that end-to-end text-to-sign translation is feasible with transformer architectures at scale.",
        "proposed": "Proposed System Relevance: This work directly informed Saṁvāk's speech-to-sign pipeline design philosophy — converting spoken English first to an intermediate ISL gloss representation (via Gemini API) and then to avatar animation tokens. While Saṁvāk uses pre-defined sign animations rather than neural pose generation, the gloss-as-intermediate-representation approach is directly inspired by progressive transformer methods.",
    },
    {
        "ref": "[6] Rastgoo, R., Kiani, K. & Escalera, S., \"Sign Language Recognition: A Deep Survey\", Expert Systems with Applications, vol. 164, 2021.",
        "objective": "Objective: This survey provides a systematic review of deep learning-based SLR methods covering 200+ papers, categorizing architectures by input modality (RGB, depth, skeleton, multi-modal), temporal modeling approach (CNN, LSTM, GRU, Transformer, 3D-CNN), and benchmark dataset performance.",
        "methodology": "Methodology: The survey identifies five input modalities: RGB video, depth maps (Kinect), optical flow, skeleton/pose sequences, and multi-modal combinations. Skeleton/pose-based approaches using LSTM or GCN (Graph Convolutional Network) on MediaPipe or OpenPose landmarks consistently outperform raw RGB approaches in cross-subject generalization, being invariant to skin color, clothing, and lighting conditions.",
        "advantages": "Advantages: The survey's comprehensive comparison definitively establishes that pose-based features extracted by lightweight detectors (MediaPipe over OpenPose) provide the optimal balance of accuracy, computational cost, and cross-demographic generalization for production SLR systems.",
        "proposed": "Proposed System Relevance: Saṁvāk adopts the survey's recommendation of using MediaPipe skeleton/pose features as the primary input modality, deliberately avoiding raw RGB or optical flow inputs that would require GPU-accelerated processing and fail to generalize across different signers' skin tones and backgrounds.",
    },
    {
        "ref": "[7] Moryossef, A. et al., \"Real-Time Sign Language Detection Using Human Pose Estimation\", Proceedings of the 1st International Workshop on Sign Language Translation and Production (SLTP), ACL 2020.",
        "objective": "Objective: This paper proposes a lightweight real-time sign language detection system that distinguishes signing segments from non-signing in continuous video streams using only 2D body pose keypoints, enabling frame-level binary segmentation for downstream recognition.",
        "methodology": "Methodology: A compact binary classifier operates on the velocity and acceleration of wrist and hand keypoints extracted by OpenPose, achieving 91% detection accuracy at 28 FPS on standard CPUs. The method demonstrates that sign language activity detection requires only hand-motion features, not full visual appearance, making it deployable on edge devices.",
        "advantages": "Advantages: The approach demonstrates that high-quality sign detection requires only kinematic features (keypoint velocities) rather than full image processing, validating the MediaPipe-only feature extraction strategy as both accurate and computationally efficient.",
        "proposed": "Proposed System Relevance: Saṁvāk's sign.js implements a 30-frame buffer with a minimum-motion pre-filter that discards near-static frames before classification, inspired by this work. The pre-filter prevents false positives from background movement and ensures that only active signing segments are fed to the TF.js LSTM classifier.",
    },
    {
        "ref": "[8] Özdemir, O., Kindiroglu, A.A. & Akarun, L., \"Isolated Sign Language Recognition with Multi-Scale Features Using LSTM\", IEEE Signal Processing Letters, 2020.",
        "objective": "Objective: This paper addresses isolated ISL recognition (specifically Turkish Sign Language) using a multi-scale LSTM architecture that simultaneously processes hand shape, hand motion, and body pose features at different temporal resolutions, achieving state-of-the-art recognition on the BosphorusSign dataset.",
        "methodology": "Methodology: Three parallel LSTM branches process: (1) hand shape features (21 MediaPipe keypoints per hand per frame); (2) hand motion features (frame-to-frame displacement of palm center); and (3) body pose features (upper body OpenPose keypoints). Branch outputs are concatenated and classified by a final dense layer. Training uses data augmentation (temporal jitter, spatial perturbation) to improve robustness.",
        "advantages": "Advantages: The multi-scale architecture achieves 93.8% accuracy on a 226-class sign dataset, demonstrating that combining pose, motion, and shape features outperforms single-modality approaches. The temporal jitter augmentation technique significantly reduces overfitting on small sign datasets.",
        "proposed": "Proposed System Relevance: Saṁvāk's 258-dimensional MediaPipe Holistic feature vector implicitly captures all three modalities — hand shape (42 hand keypoints × 3D), body pose (33 landmarks × 4D), and motion implicitly across the 30-frame sequence — aligning with this multi-scale philosophy in a unified single-LSTM architecture rather than three parallel branches, achieving 94.6% accuracy on 40 ISL classes.",
    },
]

for paper in lit_papers:
    p_ref = doc.add_paragraph()
    p_ref.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p_ref.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_before = Pt(12)
    pf.space_after = Pt(0)
    r_ref = p_ref.add_run(paper["ref"])
    _set_font(r_ref, 12, bold=True)

    for key in ["objective", "methodology", "advantages", "proposed"]:
        body(paper[key])

pb()

# ═══════════════════════════════════════════════════════════════════════════════
# 3. SYSTEM ANALYSIS
# ═══════════════════════════════════════════════════════════════════════════════
chapter("3.  SYSTEM ANALYSIS")

subhead("3.1 Existing System:")
body(
    'Four categories of existing systems address sign language recognition and translation, '
    'each with significant limitations that Saṁvāk addresses. First, hardware-based systems '
    'such as CyberGlove III (Immersion Corp.) and P5 Data Glove use resistive bend sensors '
    'on each finger to measure joint angles, providing accurate joint telemetry but costing '
    '$2,000–$5,000 per unit and requiring physical wiring that restricts natural signing. '
    'Second, depth-camera systems based on Microsoft Kinect v2 capture full skeletal pose '
    'and depth maps but require the signer to remain within a 0.5–4.5m range, cost $250–$400, '
    'and were discontinued with no active support. Third, desktop software ASL applications '
    '(HandTalk, SignAll) use RGB cameras and machine learning models trained exclusively on '
    'American Sign Language with English-only output and no speech-to-sign reversal. Fourth, '
    'static image classifiers published in academic repositories recognize isolated ASL '
    'letter shapes (fingerspelling) from single frames using CNN, lacking temporal modeling '
    'needed for dynamic phrase recognition and providing no integration with speech pipelines.'
)

subhead("3.2 Disadvantages:")
body(
    '1. Hardware Dependency: All production-quality systems require sensor gloves ($2,000–$5,000) '
    'or discontinued depth cameras (Kinect), making them inaccessible to the general public.\n'
    '2. ASL-Only Vocabulary: No commercially available system trains on Indian Sign Language; '
    'ISL has entirely different signs, grammar (OSV order), and handshapes from ASL.\n'
    '3. Unidirectional Translation: All existing systems perform only sign-to-text; no system '
    'allows a hearing user to generate ISL signs from speech, preventing true conversation.\n'
    '4. Single-Language Output: Output is universally English-only, excluding Telugu, Tamil, '
    'Kannada, Malayalam, Hindi, Spanish, and French speakers from participation.\n'
    '5. No ISL Grammar Engine: Available speech-to-sign tools perform naive English word-to-sign '
    'mapping, generating grammatically incorrect sign sequences that ISL users cannot understand.\n'
    '6. No Learning Component: Existing systems focus solely on translation and provide no '
    'mechanism for hearing users to learn ISL vocabulary or grammar.\n'
    '7. No Cloud Accessibility: All known systems require local installation on Windows desktops '
    'with GPU acceleration; no web-based, cross-platform, mobile-accessible deployment exists.'
)

subhead("3.3 Proposed System:")
body(
    'Saṁvāk is a web-based, bidirectional, hardware-free Indian Sign Language translation '
    'system that resolves all seven limitations of existing tools. Using only a standard '
    'consumer webcam and microphone accessible through any modern web browser, Saṁvāk enables: '
    'real-time ISL recognition via MediaPipe + TF.js LSTM (40 classes, 94.6% accuracy, 30-frame '
    'sequences, 258 features) running entirely client-side; grammatically correct speech-to-ISL '
    'conversion via Gemini API with OSV gloss generation and Three.js 3D avatar animation; '
    'multilingual output in 8 languages via deep-translator and gTTS; interactive ISL dictionary '
    'powered by the 14,674-clip iSign retrieval index (132 MB cosine similarity index); gamified '
    'XP-based learning; and full user account management with Flask-Login authentication and '
    'SQLite history logging. The application is deployed on Render.com and is accessible from '
    'any desktop or mobile device with an internet connection.'
)

subhead("3.4 Advantages:")
body(
    '1. Zero Hardware Cost: Runs on any device with a standard webcam — no gloves, depth cameras, '
    'or GPU required.\n'
    '2. Real-Time In-Browser Inference: TF.js LSTM classification at sub-200ms per 30-frame window '
    'entirely in the browser using WebGL acceleration.\n'
    '3. Bidirectional Communication: Both sign-to-speech and speech-to-sign enable full conversations '
    'between Deaf and hearing users.\n'
    '4. Grammatically Correct ISL Gloss: Gemini API generates OSV-ordered ISL gloss sequences with '
    'proper article dropping, auxiliary elimination, and WH-word repositioning.\n'
    '5. Multilingual Support: Output available in 8 languages covering 90%+ of Indian spoken languages.\n'
    '6. ISL-Specific Training Data: Model trained on iSign (ACL 2024), the only publicly available '
    'large-scale ISL dataset with 14,674 clips across 101 phrases.\n'
    '7. Cloud Deployment: Accessible from any browser via Render.com with no installation.\n'
    '8. Gamified Learning: XP-based learning module encourages hearing users to learn ISL, expanding '
    'the community of fluent signers.'
)

subhead("3.5 Work Flow:")
body('Sign-to-Speech Pipeline (8 steps):')
body(
    '1. User positions hands in webcam frame and performs an ISL gesture.\n'
    '2. Browser captures video frames via HTML5 <video> element at up to 30 FPS.\n'
    '3. MediaPipe Holistic JavaScript API extracts 258-dimensional landmark vector from each frame '
    '(33 pose × 4 + 21 left-hand × 3 + 21 right-hand × 3).\n'
    '4. sign.js buffers 30 consecutive landmark vectors into a (1, 30, 258) TF.js tensor.\n'
    '5. TF.js LSTM model (tfjs_lstm/ bundle) performs in-browser inference; argmax of softmax output '
    'selects the predicted ISL class from 40 labels.\n'
    '6. If confidence exceeds threshold, the predicted phrase is sent via WebSocket to Flask backend.\n'
    '7. Flask translates phrase using deep-translator to the user-selected target language.\n'
    '8. gTTS synthesizes audio; browser plays the mp3 file and displays the translated text.'
)
blank(1)
body('Speech-to-Sign Pipeline (7 steps):')
body(
    '1. User speaks into the device microphone; browser captures audio via Web Speech API or '
    'uploads an audio file to Flask /speech/upload endpoint.\n'
    '2. speech.py receives the audio, converts to WAV using pydub, and transcribes via '
    'SpeechRecognition with Google STT (language: en-IN).\n'
    '3. Transcribed English text is passed to grammar_helper.py, which applies preliminary ISL '
    'grammar rules: drop articles and auxiliary verbs, expand contractions, identify time markers.\n'
    '4. The preprocessed text is sent to the Gemini API (gemini-1.5-flash) with a specialized '
    'ISL OSV grammar prompt; the API returns a JSON array of ordered ISL gloss tokens.\n'
    '5. Flask returns the gloss token array to sign.js via HTTP response.\n'
    '6. sign.js interprets each gloss token: known tokens trigger Three.js 3D avatar animations '
    'using pre-defined joint-angle keyframes; unknown tokens trigger fingerspelling animation '
    'via fingerspell_recognizer.py character-by-character signing.\n'
    '7. The 3D avatar (Three.js r128) renders the complete ISL sign sequence in the browser WebGL canvas.'
)
pb()

# ═══════════════════════════════════════════════════════════════════════════════
# 4. REQUIREMENT ANALYSIS
# ═══════════════════════════════════════════════════════════════════════════════
chapter("4.  REQUIREMENT ANALYSIS")

subhead("4.1 Functional and Non-Functional Requirements:")
body('Functional Requirements (10 items):', space_before=6)
body(
    '1. The system shall capture live video from the user\'s webcam and detect hand and pose '
    'landmarks in real time using the MediaPipe Holistic JavaScript API.\n'
    '2. The system shall extract 258-dimensional feature vectors per frame and classify ISL '
    'gestures into 40 predefined phrases using the TF.js LSTM model with ≥90% top-1 accuracy.\n'
    '3. The system shall convert recognized ISL phrases to speech in the user-selected language '
    'using gTTS with Web Speech API as fallback.\n'
    '4. The system shall accept voice input via microphone or uploaded audio file, transcribe '
    'it using Google Speech-to-Text (SpeechRecognition library), and display the transcription.\n'
    '5. The system shall convert transcribed text to grammatically correct ISL gloss sequences '
    'using the Gemini API with OSV reordering and animate the result via the 3D avatar.\n'
    '6. The system shall support multilingual text output and TTS in English, Hindi, Telugu, '
    'Tamil, Kannada, Malayalam, Spanish, and French.\n'
    '7. The system shall provide user registration, login, session management, and secure '
    'logout using Flask-Login and bcrypt password hashing.\n'
    '8. The system shall maintain a complete translation history for authenticated users, '
    'storing each session\'s recognized phrases, timestamps, and target languages in SQLite.\n'
    '9. The system shall provide an ISL dictionary allowing users to search and browse sign '
    'video clips from the iSign dataset using the cosine similarity retrieval index.\n'
    '10. The system shall implement a gamified ISL learning module with XP-based progress '
    'tracking, level advancement thresholds, and a user leaderboard dashboard.'
)
body('Non-Functional Requirements (6 items):', space_before=6)
body(
    '1. Performance: Full sign-to-speech pipeline (frame capture → landmark extraction → '
    'LSTM inference → translation) shall complete within 200ms per 30-frame window.\n'
    '2. Security: User passwords shall be hashed using bcrypt with a salt factor of 12; '
    'Flask sessions shall use signed cookies with a SECRET_KEY; all forms shall be CSRF-protected.\n'
    '3. Portability: The frontend shall render correctly on Chrome 90+, Firefox 88+, Safari 14+, '
    'and Edge 90+ on desktop and mobile (iOS Safari, Android Chrome) without any plugins.\n'
    '4. Reliability: The system shall gracefully handle camera unavailability, microphone '
    'permission denial, API timeouts, and network disconnection with user-facing error messages.\n'
    '5. Maintainability: The backend shall follow a modular Flask blueprint architecture '
    'with separate blueprints for authentication, sign recognition, speech processing, '
    'dictionary, and learning features.\n'
    '6. Scalability: The Render.com deployment shall support horizontal scaling via Gunicorn '
    'workers and Eventlet async workers for WebSocket connections.'
)

subhead("4.2 Hardware Requirements:")
make_table(
    ["S.No.", "Component", "Minimum Specification"],
    [
        ("1", "Processor", "Intel Core i3 (8th Gen) / AMD Ryzen 3 3200U or equivalent"),
        ("2", "RAM", "8 GB DDR4 minimum (16 GB recommended for development)"),
        ("3", "Hard Disk", "256 GB SSD (model files: sign_language.h5 = 1.5 MB, retrieval index = 132 MB)"),
        ("4", "Webcam", "720p (1280×720) minimum, 30 FPS — standard integrated or USB webcam"),
        ("5", "Microphone", "Any standard built-in or external microphone (8 kHz+ sample rate)"),
        ("6", "Monitor", "1366×768 minimum resolution; 16:9 aspect ratio recommended"),
        ("7", "Network", "Broadband Internet connection (≥5 Mbps) for Gemini API and gTTS calls"),
    ],
    col_widths=[0.6, 1.8, 4.2]
)

subhead("4.3 Software Requirements:")
make_table(
    ["S.No.", "Component", "Version / Specification"],
    [
        ("1", "Operating System", "Windows 10/11, macOS 12+, or Ubuntu 20.04+"),
        ("2", "Programming Language", "Python 3.10+ (backend); JavaScript ES2020+ (frontend)"),
        ("3", "Web Framework", "Flask 3.x with Jinja2 templating and Werkzeug WSGI"),
        ("4", "Frontend", "HTML5, Vanilla CSS3, JavaScript ES6+ (no React/Vue)"),
        ("5", "ML Framework (Server)", "TensorFlow 2.15+ / Keras; model: sign_language.h5"),
        ("6", "ML Framework (Browser)", "TensorFlow.js 4.x; model bundle: tfjs_lstm/"),
        ("7", "Pose Estimation (Browser)", "Google MediaPipe Holistic JavaScript 0.5+ (CDN)"),
        ("8", "NLP", "spaCy 3.x with en_core_web_sm model pipeline"),
        ("9", "Generative AI", "Google Gemini API (gemini-1.5-flash) via google-generativeai"),
        ("10", "Speech Recognition", "SpeechRecognition 3.10+ with Google STT backend"),
        ("11", "Text-to-Speech", "gTTS (Google TTS) 2.x; Web Speech API fallback"),
        ("12", "3D Rendering", "Three.js r128 (CDN) with WebGL renderer"),
        ("13", "Database", "SQLite 3 via SQLAlchemy 2.x ORM; Flask-Login 0.6+"),
        ("14", "Real-Time", "Flask-SocketIO 5.x + Eventlet 0.33+ for WebSocket"),
    ],
    col_widths=[0.6, 2.2, 3.8]
)

subhead("4.4 Architecture:")
body(
    'Saṁvāk follows a hybrid client-server architecture that deliberately offloads all '
    'computationally intensive ML inference to the client browser using TF.js and MediaPipe '
    'JavaScript, reserving the Flask server for authentication, database operations, grammar '
    'processing, and API coordination. The client-side layer handles: webcam capture and '
    'streaming via HTML5 Video API; MediaPipe Holistic landmark extraction (258D per frame); '
    '30-frame tensor construction and TF.js LSTM classification; Three.js 3D avatar animation; '
    'and Web Speech API TTS fallback. The server-side Flask layer provides: user management '
    'via Flask-Login + bcrypt; ISL grammar processing via grammar_helper.py and Gemini API; '
    'cosine similarity retrieval via isign_retrieval.py against the 132 MB isign_retrieval_'
    'index.npz; gTTS audio synthesis; SQLAlchemy ORM for User, Translation, UserProgress, '
    'UserPreference, and ContactMessage models; and Flask-SocketIO WebSocket for real-time '
    'frame prediction callbacks. The deployment tier uses Render.com free-tier cloud with '
    'Gunicorn WSGI server and Eventlet async workers.'
)
img_arch = os.path.join(DIAGRAMS, "deployment_diagram.png")
insert_image(img_arch, width=6.0, caption="Figure 4.1: System Architecture — Client-Server Deployment on Render.com")
pb()

print("Chapters 1-4 done.")

# ── Save checkpoint ──
doc.save(OUT)
print(f"Checkpoint saved: {OUT}")
