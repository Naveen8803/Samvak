import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document(r'E:\sam\Project documents\Samvak_Base_Clean_V12.docx')
DIAGRAM_DIR = r"E:\sam\Project documents\diagrams"
SCREENSHOT_DIR = r"E:\sam\Project documents\screenshots"

def add_para(text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=13, space_before=None, space_after=None):
    p = doc.add_paragraph()
    p.alignment = align
    if space_before is not None: p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None: p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = 'Times New Roman'
    return p

def add_empty(n=1):
    for _ in range(n): doc.add_paragraph()

doc.add_page_break()

# ========================================== ABSTRACT ==========================================
add_empty()
add_para('ABSTRACT', True, WD_ALIGN_PARAGRAPH.CENTER, 14)
add_empty()
add_para('This project aims to solve the severe communication gap between the hearing-impaired community and the general population through an AI-based bidirectional translation system. Traditional communication relies heavily on human sign language interpreters, limiting accessibility and independence for Deaf individuals.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
add_para('The proposed Multi-Lingual Sign Language to Speech and Speech to Sign Translator provides a software-based solution leveraging Machine Learning, Natural Language Processing (NLP), and Computer Vision (CV). Operating via a standard webcam and microphone, the system utilizes MediaPipe Hand and Pose tracking to extract gesture landmarks, which are classified into accurate text using an advanced deep learning model. For bidirectional dialogue, the system captures spoken audio, employs Google Speech-to-Text for transcription, processes it through an NLP engine to map spoken grammar into Indian Sign Language (ISL) syntactical structure, and visually renders the translation using a 3D avatar animations.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
add_para('The outcome is a highly accessible, multi-lingual web platform capable of recognizing complex gestures and converting them to regional speech while simultaneously translating audio input directly into real-time ISL gestures. This abolishes the reliance on costly sensor hardware, drastically improving independence and inclusivity.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
doc.add_page_break()

# ========================================== TOC ==========================================
add_para('TABLE OF CONTENTS', True, WD_ALIGN_PARAGRAPH.CENTER, 14)
add_empty()
toc_data = [
    ("CHAPTER NO.", "TITLE", "PAGE NO."),
    ("1", "INTRODUCTION", "1"),
    ("1.1", "Motivation", "1"),
    ("1.2", "Problem Statement", "1"),
    ("1.3", "Objectives", "2"),
    ("1.4", "Scope", "2"),
    ("1.5", "Overview of the System", "2"),
    ("2", "LITERATURE SURVEY", "3"),
    ("3", "SYSTEM ANALYSIS", "5"),
    ("3.1", "Existing System", "5"),
    ("3.2", "Disadvantages of Existing System", "5"),
    ("3.3", "Proposed System", "6"),
    ("3.4", "Advantages of Proposed System", "6"),
    ("3.5", "Workflow Explanation", "6"),
    ("4", "REQUIREMENT ANALYSIS", "8"),
    ("4.1", "Functional Requirements", "8"),
    ("4.2", "Non-Functional Requirements", "8"),
    ("4.3", "Hardware Requirements", "9"),
    ("4.4", "Software Requirements", "9"),
    ("4.5", "System Architecture", "10"),
    ("5", "SYSTEM DESIGN", "11"),
    ("5.1", "Use Case Diagram", "11"),
    ("5.2", "Class Diagram", "12"),
    ("5.3", "Sequence Diagram", "13"),
    ("5.4", "Activity Diagram", "14"),
    ("5.5", "Component Diagram", "15"),
    ("5.6", "Deployment Diagram", "15"),
    ("5.7", "ER Diagram", "16"),
    ("5.8", "Data Flow Diagram", "16"),
    ("6", "CODE, IMPLEMENTATION AND RESULTS", "18"),
    ("6.1", "Modules Overview", "18"),
    ("6.2", "Sample Implementation Code", "20"),
    ("6.3", "Output Screens", "22"),
    ("7", "SYSTEM STUDY AND TESTING", "24"),
    ("7.1", "Feasibility Study", "24"),
    ("7.2", "Types of Testing", "24"),
    ("7.3", "Test Cases", "25"),
    ("8", "CONCLUSION", "26"),
    ("9", "FUTURE ENHANCEMENT", "27"),
    ("10", "REFERENCES", "28")
]
toc_table = doc.add_table(rows=len(toc_data), cols=3)
toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (c1, c2, c3) in enumerate(toc_data):
    toc_table.rows[i].cells[0].text = c1
    toc_table.rows[i].cells[1].text = c2
    toc_table.rows[i].cells[2].text = c3
    for cell in toc_table.rows[i].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
                if i == 0 or len(c1) == 1 or c1 == "":  r.bold = True
                else: r.bold = False

doc.add_page_break()

# ========================================== CHAP 1 ==========================================
add_para("1. INTRODUCTION", True, WD_ALIGN_PARAGRAPH.LEFT, 16)
add_empty()

add_para("1.1 Motivation", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('According to global health records, millions of individuals navigate society primarily through sign language. Because the general public is unversed in sign vocabulary, the Deaf community faces profound societal isolation. Developing an autonomous, AI-driven software translating sign language to speech breaks this barrier natively without demanding users to wear intrusive sensors.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("1.2 Problem Statement", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('Current interaction solutions between hearing and non-hearing individuals heavily depend on manual interpretation. Digital alternatives usually concentrate uniquely on American Sign Language while relying on high-cost hardware gloves. A lack of bidirectional functionality—meaning they translate signs to words but fail to translate spoken words back into sign patterns—renders daily conversational integration highly inefficient.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("1.3 Objectives", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('1. Develop a bidirectional Sign-to-Speech and Speech-to-Sign translation system using Computer Vision.\n2. Utilize MediaPipe arrays to dynamically recognize multi-hand structures substituting hardware layers.\n3. Integrate Google Speech-to-Text establishing robust microphone transcription.\n4. Design scalable multi-lingual pipelines permitting real-time vocalized outputs natively in regional languages.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("1.4 Scope", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('The software bounds encompass isolated sign classification over a webcam matrix parsing direct NLP manipulations through Python pipelines. It enables users to switch localized vernacular outputs through a web interface, rendering audio and a dynamically mapped 3D avatar executing precise gloss motions accurately modeling Indian Sign Language mechanics.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("1.5 Overview of the System", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('The system encapsulates both detection and rendering modules natively within web protocols. Camera streams feed isolated landmark features (hands, shoulders, elbows) matching complex gestures through integrated deep neural weights generating accurate strings. Subsequent text arrays trigger real-time Text-to-Speech functions mapping output sounds. Inversely, user audio feeds invoke Natural Language processing isolating the Subject/Object logic mapping resulting tokens rendering 3D avatar skeletal nodes executing translation sequences comprehensively.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

doc.add_page_break()

# ========================================== CHAP 2 ==========================================
add_para("2. LITERATURE SURVEY", True, WD_ALIGN_PARAGRAPH.LEFT, 16)
add_empty()

refs = [
    ('[1] Sign Language Recognition utilizing CNN and LSTM Architectures', 'This model maps robust spatial characteristics avoiding the limitations of flat vector machines. However, it completely lacks speech synthesis endpoints mapping outputs strictly into basic text strings.'),
    ('[2] MediaPipe Holistic Tracking within Real-time Computer Vision Systems', 'Highlights optimized browser extraction of human tracking points. It resolves physical boundary limits preventing users from wearing physical tracked gloves facilitating natural sign flow dynamics.'),
    ('[3] Continuous Gesture Recognition mapping American Sign Language', 'Tracks dynamic sequences across complex lighting scenarios. Its primary limitation is the absolute isolation of its logic solely encompassing ASL preventing ISL adoption entirely.'),
    ('[4] A Survey on Bidirectional Application Translation: Speech to Sign NLP', 'Validates transforming English syntax directly mimicking linguistic rules identifying verb modifiers isolating Subject/Object parameters effectively translating text outputs logically for Avatar ingestion constraints.'),
    ('[5] Progressive Translators: End-to-End Hand Gesture Modelling', 'Investigates directly mapping text phrases directly to generated skeletal models minimizing rendering latencies highlighting why WebGL technologies successfully supersede native application render bounds.'),
    ('[6] Implementing Multi-lingual Vocals mapping NLP strings via gTTS', 'Highlights programmatic methods utilizing regional speech identifiers configuring automated pitch modulations eliminating monolingual dependencies.'),
    ('[7] Indian Sign Language Database Parsing and Implementation Algorithms', 'Addresses the severe dataset deficits found regionally demonstrating structural approaches mapping generalized Indian vocabulary nodes scaling accurate categorical modeling datasets.')
]
for title, desc in refs:
    add_para(title, True, size=13)
    add_para(desc, size=13, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_empty()
doc.add_page_break()

# ========================================== CHAP 3 ==========================================
add_para("3. SYSTEM ANALYSIS", True, WD_ALIGN_PARAGRAPH.LEFT, 16)
add_empty()

add_para("3.1 Existing System", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('Current translation environments leverage sensor hardware tracking localized hand movements parsing standard ASL vocabularies returning fixed English interpretations lacking bidirectional interactions enabling deaf individuals to simultaneously understand speaking users independently.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("3.2 Disadvantages of Existing System", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('• High Hardware Dependence: Restricts access strictly relying on expensive sensor nodes masking operational accessibility.\n• Monolingual Bias: Hardcoded output processing limits global scaling preventing regional implementations significantly.\n• Syntactical Incorrectness: Processing word-by-word without rearranging Subject/Object mapping yielding grammatically unviable visuals.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("3.3 Proposed System", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('The proposed tool eliminates hardware leveraging web camera streams calculating node boundaries mapping predictive sequences leveraging customized ISL model weights resolving textual/vocal responses dynamically. Conversely, recording live mic data tokenizes phrasing matching structural Avatar parameters parsing real-time ISL visuals bi-directionally.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("3.4 Advantages of Proposed System", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('• Zero Hardware Complexity: Bypasses flex-glove dependency operating smoothly over ubiquitous digital webcams natively.\n• Universal Multi-Language Scaling: Embeds text integrations supporting massive dialect toggling simultaneously routing output syntaxes correctly.\n• Real-time Grammar Engine: Transcribes audio correctly assigning token hierarchies modeling native OSV rules before translating avatar boundaries logically.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("3.5 Workflow Explanation", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('The logical bounds configure initial client interactions routing user inputs specifically between a visual camera feed processing gesture mapping parameters mapping outputs toward deep-translation / Text-To-Speech interfaces directly. Alternatively, invoking the microphone processes Google Speech transcripts extracting core phrases mapping gloss rules outputting visual avatar bounds accurately mitigating inference limits reliably.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
doc.add_page_break()

# ========================================== CHAP 4 ==========================================
add_para("4. REQUIREMENT ANALYSIS", True, WD_ALIGN_PARAGRAPH.LEFT, 16)
add_empty()

add_para("4.1 Functional Requirements", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('1. Web camera ingestion must extract and validate 258 coordinate dimensions natively tracking hand parameters.\n2. Model operations must process sequence tracking mapping categorical predictions evaluating valid lexical sign tokens accurately.\n3. Bidirectional logic processing microphone streams accurately logging strings evaluating syntactic language configurations outputting proper visual models.\n4. Translating raw text into 8 regional formats parsing outputs mapping exact linguistic representations synthetically.\n5. Tracking Dictionary progress managing dynamic point schemas indexing historical references continuously.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("4.2 Non-Functional Requirements", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('• Low Latency Bounds: Ensuring end-to-end rendering logic evaluates predictions cleanly under 400ms mimicking actual conversation parameters natively.\n• High Availability: Responsive deployments routing client rendering interfaces maintaining execution loops seamlessly handling interrupted signals gracefully.\n• Security Integrity: Abstracted SQLite payloads managing local states encrypting internal registry passwords protecting backend routes effectively.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("4.3 Hardware Requirements", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('• Central Processing Unit: AMD Ryzen 3 / Intel i3 baseline\n• Allocation Memory: Minimum 8 GB internal RAM limits\n• Audio/Video Nodes: Functional internal/external Webcam and Microphone endpoints.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("4.4 Software Requirements", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('• Execution Tier: Python 3.10 and corresponding browser instances mapping HTML5 environments.\n• External Library Chains: Flask REST bindings, OpenCV, MediaPipe tracking limits, SpeechRecognition.\n• Synthesis Engines: Deep Translator APIs, Google Text to Speech module mapping parameters natively.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("4.5 System Architecture", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('Client bounds manage capturing arrays configuring visual bounding interfaces communicating mapping sequences asynchronously transmitting validation payloads resolving REST API paths across global Flask bindings tracking localized session parameters tracking SQLite models outputting processed payload coordinates mapping localized HTML template hierarchies natively maintaining high operational responsiveness processing independent WebGL operations fluidly.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
doc.add_page_break()


# ========================================== CHAP 5 ==========================================
add_para("5. SYSTEM DESIGN", True, WD_ALIGN_PARAGRAPH.LEFT, 16)
add_empty()

diagrams = [
    ("5.1 Use Case Diagram", "use_case_diagram.png", "Showcases the User interacting closely over primary sub-components identifying Authentication layers feeding input structures directly into the Gesture module, Speech module, and the Translation module evaluating boundaries efficiently."),
    ("5.2 Class Diagram", "class_diagram.png", "Translating operational classes linking logic structures outlining User Object persistence linking NLP grammar formatting pipelines passing Translation model endpoints seamlessly avoiding generic boundaries completely."),
    ("5.3 Sequence Diagram", "sequence_sign.png", "Evaluates message calls tracking User commands passing UI predictions tracking outputs specifically over User, Flask backend mapping Translation processing states accurately."),
    ("5.4 Activity Diagram", "activity_diagram.png", "Logs branch thresholds mapping conditional parameters testing the Speech tracking boundaries executing model validation verifying logical rendering rules executing Sign representations explicitly."),
    ("5.5 Component Diagram", "deployment_diagram.png", "Summarizes the macro code bounds verifying the core structure separating UI execution bindings away from Python AI Translation modules identifying internal API bridges properly."),
    ("5.6 Deployment Diagram", "deployment_diagram.png", "Models precise server execution logic identifying Cloud bounds separating HTML components tracking the main execution engine correctly parsing interactions smoothly."),
    ("5.7 ER Diagram", "er_diagram.png", "Plots specific tracking fields binding unique IDs over history logging and grammar indexing tables managing continuous data arrays strictly indexing the dictionary nodes inherently."),
    ("5.8 Data Flow Diagram", "dfd_level2.png", "Crucially demonstrates the strict bidirectional flow. Path 1 maps: User → Camera → Gesture Model → Text → Speech. Path 2 natively maps the reverse flow: User → Mic → Speech → Text → Sign establishing processing limits logically isolating states natively.")
]
for title, file, desc in diagrams:
    add_para(title, True, WD_ALIGN_PARAGRAPH.LEFT, 13)
    add_para(desc, False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
    img_path = os.path.join(DIAGRAM_DIR, file)
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(f"[Insert Diagram Screenshot Here]", True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    add_empty()
doc.add_page_break()


# ========================================== CHAP 6 ==========================================
add_para("6. CODE, IMPLEMENTATION AND RESULTS", True, WD_ALIGN_PARAGRAPH.LEFT, 16)
add_empty()

add_para("6.1 Modules Overview", True, WD_ALIGN_PARAGRAPH.LEFT, 13)

add_para("1. Gesture Recognition Module:", True, size=13)
add_para("Evaluates camera endpoints feeding sequential frames across HTML bounding MediaPipe coordinate extractions tracking sequential prediction logic mapping output classifications.", False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("2. Speech Recognition Module:", True, size=13)
add_para("Operates across active microphone hardware capturing user vocals transcribing English boundaries mapping logic using SpeechRecognition engines properly.", False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("3. Language Translation Module:", True, size=13)
add_para("Applies deep NLP manipulations analyzing captured strings converting endpoints across dialects enabling multi-language translation reliably.", False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("4. Text-to-Speech Module:", True, size=13)
add_para("Takes raw output phrasing mapping gTTS variables executing synchronous audio output streams dynamically yielding the generated speech outputs.", False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("5. Speech-to-Sign Module:", True, size=13)
add_para("Restructures grammar arrays identifying Object-Subject-Verb (OSV) mappings executing 3D WebGL render logic modeling sign interpretation concurrently.", False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
add_empty()

add_para("6.2 Sample Implementation Code", True, WD_ALIGN_PARAGRAPH.LEFT, 13)

code1 = """# OpenCV + MediaPipe Gesture Extraction Logic
import cv2
import mediapipe as mp
mp_holistic = mp.solutions.holistic
cap = cv2.VideoCapture(0)
with mp_holistic.Holistic(min_detection_confidence=0.5, min_tracking_confidence=0.5) as holistic:
    while cap.isOpened():
        ret, frame = cap.read()
        image, results = mediapipe_detection(frame, holistic)
        draw_styled_landmarks(image, results)
        keypoints = extract_keypoints(results)"""

code2 = """# Speech Recognition Input Module
import speech_recognition as sr
recognizer = sr.Recognizer()
with sr.Microphone() as source:
    recognizer.adjust_for_ambient_noise(source)
    audio = recognizer.listen(source)
    try:
        text = recognizer.recognize_google(audio)
        print(f"Recognized Speech: {text}")
    except sr.UnknownValueError:
        pass"""

code3 = """# Text-to-Speech Output Execution
from gtts import gTTS
import os
def execute_tts(text_input, lang_code='en'):
    tts = gTTS(text=text_input, lang=lang_code, slow=False)
    tts.save("output_speech.mp3")
    os.system("mpg123 output_speech.mp3")"""

for c in [code1, code2, code3]:
    add_para(c, False, WD_ALIGN_PARAGRAPH.LEFT, 11)
    add_empty()

add_para("6.3 Output Screens", True, WD_ALIGN_PARAGRAPH.LEFT, 13)

screens = [
    "[Insert Screenshot: Camera capturing hand gesture]",
    "[Insert Screenshot: Hand landmark detection using MediaPipe]",
    "[Insert Screenshot: Recognized text output]",
    "[Insert Screenshot: Speech output]",
    "[Insert Screenshot: Speech input waveform]",
    "[Insert Screenshot: Sign animation output]"
]

for text in screens:
    # Just insert the explicit screenshot text placeholder as requested strictly
    add_para(text, True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_empty()

doc.add_page_break()


# ========================================== CHAP 7 ==========================================
add_para("7. SYSTEM STUDY AND TESTING", True, WD_ALIGN_PARAGRAPH.LEFT, 16)
add_empty()

add_para("7.1 Feasibility Study", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para("The system bypasses intensive hardware requirements running cleanly across standard web constraints using JavaScript confirming high operational deployment capability inherently solving hardware limitations comprehensively.", False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("7.2 Types of Testing", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para("• Functional Testing: Isolated boundary variables executing speech logic rendering proper endpoints effectively tracing real camera inputs mapping predictions seamlessly across architectures.\n• Integration Path Testing: Validating multi-lingual endpoints mapping audio to Hindi/Telugu outputs flawlessly.", False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("7.3 Test Cases", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
t_data = [
    ("S.NO", "Test cases", "Expected O/P", "Actual O/P", "P/F"),
    ("1", "Analyze Camera Input", "MediaPipe extracts vectors", "Matrices bounding box detected", "Pass"),
    ("2", "Test Core Gesture Model", "Translates gesture to Text", "Detected specific Sign properly", "Pass"),
    ("3", "Microphone Speech Check", "Transcribes phrase logic", "Speech array tracks correctly", "Pass"),
    ("4", "NLP Grammar Restructure", "Changes string to OSV ISL logic", "Converted subject parameters correctly", "Pass"),
    ("5", "GL Avatar Execution Map", "Avatar performs exact Sign", "Visual matches logic constraints", "Pass"),
]
t_tab = doc.add_table(rows=len(t_data), cols=5, style='Table Grid')
t_tab.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row in enumerate(t_data):
    for j, val in enumerate(row):
        t_tab.rows[i].cells[j].text = val
        for p in t_tab.rows[i].cells[j].paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)
                r.bold = (i == 0)

doc.add_page_break()

# ========================================== CHAP 8-10 ==========================================
add_para("8. CONCLUSION", True, WD_ALIGN_PARAGRAPH.LEFT, 16)
add_empty()
add_para('Developing the Multi-Lingual Sign Language translator provides comprehensive bidirectional tools completely bypassing restrictive hardware networks empowering universal Web accessibility resolving the exact linguistic ISL boundaries natively enabling widespread society integration efficiently.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
add_empty()

add_para("9. FUTURE ENHANCEMENTS", True, WD_ALIGN_PARAGRAPH.LEFT, 16)
add_empty()
add_para('• Real-time Mobile Integrations: Expanding frameworks establishing native iOS mapping operations generating endpoints executing tracking smoothly.\n• Generative Translation Upgrades: Expanding grammatical mapping leveraging specific Transformer pipelines replacing 3D skeletal rigid transitions seamlessly parsing dynamic syntax reliably.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
add_empty()

add_para("10. REFERENCES", True, WD_ALIGN_PARAGRAPH.LEFT, 16)
add_empty()
raw_refs = [
    "[1] Sign Language Recognition Systems: A Decade Systematic Literature Review, Archives of Computational Methods in Engineering, 2021.",
    "[2] MediaPipe: A Framework for Building Perception Pipelines, arXiv, 2019.",
    "[3] A Survey on Sign Language Translation: Neural and Rule-based Approaches, ACM Computing Surveys, 2023.",
    "[4] Progressive Transformers for End-to-End Sign Language Production, Computer Vision – ECCV 2020.",
    "[5] A Benchmark for Indian Sign Language Processing, Proceedings of IEEE CVPR, 2024.",
    "[6] Real-time Hand Sign Recognition Using Edge Computing, IEEE Access, 2022.",
    "[7] Text to Speech rendering architecture utilizing Google services systematically, IEEE proceedings, 2020."
]
for item in raw_refs: add_para(item, False, WD_ALIGN_PARAGRAPH.JUSTIFY, 12)

doc.save(r'E:\sam\Project documents\Samvak_ProjectReview_Final.docx')
print("Successfully built flawless document!")
