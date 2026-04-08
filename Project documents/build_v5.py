import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml

# Load existing content logic (or we can just define a new document structure)
doc = Document()

# --- Page Borders ---
s = doc.sections[0]
s.top_margin = Cm(0.81)
s.bottom_margin = Cm(0.49)
s.left_margin = Cm(2.33)
s.right_margin = Cm(2.50)

sectPr = s._sectPr
pgBorders = parse_xml(
    '<w:pgBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:offsetFrom="page">'
    '<w:top w:val="thinThickSmallGap" w:sz="24" w:space="24" w:color="auto"/>'
    '<w:left w:val="thinThickSmallGap" w:sz="24" w:space="24" w:color="auto"/>'
    '<w:bottom w:val="thickThinSmallGap" w:sz="24" w:space="24" w:color="auto"/>'
    '<w:right w:val="thickThinSmallGap" w:sz="24" w:space="24" w:color="auto"/>'
    '</w:pgBorders>'
)
sectPr.append(pgBorders)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(13)
style.paragraph_format.line_spacing = 1.5

def add_para(text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=13, space_before=None, space_after=None):
    p = doc.add_paragraph()
    p.alignment = align
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = 'Times New Roman'
    return p

def add_empty(n=1):
    for _ in range(n):
        doc.add_paragraph()

LOGO_PATH = r"E:\sam\Project documents\extracted_images\image_0.png"
DIAGRAM_DIR = r"E:\sam\Project documents\diagrams"
SCREENSHOT_DIR = r"E:\sam\Project documents\screenshots"

# --- Page 1: Cover ---
add_empty(3)
add_para('ATTRIBUTE-BASED DATA SHARING SCHEME REVISITED IN CLOUD COMPUTING' if False else 'MULTI-LINGUAL SIGN LANGUAGE TO SPEECH AND SPEECH TO SIGN TRANSLATOR', True, WD_ALIGN_PARAGRAPH.CENTER, 16)
add_empty(1)
add_para('A Project Report', False, WD_ALIGN_PARAGRAPH.CENTER, 13)
add_para('Submitted in the partial fulfillment of the requirements for', False, WD_ALIGN_PARAGRAPH.CENTER, 13)
add_empty(1)
add_para('Master of Computer Applications', True, WD_ALIGN_PARAGRAPH.CENTER, 16)
add_para('In', False, WD_ALIGN_PARAGRAPH.LEFT, 13).alignment = WD_ALIGN_PARAGRAPH.CENTER
add_para('Department of Computer Science and Applications', True, WD_ALIGN_PARAGRAPH.CENTER, 15)
add_para('By', False, WD_ALIGN_PARAGRAPH.LEFT, 13).alignment = WD_ALIGN_PARAGRAPH.CENTER
add_empty(1)
add_para('Uppalapati Naveen Varma', True, WD_ALIGN_PARAGRAPH.CENTER, 16)
add_empty(1)
add_para('(2401600155)', True, WD_ALIGN_PARAGRAPH.CENTER, 13)
add_para('Under the supervision of', False, WD_ALIGN_PARAGRAPH.CENTER, 13)
add_para('Mrs. Swathi Voddi', True, WD_ALIGN_PARAGRAPH.CENTER, 13)
add_para('Assistant Professor', True, WD_ALIGN_PARAGRAPH.CENTER, 13)
add_empty(3)
if os.path.exists(LOGO_PATH):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(LOGO_PATH, width=Inches(1.8))
add_para('Department of Computer Science and Applications', False, WD_ALIGN_PARAGRAPH.CENTER, 16)
add_para('K L E F, Green Fields,', False, WD_ALIGN_PARAGRAPH.CENTER, 12)
add_para('Vaddeswaram, Guntur, Andhra Pradesh, India- 522502.', True, WD_ALIGN_PARAGRAPH.CENTER, 12)
add_para('2025- 26', False, WD_ALIGN_PARAGRAPH.CENTER, 12)
doc.add_page_break()

# --- Page 2: Declaration ---
add_empty(4)
add_para('KONERU LAKSHMAIAH EDUCATION FOUNDATION DEPARTMENT OF COMPUTER SCIENCE AND APPLICATIONS', True, WD_ALIGN_PARAGRAPH.CENTER, 12)
add_empty(2)
add_para('DECLARATION', True, WD_ALIGN_PARAGRAPH.CENTER, 13)
add_empty(2)
add_para('The Project Report entitled "MULTI-LINGUAL SIGN LANGUAGE TO SPEECH AND SPEECH TO SIGN TRANSLATOR" is a record of bonafide work carried out by me under the guidance of Mrs. Swathi Voddi, Assistant Professor, Department of Computer Science and Applications, Koneru Lakshmaiah Education Foundation, Vaddeswaram.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
add_empty(1)
add_para('I hereby declare that this project work is original and has not been submitted to any other university or institution for the award of any degree or diploma. All the information furnished in this project report is genuine to the best of my knowledge and belief.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
add_empty(6)
add_para('Signature of the Student                                Uppalapati Naveen Varma (2401600155)', True, WD_ALIGN_PARAGRAPH.LEFT, 13)
doc.add_page_break()

# --- Page 3: Certificate ---
add_empty(4)
add_para('KONERU LAKSHMAIAH EDUCATION FOUNDATION DEPARTMENT OF COMPUTER SCIENCE AND APPLICATIONS', True, WD_ALIGN_PARAGRAPH.CENTER, 12)
add_empty(2)
add_para('CERTIFICATE', True, WD_ALIGN_PARAGRAPH.CENTER, 13)
add_empty(2)
add_para('This is to certify that the Project Report entitled "MULTI-LINGUAL SIGN LANGUAGE TO SPEECH AND SPEECH TO SIGN TRANSLATOR" is a bonafide record of the work done by Uppalapati Naveen Varma (2401600155) in partial fulfillment of the requirements for the award of the degree of Master of Computer Applications from the Department of Computer Science and Applications, Koneru Lakshmaiah Education Foundation, Vaddeswaram, during the academic year 2025–26.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
add_empty(6)

p = add_para('Signature of the Supervisor', True, WD_ALIGN_PARAGRAPH.LEFT, 13)
p.add_run(' '*60 + 'Signature of the HOD').bold = True

p2 = add_para('(Mrs. Swathi Voddi)', True, WD_ALIGN_PARAGRAPH.LEFT, 13)
p2.add_run(' '*78 + '(Dr. Ch. Kiran Kumar)').bold = True

p3 = add_para('Assistant Professor', True, WD_ALIGN_PARAGRAPH.LEFT, 13)
p3.add_run(' '*90 + 'Professor & HOD').bold = True

add_empty(4)
add_para('Signature of the Examiner', True, WD_ALIGN_PARAGRAPH.LEFT, 13)
doc.add_page_break()

# --- Page 4: Acknowledgement ---
add_empty(2)
add_para('ACKNOWLEDGEMENT', True, WD_ALIGN_PARAGRAPH.CENTER, 14)
add_empty(1)
add_para('The satisfaction that accompanies the successful completion of any task would be incomplete without the mention of people who made it possible, whose constant guidance and encouragement crown all the efforts with success.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
add_para('I am very thankful to my project guide Mrs. Swathi Voddi, Assistant Professor, for her continuous support, encouragement, and invaluable guidance in completing this project. Her technical expertise in the domain of machine learning and computer vision was instrumental in shaping this work.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
add_para('I express my heartfelt gratitude to Dr. Ch. Kiran Kumar, Head of the Department of Computer Science and Applications, for providing me with the opportunity and facilities to carry out this project successfully.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
add_para('I express my sincere thanks to all the faculty members of the Department of Computer Science and Applications for imparting the knowledge that laid the foundation for this project.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
add_para('Last but not the least, I thank all Teaching and Non-Teaching Staff of our department and especially my classmates and friends who directly or indirectly helped me in completing this project.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
add_empty(4)
p = add_para('Signature of the Student\nUppalapati Naveen Varma\n(2401600155)', True, WD_ALIGN_PARAGRAPH.RIGHT, 13)
doc.add_page_break()

# --- Page 5: Abstract ---
add_empty(2)
add_para('ABSTRACT', True, WD_ALIGN_PARAGRAPH.CENTER, 14)
add_empty(1)
add_para('This project presents the design and implementation of Samvak — a Multi-Lingual Sign Language to Speech and Speech to Sign Translator, a real-time, web-based application that bridges the communication gap between the hearing-impaired community and the general population. The system operates in two primary modes: (1) Sign-to-Text/Speech, where a webcam captures hand gestures performed in Indian Sign Language (ISL), processes them through a deep learning pipeline using MediaPipe for landmark extraction and an LSTM-based Temporal Convolutional Network (TCN) for classification, and produces text and spoken output in eight languages; and (2) Speech-to-Sign, where spoken or typed input is transcribed, translated into ISL gloss sequences using rule-based grammar transformation, and rendered as sign animations through a real-time 3D avatar in the browser.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
add_para('The application is built on a Flask backend with a responsive HTML/CSS/JavaScript frontend. MediaPipe Holistic provides 258-dimensional pose and hand feature vectors from each video frame, which are fed into a trained LSTM-TCN model capable of recognizing 40 ISL phrases with 94.6% top-1 accuracy. For speech-to-sign translation, the system employs the SpeechRecognition library for automatic speech recognition, spaCy for natural language tokenization, and a local ISL grammar engine that applies Object-Subject-Verb (OSV) reordering, article dropping, and WH-word repositioning to generate correct ISL gloss sequences. A 3D humanoid avatar, rendered using Three.js, performs the corresponding sign gestures in real time.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
add_para('The system supports multilingual output in English, Hindi, Telugu, Tamil, Kannada, Malayalam, Spanish, and French using the Google Translate API and gTTS for text-to-speech synthesis. Additional features include an interactive ISL dictionary, a gamified learning module with XP-based progress tracking, user authentication, a personalized dashboard, and translation history logging. The application is deployed on Render for cloud accessibility.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
add_empty(1)
p = add_para('Keywords: ', True, WD_ALIGN_PARAGRAPH.LEFT, 13)
p.add_run('Indian Sign Language, gesture recognition, speech recognition, MediaPipe, LSTM, real-time translation, 3D avatar, multilingual, accessibility').bold = False
doc.add_page_break()

# --- Page 6: TABLE OF CONTENTS ---
add_para('TABLE OF CONTENTS', True, WD_ALIGN_PARAGRAPH.CENTER, 14)
add_empty(1)
toc_data = [
    ("CHAPTER NO.", "TITLE", "PAGE NO."),
    ("1", "INTRODUCTION", "1"),
    ("1.1", "Motivation", "1"),
    ("1.2", "Problem Statement", "1"),
    ("1.3", "Objective of the Project", "2"),
    ("1.4", "Scope", "2"),
    ("1.5", "Project Introduction", "2"),
    ("2", "LITERATURE SURVEY", "3"),
    ("2.1", "Related Work", "3"),
    ("3", "SYSTEM ANALYSIS", "5"),
    ("3.1", "Existing System", "5"),
    ("3.2", "Disadvantages of Existing System", "5"),
    ("3.3", "Proposed System", "5"),
    ("3.4", "Advantages of Proposed System", "6"),
    ("3.5", "Workflow of Proposed System", "6"),
    ("4", "REQUIREMENT ANALYSIS", "8"),
    ("4.1", "Functional and Non-Functional Requirements", "8"),
    ("4.2", "Hardware Requirements", "8"),
    ("4.3", "Software Requirements", "9"),
    ("4.4", "Architecture", "9"),
    ("5", "SYSTEM DESIGN", "10"),
    ("5.1", "Input and Output Design", "10"),
    ("5.2", "UML Diagrams", "11"),
    ("5.3", "DFD Diagrams", "14"),
    ("6", "CODE, IMPLEMENTATION AND RESULT", "16"),
    ("7", "SYSTEM STUDY AND TESTING", "26"),
    ("7.1", "Feasibility Study", "26"),
    ("7.2", "System Testing", "27"),
    ("7.3", "Types of Tests", "27"),
    ("8", "CONCLUSION", "29"),
    ("9", "FUTURE ENHANCEMENT", "30"),
    ("", "REFERENCES", "31")
]
toc_table = doc.add_table(rows=len(toc_data), cols=3)
toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (c1, c2, c3) in enumerate(toc_data):
    toc_table.rows[i].cells[0].text = c1
    toc_table.rows[i].cells[1].text = c2
    toc_table.rows[i].cells[2].text = c3
    for cell in toc_table.rows[i].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
                if i == 0 or len(c1) == 1 or c1 == "":  # Main chapters bold
                    r.bold = True
                else:
                    r.bold = False

doc.add_page_break()

# We save the file using doc.save() at the end.
doc.save(r"E:\sam\Project documents\Samvak_ProjectReview_Final.docx")
print("Done creating the first skeleton.")
