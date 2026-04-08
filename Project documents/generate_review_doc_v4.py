"""
Generate the complete Project Review Document as a formatted Word (.docx) file.
Following the exact formatting of the reference document.
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import qn
import os

doc = Document()

# ── Margins & Page Borders Setup ──
section = doc.sections[0]
section.top_margin = Cm(0.81)
section.bottom_margin = Cm(0.49)
section.left_margin = Cm(2.33)
section.right_margin = Cm(2.50)

def add_page_border(section):
    sectPr = section._sectPr
    pgBorders = parse_xml(
        '<w:pgBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' w:offsetFrom="page">'
        '<w:top w:val="thinThickSmallGap" w:sz="24" w:space="24" w:color="auto"/>'
        '<w:left w:val="thinThickSmallGap" w:sz="24" w:space="24" w:color="auto"/>'
        '<w:bottom w:val="thickThinSmallGap" w:sz="24" w:space="24" w:color="auto"/>'
        '<w:right w:val="thickThinSmallGap" w:sz="24" w:space="24" w:color="auto"/>'
        '</w:pgBorders>'
    )
    sectPr.append(pgBorders)

add_page_border(section)

# ── Style Setup ──
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(13)
pf = style.paragraph_format
pf.line_spacing = 1.5

def add_para(text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=13, spacing_after=None, space_before=None):
    p = doc.add_paragraph()
    p.alignment = align
    if spacing_after is not None:
        p.paragraph_format.space_after = Pt(spacing_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_empty_lines(n=1):
    for _ in range(n):
        doc.add_paragraph()

LOGO_PATH = r"E:\sam\Project documents\extracted_images\image_0.png"
DIAGRAM_DIR = r"E:\sam\Project documents\diagrams"
SCREENSHOT_DIR = r"E:\sam\Project documents\screenshots"


# ═══════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════
add_empty_lines(3)
add_para("MULTI-LINGUAL SIGN LANGUAGE TO SPEECH AND SPEECH TO SIGN TRANSLATOR", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
add_empty_lines(1)
add_para("A Project Report", align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
add_para("Submitted in the partial fulfillment of the requirements for", align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
add_empty_lines(1)
add_para("Master of Computer Applications", bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=15)
add_para("In", align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
add_para("Department of Computer Science and Applications", bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
add_para("By", align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
add_empty_lines(1)
add_para("Uppalapati Naveen Varma", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
add_empty_lines(1)
add_para("(2401600155)", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
add_para("Under the supervision of", align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
add_para("Mrs. Swathi Voddi", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
add_para("Assistant Professor", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
add_empty_lines(4)

if os.path.exists(LOGO_PATH):
    lp = doc.add_paragraph()
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp.add_run().add_picture(LOGO_PATH, width=Inches(1.8))

add_para("Department of Computer Science and Applications", align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
add_para("K L E F, Green Fields,", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_para("Vaddeswaram, Guntur, Andhra Pradesh, India- 522502.", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_para("2025- 26", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# DECLARATION
# ═══════════════════════════════════════════════════════════════
add_empty_lines(3)
add_para("DECLARATION", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
add_empty_lines(1)
add_para(
    'The Project Report entitled "MULTI-LINGUAL SIGN LANGUAGE TO SPEECH AND SPEECH TO SIGN TRANSLATOR" '
    'is a record of bonafide work carried out by me under the guidance of Mrs. Swathi Voddi, Assistant Professor, '
    'Department of Computer Science and Applications, Koneru Lakshmaiah Education Foundation, Vaddeswaram.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)
add_empty_lines(1)
add_para(
    'I hereby declare that this project work is original and has not been submitted to any other university or '
    'institution for the award of any degree or diploma. All the information furnished in this project report is '
    'genuine to the best of my knowledge and belief.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)
add_empty_lines(6)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.add_run("Signature of the Student ".ljust(60) + "Uppalapati Naveen Varma (2401600155)").font.bold = True

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# CERTIFICATE
# ═══════════════════════════════════════════════════════════════
add_empty_lines(5)
add_para("KONERU LAKSHMAIAH EDUCATION FOUNDATION DEPARTMENT OF COMPUTER SCIENCE AND APPLICATIONS", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_empty_lines(2)
add_para("CERTIFICATE", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
add_empty_lines(2)
add_para(
    'This is to certify that the Project Report entitled "MULTI-LINGUAL SIGN LANGUAGE TO SPEECH AND SPEECH TO SIGN TRANSLATOR" '
    'is a bonafide record of the work done by Uppalapati Naveen Varma (2401600155) in partial fulfillment of the requirements '
    'for the award of the degree of Master of Computer Applications from the Department of Computer Science and Applications, '
    'Koneru Lakshmaiah Education Foundation, Vaddeswaram, during the academic year 2025–26.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)
add_empty_lines(5)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("Signature of the Supervisor".ljust(60) + "Signature of the HOD")
r.font.bold = True

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
p2.add_run("(Mrs. Swathi Voddi)".ljust(75) + "(Dr. Ch. Kiran Kumar)")

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
p3.add_run("Assistant Professor".ljust(80) + "Professor & HOD")

add_empty_lines(3)
add_para("Signature of the Examiner", bold=True)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# ACKNOWLEDGEMENT
# ═══════════════════════════════════════════════════════════════
add_empty_lines(1)
add_para("ACKNOWLEDGEMENT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
add_empty_lines(1)

add_para(
    'The satisfaction that accompanies the successful completion of any task would be incomplete without the mention '
    'of people who made it possible, whose constant guidance and encouragement crown all the efforts with success.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)
add_para(
    'I am very thankful to my project guide Mrs. Swathi Voddi, Assistant Professor, for her continuous support, '
    'encouragement, and invaluable guidance in completing this project. Her technical expertise in the domain of '
    'machine learning and computer vision was instrumental in shaping this work.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)
add_para(
    'I express my heartfelt gratitude to Dr. Ch. Kiran Kumar, Head of the Department of Computer Science and Applications, '
    'for providing me with the opportunity and facilities to carry out this project successfully.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)
add_para(
    'I express my sincere thanks to all the faculty members of the Department of Computer Science and Applications '
    'for imparting the knowledge that laid the foundation for this project.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)
add_para(
    'Last but not the least, I thank all Teaching and Non-Teaching Staff of our department and especially my '
    'classmates and friends who directly or indirectly helped me in completing this project.',
    align=WD_ALIGN_PARAGRAPH.JUSTIFY
)
add_empty_lines(4)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.add_run("Signature of the Student\nUppalapati Naveen Varma\n(2401600155)").font.bold = True

doc.add_page_break()
