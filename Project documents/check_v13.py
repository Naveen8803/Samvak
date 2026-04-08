from docx import Document
import os

path = r'E:\sam\Project documents\Samvak_Review_Final_V13.docx'
doc = Document(path)

imgs  = [r for r in doc.part.rels.values() if 'image' in r.reltype]
codes = [p for p in doc.paragraphs if p.runs and p.runs[0].font.name == 'Courier New']
words = sum(len(p.text.split()) for p in doc.paragraphs if p.text.strip())
sec   = doc.sections[0]

print('File size   :', round(os.path.getsize(path)/1024), 'KB')
print('Paragraphs  :', len(doc.paragraphs))
print('Images      :', len(imgs))
print('Code blocks :', len(codes))
print('Tables      :', len(doc.tables))
print('Word count~ :', words)
print('Left margin :', round(sec.left_margin.inches, 3), 'inches (spec: 0.920)')
print('Top margin  :', round(sec.top_margin.inches, 3), 'inches (spec: 0.320)')
print('Page width  :', round(sec.page_width.inches, 2), 'inches (spec: 8.50)')

print()
print('=== 16pt CHAPTER HEADINGS ===')
for p in doc.paragraphs:
    t = p.text.strip()
    if t and p.runs and p.runs[0].font.size and p.runs[0].font.size.pt == 16:
        print(' ', t[:80])

print()
print('=== 14pt SUB-HEADINGS (first 20) ===')
count = 0
for p in doc.paragraphs:
    t = p.text.strip()
    if t and p.runs and p.runs[0].font.size and p.runs[0].font.size.pt == 14 and p.runs[0].bold:
        print(' ', t[:80])
        count += 1
        if count >= 20:
            break

print()
print('=== TABLE HEADERS ===')
for i, tb in enumerate(doc.tables):
    header = [c.text[:25] for c in tb.rows[0].cells] if tb.rows else []
    print(f'  Table {i+1} ({len(tb.rows)} rows x {len(tb.columns)} cols): {header}')
