import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document(r'E:\sam\Project documents\Samvak_Temp1.docx')
DIAGRAM_DIR = r"E:\sam\Project documents\diagrams"

def add_para(text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=13, space_before=None, space_after=None):
    p = doc.add_paragraph()
    p.alignment = align
    if space_before is not None: p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None: p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = 'Times New Roman'
    return p

def add_empty(n=1):
    for _ in range(n): doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# 2. LITERATURE SURVEY
# ═══════════════════════════════════════════════════════════════
add_para("2. LITERATURE SURVEY", True, WD_ALIGN_PARAGRAPH.LEFT, 16)
add_empty()

refs = [
    ('[1] Sign Language Recognition utilizing CNN and LSTM Architectures', 'This model maps robust spatial characteristics avoiding the limitations of flat vector machines. However, it completely lacks speech synthesis endpoints mapping outputs strictly into basic text strings.'),
    ('[2] MediaPipe Holistic Tracking within Real-time Computer Vision Systems', 'Highlights optimized browser extraction of human tracking points. It resolves physical boundary limits preventing users from wearing physical tracked gloves facilitating natural sign flow dynamics.'),
    ('[3] Continuous Gesture Recognition mapping American Sign Language', 'Tracks dynamic sequences across complex lighting scenarios. Its primary limitation is the absolute isolation of its logic solely encompassing ASL preventing ISL adoption entirely.'),
    ('[4] A Survey on Bidirectional Application Translation: Speech to Sign NLP', 'Validates transforming English syntax directly mimicking linguistic rules identifying verb modifiers isolating Subject/Object parameters effectively translating text outputs logically for Avatar ingestion constraints.'),
    ('[5] Progressive Translators: End-to-End Hand Gesture Modelling', 'Investigates directly mapping text phrases directly to generated skeletal models minimizing rendering latencies highlighting why WebGL technologies successfully supersede native application render bounds.'),
    ('[6] Implementing Multi-lingual Vocals mapping NLP strings via gTTS', 'Highlights programmatic methods utilizing regional speech identifiers configuring automated pitch modulations eliminating monolingual dependencies.'),
    ('[7] Indian Sign Language Database Parsing and Implementation Algorithms', 'Addresses the severe dataset deficits found regionally demonstrating structural approaches mapping generalized Indian vocabulary nodes scaling accurate categorical modeling datasets.')
]

for title, desc in refs:
    add_para(title, True, size=13)
    add_para(desc, size=13, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_empty()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 3. SYSTEM ANALYSIS
# ═══════════════════════════════════════════════════════════════
add_para("3. SYSTEM ANALYSIS", True, WD_ALIGN_PARAGRAPH.LEFT, 16)
add_empty()

add_para("3.1 Existing System", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('Current translation environments leverage sensor hardware tracking localized hand movements parsing standard ASL vocabularies returning fixed English interpretations lacking bidirectional interactions enabling deaf individuals to simultaneously understand speaking users independently.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("3.2 Disadvantages of Existing System", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('• High Hardware Dependence: Restricts access strictly relying on expensive sensor nodes masking operational accessibility.\n• Monolingual Bias: Hardcoded output processing limits global scaling preventing regional implementations significantly.\n• Syntactical Incorrectness: Processing word-by-word without rearranging Subject/Object mapping yielding grammatically unviable visuals.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("3.3 Proposed System", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('The proposed tool eliminates hardware leveraging web camera streams calculating node boundaries mapping predictive sequences leveraging customized ISL model weights resolving textual/vocal responses dynamically. Conversely, recording live mic data tokenizes phrasing matching structural Avatar parameters parsing real-time ISL visuals bi-directionally.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("3.4 Advantages of Proposed System", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('• Zero Hardware Complexity: Bypasses flex-glove dependency operating smoothly over ubiquitous digital webcams natively.\n• Universal Multi-Language Scaling: Embeds text integrations supporting massive dialect toggling simultaneously routing output syntaxes correctly.\n• Real-time Grammar Engine: Transcribes audio correctly assigning token hierarchies modeling native OSV rules before translating avatar boundaries logically.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("3.5 Workflow Explanation", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('The logical bounds configure initial client interactions routing user inputs specifically between a visual camera feed processing gesture mapping parameters mapping outputs toward deep-translation / Text-To-Speech interfaces directly. Alternatively, invoking the microphone processes Google Speech transcripts extracting core phrases mapping gloss rules outputting visual avatar bounds accurately mitigating inference limits reliably.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 4. REQUIREMENT ANALYSIS
# ═══════════════════════════════════════════════════════════════
add_para("4. REQUIREMENT ANALYSIS", True, WD_ALIGN_PARAGRAPH.LEFT, 16)
add_empty()

add_para("4.1 Functional Requirements", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('1. Web camera ingestion must extract and validate 258 coordinate dimensions natively tracking hand parameters.\n2. Model operations must process sequence tracking mapping categorical predictions evaluating valid lexical sign tokens accurately.\n3. Bidirectional logic processing microphone streams accurately logging strings evaluating syntactic language configurations outputting proper visual models.\n4. Translating raw text into 8 regional formats parsing outputs mapping exact linguistic representations synthetically.\n5. Tracking Dictionary progress managing dynamic point schemas indexing historical references continuously.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("4.2 Non-Functional Requirements", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('• Low Latency Bounds: Ensuring end-to-end rendering logic evaluates predictions cleanly under 400ms mimicking actual conversation parameters natively.\n• High Availability: Responsive deployments routing client rendering interfaces maintaining execution loops seamlessly handling interrupted signals gracefully.\n• Security Integrity: Abstracted SQLite payloads managing local states encrypting internal registry passwords protecting backend routes effectively.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("4.3 Hardware Requirements", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('• Central Processing Unit: AMD Ryzen 3 / Intel i3 baseline\n• Allocation Memory: Minimum 8 GB internal RAM limits\n• Audio/Video Nodes: Functional internal/external Webcam and Microphone endpoints.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("4.4 Software Requirements", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('• Execution Tier: Python 3.10 and corresponding browser instances mapping HTML5 environments.\n• External Library Chains: Flask REST bindings, OpenCV, MediaPipe tracking limits, SpeechRecognition.\n• Synthesis Engines: Deep Translator APIs, Google Text to Speech module mapping parameters natively.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("4.5 System Architecture", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('Client bounds manage capturing arrays configuring visual bounding interfaces communicating mapping sequences asynchronously transmitting validation payloads resolving REST API paths across global Flask bindings tracking localized session parameters tracking SQLite models outputting processed payload coordinates mapping localized HTML template hierarchies natively maintaining high operational responsiveness processing independent WebGL operations fluidly.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 5. SYSTEM DESIGN
# ═══════════════════════════════════════════════════════════════
add_para("5. SYSTEM DESIGN", True, WD_ALIGN_PARAGRAPH.LEFT, 16)
add_empty()

diagrams = [
    ("5.1 Use Case Diagram", "use_case_diagram.png", "Showcases user boundaries interacting completely over primary sub-components identifying Authentication layers feeding input structures directly initiating the translation pipelines evaluating logic parameters."),
    ("5.2 Class Diagram", "class_diagram.png", "Translating operational classes linking logic structures outlining User Object persistence linking NLP grammar formatting boundaries passing model endpoints seamlessly."),
    ("5.3 Sequence Diagram", "sequence_sign.png", "Evaluates mapping structures tracing Client boundaries initiating capture constraints passing prediction loops interacting comprehensively across logical translation processing states generating outputs properly."),
    ("5.4 Activity Diagram", "activity_diagram.png", "Logs branch thresholds mapping conditional parameters generating accurate loop limits extracting frame sequences executing fallbacks predicting outputs correctly tracking system limits."),
    ("5.5 Component Diagram", "deployment_diagram.png", "Identifies distinct logical blocks linking Web Client extraction modules passing processing structures into AI/NLP parsing engines configuring translation pipelines mapping backend states precisely."),
    ("5.6 Deployment Diagram", "deployment_diagram.png", "Evaluates physical machine constraints managing browser architectures feeding logical server operations parsing SQLite persistence files evaluating deployment hosts natively."),
    ("5.7 ER Diagram", "er_diagram.png", "Plots data structures extracting Entity references outlining logic handling login records mapping Translation operations logging specific points configuring dictionary operations indexing natively."),
    ("5.8 Data Flow Diagram", "dfd_level2.png", "Demonstrates the bidirectional data handling mapping specifically: User → Camera → Gesture Model → Text → Speech AND User → Mic → Speech → Text → Sign establishing processing operations logically tracking continuous state interactions isolating pipeline sequences.")
]

for title, file, desc in diagrams:
    add_para(title, True, WD_ALIGN_PARAGRAPH.LEFT, 13)
    add_para(desc, False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
    img_path = os.path.join(DIAGRAM_DIR, file)
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(f"[Insert Diagram Screenshot Here]", True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    add_empty()

doc.save(r'E:\sam\Project documents\Samvak_Temp2.docx')
