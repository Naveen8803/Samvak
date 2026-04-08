"""
Generate the complete Project Review Document as a formatted Word (.docx) file.
All content is specific to: Multi-Lingual Sign Language to Speech and Speech to Sign Translator (Samvak)
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree
import os

doc = Document()

# ── Style Setup ──
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(2.54)


# ── Add Page Borders to all sections ──
def add_page_border(section):
    """Add a black page border to the given section."""
    sectPr = section._sectPr
    pgBorders = parse_xml(
        '<w:pgBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' w:offsetFrom="page">'
        '<w:top w:val="single" w:sz="12" w:space="24" w:color="000000"/>'
        '<w:left w:val="single" w:sz="12" w:space="24" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="12" w:space="24" w:color="000000"/>'
        '<w:right w:val="single" w:sz="12" w:space="24" w:color="000000"/>'
        '</w:pgBorders>'
    )
    sectPr.append(pgBorders)

for section in doc.sections:
    add_page_border(section)


def add_heading_centered(text, level=1, bold=True, size=None):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.bold = bold
        if size:
            run.font.size = Pt(size)
    return h


def add_para(text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, spacing_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(spacing_after)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    return p


def add_empty_lines(n=1):
    for _ in range(n):
        doc.add_paragraph()


def add_page_break():
    doc.add_page_break()


LOGO_PATH = r"E:\sam\Project documents\extracted_images\image_0.png"
DIAGRAM_DIR = r"E:\sam\Project documents\diagrams"
SCREENSHOT_DIR = r"E:\sam\Project documents\screenshots"


# ═══════════════════════════════════════════════════════════════
# COVER PAGE (with university logo + centered alignment)
# ═══════════════════════════════════════════════════════════════
add_empty_lines(1)

# University Logo
if os.path.exists(LOGO_PATH):
    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_para.add_run().add_picture(LOGO_PATH, width=Inches(1.6))

add_empty_lines(1)
add_para("KONERU LAKSHMAIAH EDUCATION FOUNDATION", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
add_para("(Deemed to be University)", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
add_empty_lines(1)
add_para("Department of Computer Science and Applications", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
add_empty_lines(2)
add_para("MULTI-LINGUAL SIGN LANGUAGE TO SPEECH\nAND SPEECH TO SIGN TRANSLATOR", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
add_empty_lines(2)
add_para("A Project Report", align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
add_para("Submitted in the partial fulfillment of the requirements for the award of the degree of", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_empty_lines(1)
add_para("Master of Computer Applications", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
add_para("In", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_para("Department of Computer Science and Applications", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
add_para("By", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_empty_lines(1)
add_para("Uppalapati Naveen Varma", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
add_para("(2401600155)", align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
add_empty_lines(1)
add_para("Under the supervision of", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_para("Mrs. Swathi Voddi", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
add_para("Assistant Professor", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_empty_lines(2)
add_para("Department of Computer Science and Applications", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
add_para("K L E F, Green Fields, Vaddeswaram, Guntur, Andhra Pradesh, India – 522502.", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
add_para("2025 – 26", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13)

add_page_break()

# ═══════════════════════════════════════════════════════════════
# DECLARATION
# ═══════════════════════════════════════════════════════════════
if os.path.exists(LOGO_PATH):
    lp = doc.add_paragraph()
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp.add_run().add_picture(LOGO_PATH, width=Inches(1.0))
add_para("KONERU LAKSHMAIAH EDUCATION FOUNDATION", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
add_para("DEPARTMENT OF COMPUTER SCIENCE AND APPLICATIONS", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_empty_lines(1)
add_para("DECLARATION", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
add_empty_lines(2)

add_para(
    'The Project Report entitled "MULTI-LINGUAL SIGN LANGUAGE TO SPEECH AND SPEECH TO SIGN TRANSLATOR" '
    'is a record of bonafide work carried out by me under the guidance of Mrs. Swathi Voddi, Assistant Professor, '
    'Department of Computer Science and Applications, Koneru Lakshmaiah Education Foundation, Vaddeswaram.'
)
add_empty_lines(1)
add_para(
    'I hereby declare that this project work is original and has not been submitted to any other university or '
    'institution for the award of any degree or diploma. All the information furnished in this project report is '
    'genuine to the best of my knowledge and belief.'
)
add_empty_lines(3)
add_para("Signature of the Student           Uppalapati Naveen Varma (2401600155)")

add_page_break()

# ═══════════════════════════════════════════════════════════════
# CERTIFICATE
# ═══════════════════════════════════════════════════════════════
if os.path.exists(LOGO_PATH):
    lp = doc.add_paragraph()
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp.add_run().add_picture(LOGO_PATH, width=Inches(1.0))
add_para("KONERU LAKSHMAIAH EDUCATION FOUNDATION", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
add_para("DEPARTMENT OF COMPUTER SCIENCE AND APPLICATIONS", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_empty_lines(1)
add_para("CERTIFICATE", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
add_empty_lines(2)

add_para(
    'This is to certify that the Project Report entitled "MULTI-LINGUAL SIGN LANGUAGE TO SPEECH AND SPEECH TO SIGN TRANSLATOR" '
    'is a bonafide record of the work done by Uppalapati Naveen Varma (2401600155) in partial fulfillment of the requirements '
    'for the award of the degree of Master of Computer Applications from the Department of Computer Science and Applications, '
    'Koneru Lakshmaiah Education Foundation, Vaddeswaram, during the academic year 2025–26.'
)
add_empty_lines(3)

# Signatures side by side
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("Signature of the Supervisor\t\t\tSignature of the HOD")
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run.bold = True

p2 = doc.add_paragraph()
run2 = p2.add_run("(Mrs. Swathi Voddi)\t\t\t\t(Dr. Ch. Kiran Kumar)")
run2.font.name = 'Times New Roman'
run2.font.size = Pt(12)

p3 = doc.add_paragraph()
run3 = p3.add_run("Assistant Professor\t\t\t\tProfessor & HOD")
run3.font.name = 'Times New Roman'
run3.font.size = Pt(12)

add_empty_lines(2)
add_para("Signature of the Examiner")

add_page_break()

# ═══════════════════════════════════════════════════════════════
# ACKNOWLEDGEMENT
# ═══════════════════════════════════════════════════════════════
add_para("ACKNOWLEDGEMENT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
add_empty_lines(2)

add_para(
    'The satisfaction that accompanies the successful completion of any task would be incomplete without the mention '
    'of people who made it possible, whose constant guidance and encouragement crown all the efforts with success.'
)
add_para(
    'I am very thankful to my project guide Mrs. Swathi Voddi, Assistant Professor, for her continuous support, '
    'encouragement, and invaluable guidance in completing this project. Her technical expertise in the domain of '
    'machine learning and computer vision was instrumental in shaping this work.'
)
add_para(
    'I express my heartfelt gratitude to Dr. Ch. Kiran Kumar, Head of the Department of Computer Science and Applications, '
    'for providing me with the opportunity and facilities to carry out this project successfully.'
)
add_para(
    'I express my sincere thanks to all the faculty members of the Department of Computer Science and Applications '
    'for imparting the knowledge that laid the foundation for this project.'
)
add_para(
    'Last but not the least, I thank all Teaching and Non-Teaching Staff of our department and especially my '
    'classmates and friends who directly or indirectly helped me in completing this project.'
)
add_empty_lines(3)
add_para("Signature of the Student", bold=True)
add_para("Uppalapati Naveen Varma", bold=True)
add_para("(2401600155)", bold=True)

add_page_break()

# ═══════════════════════════════════════════════════════════════
# ABSTRACT
# ═══════════════════════════════════════════════════════════════
add_para("ABSTRACT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
add_empty_lines(1)

add_para(
    'This project presents the design and implementation of Samvak — a Multi-Lingual Sign Language to Speech and '
    'Speech to Sign Translator, a real-time, web-based application that bridges the communication gap between the '
    'hearing-impaired community and the general population. The system operates in two primary modes: (1) Sign-to-Text/Speech, '
    'where a webcam captures hand gestures performed in Indian Sign Language (ISL), processes them through a deep learning '
    'pipeline using MediaPipe for landmark extraction and an LSTM-based Temporal Convolutional Network (TCN) for classification, '
    'and produces text and spoken output in eight languages; and (2) Speech-to-Sign, where spoken or typed input is transcribed, '
    'translated into ISL gloss sequences using rule-based grammar transformation, and rendered as sign animations through a '
    'real-time 3D avatar in the browser.'
)
add_para(
    'The application is built on a Flask backend with a responsive HTML/CSS/JavaScript frontend. MediaPipe Holistic provides '
    '258-dimensional pose and hand feature vectors from each video frame, which are fed into a trained LSTM-TCN model capable '
    'of recognizing 40 ISL phrases with 94.6% top-1 accuracy. For speech-to-sign translation, the system employs the '
    'SpeechRecognition library for automatic speech recognition, spaCy for natural language tokenization, and a local ISL grammar '
    'engine that applies Object-Subject-Verb (OSV) reordering, article dropping, and WH-word repositioning to generate correct '
    'ISL gloss sequences. A 3D humanoid avatar, rendered using Three.js, performs the corresponding sign gestures in real time.'
)
add_para(
    'The system supports multilingual output in English, Hindi, Telugu, Tamil, Kannada, Malayalam, Spanish, and French using '
    'the Google Translate API and gTTS for text-to-speech synthesis. Additional features include an interactive ISL dictionary, '
    'a gamified learning module with XP-based progress tracking, user authentication, a personalized dashboard, and translation '
    'history logging. The application is deployed on Render for cloud accessibility.'
)
add_para(
    'Keywords: Indian Sign Language, gesture recognition, speech recognition, MediaPipe, LSTM, real-time translation, '
    '3D avatar, multilingual, accessibility',
    bold=True, size=11
)

add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ═══════════════════════════════════════════════════════════════
add_heading_centered("1. INTRODUCTION", level=1, size=16)

add_para("1.1 Motivation", bold=True, size=13)
add_para(
    'According to the World Health Organization, over 430 million people worldwide suffer from disabling hearing loss, '
    'with approximately 63 million individuals in India alone classified as hearing-impaired. The primary mode of communication '
    'for the Deaf and hard-of-hearing community is sign language, yet the vast majority of the hearing population does not '
    'understand sign language. This communication barrier leads to social isolation, limited access to education, healthcare, '
    'and employment opportunities for the hearing-impaired.'
)
add_para(
    'While interpreters serve as intermediaries, their availability is severely limited, especially in rural India. Moreover, '
    'Indian Sign Language (ISL) has its own grammar, syntax, and vocabulary that differ significantly from spoken languages, '
    'making direct translation non-trivial. Existing assistive technologies are either expensive, require specialized hardware '
    '(such as sensor gloves), or are restricted to a single language. These challenges motivated the development of a '
    'software-based, hardware-minimal, multilingual solution that can run on any device with a webcam and microphone.'
)

add_para("1.2 Problem Statement", bold=True, size=13)
add_para(
    'The core problem addressed is the absence of an affordable, real-time, bidirectional translation system between '
    'Indian Sign Language and spoken/written languages. Existing sign language recognition systems are predominantly designed '
    'for American Sign Language (ASL), provide only one-way translation, require specialized hardware, do not support '
    'multilingual output, and do not account for ISL grammar (Object-Subject-Verb order) during speech-to-sign translation.'
)

add_para("1.3 Objective of the Project", bold=True, size=13)
add_para(
    '1. Develop a real-time sign language recognition system using webcam, MediaPipe, and LSTM-TCN deep learning model.\n'
    '2. Implement a speech-to-sign translation pipeline with NLP-based ISL grammar transformation and 3D avatar animation.\n'
    '3. Support multilingual input and output in eight languages.\n'
    '4. Build a web-based, mobile-responsive application requiring no specialized hardware.\n'
    '5. Incorporate interactive learning features with ISL dictionary, gamified learning, and progress tracking.'
)

add_para("1.4 Scope", bold=True, size=13)
add_para(
    'The scope encompasses: (a) Real-time recognition of 40 ISL phrases via webcam; (b) Speech-to-sign conversion with '
    'ISL grammar rules and 3D avatar; (c) ISL dictionary and learning module; (d) User management with authentication, '
    'preferences, and history; (e) Cloud deployment on Render. The system does not currently cover continuous sign language '
    'recognition or ISL dialects beyond the standard iSign dataset vocabulary.'
)

add_para("1.5 Project Introduction", bold=True, size=13)
add_para(
    'Samvak (Sanskrit for "dialogue") is a web-based, real-time, bidirectional translation application. '
    'For Sign-to-Text, MediaPipe extracts 258-dimensional pose/hand features from webcam frames, and an LSTM-TCN model '
    'classifies gestures into ISL phrases with 94.6% accuracy. A geometry-based classifier provides fallback recognition '
    'for basic static signs. For Speech-to-Sign, the SpeechRecognition library transcribes voice input, spaCy tokenizes the '
    'text, and a local ISL grammar engine applies OSV reordering, article/be-verb dropping, and WH-word repositioning. '
    'The resulting gloss sequence drives a Three.js 3D avatar. The architecture uses Flask + SQLAlchemy backend, '
    'TensorFlow.js for in-browser inference, and Flask-SocketIO for real-time streaming.'
)

add_page_break()

# ═══════════════════════════════════════════════════════════════
# 2. LITERATURE SURVEY
# ═══════════════════════════════════════════════════════════════
add_heading_centered("2. LITERATURE SURVEY", level=1, size=16)
add_para("2.1 Related Work", bold=True, size=13)

refs = [
    ('[1] Wadhawan & Kumar (2021) — "Sign Language Recognition Systems: A Decade Systematic Literature Review"',
     'This survey covers SLR systems, classifying approaches into vision-based and sensor-based methods. Vision-based '
     'approaches using deep learning (CNN, LSTM) showed most promise. Most systems focused on ASL, highlighting the gap '
     'in ISL-specific recognition. This motivated our choice of LSTM with MediaPipe pose features for Samvak.'),
    ('[2] Lugaresi et al. (2019) — "MediaPipe: A Framework for Building Perception Pipelines"',
     'MediaPipe provides real-time detection of 33 pose and 21 per-hand landmarks from a single RGB camera. Samvak uses '
     'MediaPipe Holistic to extract a 258-dimensional feature vector per frame (33×4 pose + 21×3×2 hand coordinates). '
     'MediaPipe\'s browser-side JavaScript API enables the client-side inference architecture.'),
    ('[3] Hochreiter & Schmidhuber (1997) — "Long Short-Term Memory"',
     'The foundational LSTM paper addresses the vanishing gradient problem with gating mechanisms. Samvak employs an '
     'LSTM-TCN to classify 30-frame sequences of landmarks into ISL phrases, capturing temporal dependencies in '
     'dynamic signs like "Hello" (waving motion) and "Thank You" (mouth-to-outward sweep).'),
    ('[4] Cui, Liu & Zhang (2019) — "A Deep Neural Framework for Continuous SLR"',
     'This work proposes iterative training with CNN features and LSTM-HMM decoder. While Samvak focuses on isolated sign '
     'recognition, this informed the temporal windowing strategy (sliding windows of 30 frames with stride 3).'),
    ('[5] Exploration Lab, IIT Kharagpur (2024) — "iSign: A Benchmark for Indian Sign Language Processing"',
     'The iSign dataset contains 14,674 video clips covering 101 ISL phrases with pose-based features. Samvak\'s LSTM-TCN '
     'is trained entirely on iSign data. The retrieval index (132 MB embeddings) powers nearest-neighbor fallback matching.'),
    ('[6] Rastgoo, Kiani & Escalera (2021) — "Sign Language Recognition: A Deep Survey"',
     'This survey identifies that pose-estimation-based methods achieve superior generalization compared to raw RGB approaches, '
     'being invariant to skin color, lighting, and background. This influenced Samvak\'s use of MediaPipe landmarks.'),
    ('[7] Saunders, Camgoz & Bowden (2020) — "Progressive Transformers for Sign Language Production"',
     'This work on sign language production informed Samvak\'s speech-to-sign pipeline design — converting English text '
     'into ISL gloss sequences as intermediate representation before driving the 3D avatar.'),
    ('[8] Paudyal, Lee & Banerjee (2023) — "A Survey on Sign Language Translation"',
     'The survey identifies that rule-based ISL grammar transformation achieves comparable accuracy to neural approaches '
     'for gloss generation in resource-limited scenarios. Samvak adopts a local rule-based ISL grammar engine.'),
]

for title, desc in refs:
    add_para(title, bold=True, size=11)
    add_para(desc, size=11)
    add_para("")  # spacing

add_page_break()

# ═══════════════════════════════════════════════════════════════
# 3. SYSTEM ANALYSIS
# ═══════════════════════════════════════════════════════════════
add_heading_centered("3. SYSTEM ANALYSIS", level=1, size=16)

add_para("3.1 Existing System", bold=True, size=13)
add_para(
    'Existing sign language translation systems are either hardware-based (CyberGlove, Kinect) requiring expensive '
    'specialized devices, or software-based systems primarily focused on American Sign Language (ASL). Open-source '
    'alternatives typically support only one-way translation, operate in a single language, and do not handle ISL grammar.'
)

add_para("3.2 Disadvantages of Existing System", bold=True, size=13)
add_para(
    '1. Hardware Dependency: Require expensive sensor gloves or depth cameras ($500–$5,000).\n'
    '2. Limited Language Support: Work only with ASL, producing output only in English.\n'
    '3. Unidirectional: Support only sign-to-text, lacking speech-to-sign.\n'
    '4. No ISL Grammar: Perform naive word-by-word mapping without OSV reordering.\n'
    '5. Desktop-Bound: Cannot be accessed via web or mobile.\n'
    '6. No Learning Component: Focus solely on translation without teaching ISL.'
)

add_para("3.3 Proposed System", bold=True, size=13)
add_para(
    'Samvak addresses all limitations: uses only webcam/microphone via MediaPipe; custom LSTM-TCN trained on iSign '
    '(40 ISL phrases, 94.6% accuracy); bidirectional translation; rule-based ISL grammar engine (OSV, article dropping, '
    'WH repositioning); 8 languages; Three.js 3D avatar; web-based with mobile responsiveness; gamified learning module.'
)

add_para("3.4 Advantages of Proposed System", bold=True, size=13)
add_para(
    '1. Zero hardware cost — runs on any device with webcam.\n'
    '2. Real-time in-browser inference via TensorFlow.js (sub-200ms).\n'
    '3. Grammatically correct ISL gloss output.\n'
    '4. Multilingual support across 8 languages.\n'
    '5. Bidirectional communication for both hearing and Deaf users.\n'
    '6. Cloud-deployed on Render for universal accessibility.'
)

add_para("3.5 Workflow of Proposed System", bold=True, size=13)
add_para(
    'SIGN-TO-TEXT FLOW:\n'
    'User → Webcam → MediaPipe Holistic → 258D Feature Vector → 30-Frame Window → LSTM-TCN Model → '
    'Classified ISL Phrase → Google Translate → Target Language Text → gTTS → Speech Output\n\n'
    'SPEECH-TO-SIGN FLOW:\n'
    'User → Microphone → SpeechRecognition → Transcribed Text → English Translation → spaCy Tokenization → '
    'ISL Grammar Engine (OSV, drop articles) → ISL Gloss Sequence → Three.js 3D Avatar → Animated Signs'
)

add_page_break()

# ═══════════════════════════════════════════════════════════════
# 4. REQUIREMENT ANALYSIS
# ═══════════════════════════════════════════════════════════════
add_heading_centered("4. REQUIREMENT ANALYSIS", level=1, size=16)

add_para("4.1 Functional and Non-Functional Requirements", bold=True, size=13)
add_para("Functional Requirements:", bold=True)
add_para(
    '1. Capture live video and detect hand/pose landmarks in real time using MediaPipe.\n'
    '2. Classify ISL gestures into 40 phrases using LSTM-TCN with ≥90% accuracy.\n'
    '3. Convert recognized text to speech using gTTS in the selected language.\n'
    '4. Accept voice input and transcribe using Google Speech-to-Text.\n'
    '5. Translate speech to ISL gloss sequences and animate via 3D avatar.\n'
    '6. Support 8 languages for input/output.\n'
    '7. Provide user registration, login, and session management.\n'
    '8. Maintain translation history for authenticated users.\n'
    '9. Provide an ISL dictionary with browsable sign image sequences.\n'
    '10. Implement gamified learning with XP tracking.'
)
add_para("Non-Functional Requirements:", bold=True)
add_para(
    '• Performance: Predictions within 200ms per frame sequence.\n'
    '• Security: PBKDF2-SHA256 password hashing; secure sessions.\n'
    '• Portability: Runs on all modern browsers across desktop and mobile.\n'
    '• Reliability: Graceful handling of camera/microphone unavailability.\n'
    '• Maintainability: Modular Flask blueprint architecture.'
)

# Hardware Requirements
add_para("4.2 Hardware Requirements", bold=True, size=13)
hw_table = doc.add_table(rows=8, cols=2, style='Table Grid')
hw_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hw_data = [
    ("Component", "Specification"),
    ("Processor", "Intel i3 / AMD Ryzen 3 or above"),
    ("RAM", "8 GB minimum"),
    ("Hard Disk", "256 GB SSD"),
    ("Webcam", "720p resolution minimum"),
    ("Microphone", "Standard built-in or external"),
    ("Monitor", "1366×768 resolution minimum"),
    ("Network", "Broadband Internet connection"),
]
for i, (c1, c2) in enumerate(hw_data):
    hw_table.rows[i].cells[0].text = c1
    hw_table.rows[i].cells[1].text = c2
    if i == 0:
        for cell in hw_table.rows[i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

add_para("")

# Software Requirements
add_para("4.3 Software Requirements", bold=True, size=13)
sw_table = doc.add_table(rows=17, cols=2, style='Table Grid')
sw_table.alignment = WD_TABLE_ALIGNMENT.CENTER
sw_data = [
    ("Component", "Specification"),
    ("Operating System", "Windows 10/11, macOS, or Linux"),
    ("Programming Language", "Python 3.10+"),
    ("Web Framework", "Flask 3.x with Jinja2"),
    ("Frontend", "HTML5, CSS3, JavaScript ES6+"),
    ("ML Framework", "TensorFlow 2.15+, TensorFlow.js 4.x"),
    ("Pose Estimation", "MediaPipe Holistic 0.10+"),
    ("Speech Recognition", "SpeechRecognition + Google STT"),
    ("NLP", "spaCy 3.x (en_core_web_sm)"),
    ("Translation", "deep-translator (Google Translate)"),
    ("Text-to-Speech", "gTTS (Google TTS)"),
    ("3D Rendering", "Three.js r160+"),
    ("Database", "SQLite 3 with SQLAlchemy ORM"),
    ("Authentication", "Flask-Login + Werkzeug"),
    ("Real-time", "Flask-SocketIO + Eventlet"),
    ("Deployment", "Render (Cloud PaaS)"),
    ("Version Control", "Git + GitHub"),
]
for i, (c1, c2) in enumerate(sw_data):
    sw_table.rows[i].cells[0].text = c1
    sw_table.rows[i].cells[1].text = c2
    if i == 0:
        for cell in sw_table.rows[i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

add_para("")

add_para("4.4 Architecture", bold=True, size=13)
add_para(
    'The system follows a client-server architecture. The browser handles real-time inference (TF.js, MediaPipe, Three.js) '
    'while the Flask backend manages authentication, translation, grammar processing, and database operations.',
    size=11
)
arch_img = os.path.join(r"E:\sam\Project documents\diagrams", "deployment_diagram.png")
if os.path.exists(arch_img):
    doc.add_picture(arch_img, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_para("Figure 4.1: System Architecture — Client-Server Deployment", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

add_page_break()

# ═══════════════════════════════════════════════════════════════
# 5. SYSTEM DESIGN
# ═══════════════════════════════════════════════════════════════
add_heading_centered("5. SYSTEM DESIGN", level=1, size=16)

add_para("5.1 Input and Output Design", bold=True, size=13)
add_para(
    'Input Design: The system accepts input through two modalities — (1) Camera Input: live video frames processed '
    'by MediaPipe Holistic to extract 258D landmark features in sliding windows of 30 frames; (2) Voice/Text Input: '
    'audio recorded via microphone, converted to WAV using pydub, transcribed via Google STT, and tokenized with spaCy.'
)
add_para(
    'Output Design: Three output forms — (1) Text: recognized phrases displayed in the selected language; '
    '(2) Speech: gTTS audio played via browser; (3) Animation: Three.js 3D avatar performing ISL sign gestures '
    'with calibrated timing and fingerspelling for unrecognized words.'
)

add_para("5.2 UML Diagrams", bold=True, size=13)

def add_diagram(title, filename, caption):
    add_para(title, bold=True, size=12)
    img_path = os.path.join(DIAGRAM_DIR, filename)
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.8))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(caption, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para("")

add_diagram("5.2.1 Use Case Diagram", "use_case_diagram.png",
            "Figure 5.1: Use Case Diagram showing user interactions with Saṁvāk system modules")

add_diagram("5.2.2 Class Diagram", "class_diagram.png",
            "Figure 5.2: Class Diagram showing data model entities and their relationships")

add_diagram("5.2.3 Sequence Diagram — Sign-to-Text", "sequence_sign.png",
            "Figure 5.3: Sequence Diagram for Sign-to-Text translation flow (User → MediaPipe → LSTM-TCN → Flask → DB)")

add_diagram("5.2.4 Sequence Diagram — Speech-to-Sign", "sequence_speech.png",
            "Figure 5.4: Sequence Diagram for Speech-to-Sign translation flow (Voice → STT → Grammar → 3D Avatar)")

add_diagram("5.2.5 Deployment Diagram", "deployment_diagram.png",
            "Figure 5.5: Deployment Diagram showing client-server architecture with Render cloud")

add_diagram("5.2.6 Activity Diagram — Sign Recognition", "activity_diagram.png",
            "Figure 5.6: Activity Diagram showing the sign recognition decision flow with geometry fallback")

add_diagram("5.2.7 ER Diagram", "er_diagram.png",
            "Figure 5.7: ER Diagram showing database schema with Users, Translations, Preferences, and Progress")

add_para("5.3 DFD Diagrams", bold=True, size=13)

add_diagram("Level 0 — Context Diagram", "dfd_level0.png",
            "Figure 5.8: DFD Level 0 — Context diagram showing external entities and system boundary")

add_diagram("Level 1 — Major Processes", "dfd_level1.png",
            "Figure 5.9: DFD Level 1 — Four major process modules with data stores")

add_diagram("Level 2 — Sign-to-Text Detailed", "dfd_level2.png",
            "Figure 5.10: DFD Level 2 — Detailed breakdown of the Sign-to-Text module processing pipeline")

add_page_break()

# ═══════════════════════════════════════════════════════════════
# 6. CODE, IMPLEMENTATION AND RESULT
# ═══════════════════════════════════════════════════════════════
add_heading_centered("6. CODE, IMPLEMENTATION AND RESULT", level=1, size=16)

add_para("6.1 Gesture Recognition — MediaPipe Landmark Extraction (sign.py)", bold=True, size=12)
add_para(
    'import mediapipe as mp\n'
    'import numpy as np\n'
    'import base64\n'
    'from PIL import Image\n\n'
    'def _get_mp_hands_detector():\n'
    '    return mp.solutions.hands.Hands(\n'
    '        static_image_mode=True,\n'
    '        max_num_hands=2,\n'
    '        min_detection_confidence=0.45,\n'
    '    )\n\n'
    'def _decode_base64_image_to_rgb(image_value):\n'
    '    payload = str(image_value).strip()\n'
    '    if payload.startswith("data:") and "," in payload:\n'
    '        payload = payload.split(",", 1)[1]\n'
    '    image_bytes = base64.b64decode(payload)\n'
    '    image = Image.open(io.BytesIO(image_bytes)).convert("RGB")\n'
    '    return np.asarray(image, dtype=np.uint8)',
    size=10, align=WD_ALIGN_PARAGRAPH.LEFT
)

add_para("6.2 LSTM-TCN Model Architecture (train_lstm.py)", bold=True, size=12)
add_para(
    'from tensorflow.keras.layers import *\n'
    'from tensorflow.keras.models import Model\n\n'
    'SEQUENCE_LENGTH = 30  # 30 frames per window\n'
    'FEATURE_SIZE = 258    # 33*4 pose + 21*3*2 hands\n\n'
    'def residual_tcn_block(x, channels, dilation_rate, dropout):\n'
    '    shortcut = Conv1D(channels, 1, padding="same")(x)\n'
    '    y = Conv1D(channels, 3, padding="same", dilation_rate=dilation_rate)(x)\n'
    '    y = BatchNormalization()(y)\n'
    '    y = Activation("relu")(y)\n'
    '    y = Dropout(dropout)(y)\n'
    '    y = Add()([shortcut, y])\n'
    '    return Activation("relu")(y)\n\n'
    'def build_model(num_classes):\n'
    '    inputs = Input(shape=(30, 258))\n'
    '    x = Conv1D(128, 1, padding="same")(inputs)\n'
    '    for d in [1, 2, 4]:\n'
    '        x = residual_tcn_block(x, 128, d, 0.2)\n'
    '    x = GlobalAveragePooling1D()(x)\n'
    '    x = Dense(128, activation="relu")(x)\n'
    '    outputs = Dense(num_classes, activation="softmax")(x)\n'
    '    return Model(inputs, outputs)',
    size=10, align=WD_ALIGN_PARAGRAPH.LEFT
)

add_para("6.3 Speech Recognition (speech.py)", bold=True, size=12)
add_para(
    'import speech_recognition as sr\n'
    'from pydub import AudioSegment\n\n'
    'recognizer = sr.Recognizer()\n'
    'audio = AudioSegment.from_file(temp_path)\n'
    'audio.export(converted_path, format="wav")\n\n'
    'with sr.AudioFile(converted_path) as source:\n'
    '    audio_data = recognizer.record(source)\n'
    '    text = recognizer.recognize_google(audio_data, language="en-IN")',
    size=10, align=WD_ALIGN_PARAGRAPH.LEFT
)

add_para("6.4 ISL Grammar Engine (grammar_helper.py)", bold=True, size=12)
add_para(
    '_DROP_WORDS = {"a", "an", "the", "is", "am", "are", "was", "were"}\n'
    '_WH_WORDS = {"what", "where", "when", "why", "how", "who"}\n\n'
    'def _apply_isl_grammar(words):\n'
    '    time_markers, wh_words, content = [], [], []\n'
    '    for w in words:\n'
    '        if w in _DROP_WORDS: continue\n'
    '        elif w in _TIME_MARKERS: time_markers.append(w)\n'
    '        elif w in _WH_WORDS: wh_words.append(w)\n'
    '        else: content.append(w)\n'
    '    # OSV reorder + build result\n'
    '    return [w.upper() for w in time_markers + content + wh_words]\n\n'
    '# "What is your name?" → ["YOUR", "NAME", "WHAT"]',
    size=10, align=WD_ALIGN_PARAGRAPH.LEFT
)

add_para("6.5 Text-to-Speech (gTTS)", bold=True, size=12)
add_para(
    'from gtts import gTTS\n\n'
    'tts = gTTS(text="Thank You", lang="en")\n'
    'tts.save("output.mp3")',
    size=10, align=WD_ALIGN_PARAGRAPH.LEFT
)

add_empty_lines(1)
add_para("RESULT — Application Screenshots", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)

screenshots = [
    "Homepage — Landing page with Sign-to-Text and Speech-to-Sign navigation cards",
    "Sign-to-Text — Camera capturing hand gesture with real-time landmark overlay",
    "Sign-to-Text — Hand landmark detection using MediaPipe (21 keypoints per hand)",
    "Sign-to-Text — Recognized text output with confidence score and language translation",
    "Speech-to-Sign — Microphone input with transcribed text display",
    "Speech-to-Sign — 3D avatar performing ISL sign gestures",
    "ISL Dictionary — Browse interface showing sign image sequences",
    "Learn Module — Interactive practice with XP progress tracking",
    "Dashboard — User statistics (XP, level, words learned)",
    "Login / Register — Authentication pages",
]

# Try to insert actual screenshots if available
screenshot_files = {
    0: "homepage.png",
    1: "sign_to_text.png",
    4: "speech_to_sign.png",
    7: "learn.png",
    9: "login.png",
}

for idx, desc in enumerate(screenshots):
    img_file = screenshot_files.get(idx)
    if img_file:
        img_path = os.path.join(SCREENSHOT_DIR, img_file)
        if os.path.exists(img_path):
            try:
                doc.add_picture(img_path, width=Inches(5.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                add_para(f"[Insert Screenshot: {desc}]", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        else:
            add_para(f"[Insert Screenshot: {desc}]", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    else:
        add_para(f"[Insert Screenshot: {desc}]", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(f"Figure 6.{idx+1}: {desc}", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para("")

add_page_break()

# ═══════════════════════════════════════════════════════════════
# 7. SYSTEM STUDY AND TESTING
# ═══════════════════════════════════════════════════════════════
add_heading_centered("7. SYSTEM STUDY AND TESTING", level=1, size=16)

add_para("7.1 Feasibility Study", bold=True, size=13)
add_para(
    'Economical Feasibility: Uses exclusively open-source technologies and free-tier cloud services. No licensing '
    'fees or specialized equipment needed.'
)
add_para(
    'Technical Feasibility: Leverages well-established technologies — MediaPipe, TensorFlow, Flask. The iSign dataset '
    'provides sufficient training data for 40 ISL phrases. All dependencies available via pip/npm.'
)
add_para(
    'Social Feasibility: Directly addresses a critical social need for the hearing-impaired community. Web-based '
    'architecture ensures accessibility. Multilingual support (8 languages) enables broad adoption. The learning '
    'module promotes ISL literacy among hearing users.'
)

add_para("7.2 Types of Tests", bold=True, size=13)

add_para("7.2.1 Unit Testing", bold=True, size=12)
add_para(
    '• Grammar Engine: Verified english_to_isl_glosses() correctly transforms sentences.\n'
    '• Model Inference: Verified LSTM-TCN produces valid probabilities for (1, 30, 258) inputs.\n'
    '• Geometry Classifier: Verified static hand configurations classify correctly.\n'
    '• Database CRUD: Verified operations for User, Translation, UserProgress models.'
)

add_para("7.2.2 Integration Testing", bold=True, size=12)
add_para(
    '• Sign-to-Text: Frame → Landmarks → Prediction → Translation → Speech.\n'
    '• Speech-to-Sign: Audio → STT → Grammar → Avatar tokens.\n'
    '• Auth Flow: Register → Login → Session → Preferences → Logout.'
)

add_para("7.2.3 Functional Testing — Test Cases", bold=True, size=12)

# Test cases table
tc_table = doc.add_table(rows=16, cols=5, style='Table Grid')
tc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
tc_headers = ["Test ID", "Description", "Input", "Expected Output", "Status"]
for j, h in enumerate(tc_headers):
    tc_table.rows[0].cells[j].text = h
    for p in tc_table.rows[0].cells[j].paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)

tc_data = [
    ("TC-01", "Sign 'Thank You'", "Hand gesture", "Text: Thank You", "Pass"),
    ("TC-02", "Sign 'Hello'", "Wave gesture", "Text: Hello", "Pass"),
    ("TC-03", "Sign 'Yes'", "Thumbs up", "Text: Yes", "Pass"),
    ("TC-04", "Speech 'How are you'", "English voice", "Avatar: HOW YOU", "Pass"),
    ("TC-05", "Hindi speech input", "Hindi voice", "Avatar signs ISL", "Pass"),
    ("TC-06", "Telugu speech input", "Telugu voice", "Avatar signs ISL", "Pass"),
    ("TC-07", "User registration", "Valid credentials", "Account created", "Pass"),
    ("TC-08", "Duplicate email", "Existing email", "Error message", "Pass"),
    ("TC-09", "Invalid login", "Wrong password", "Error displayed", "Pass"),
    ("TC-10", "Translation history", "3 translations", "3 entries shown", "Pass"),
    ("TC-11", "Dictionary search", "Search 'hello'", "Frames displayed", "Pass"),
    ("TC-12", "XP tracking", "Learn exercise", "XP +10", "Pass"),
    ("TC-13", "Multilingual output", "Sign → Telugu", "Telugu text+speech", "Pass"),
    ("TC-14", "No hand detected", "Blank camera", "Status message", "Pass"),
    ("TC-15", "Low confidence", "Ambiguous sign", "No false positive", "Pass"),
]
for i, (tid, desc, inp, exp, status) in enumerate(tc_data):
    row = tc_table.rows[i + 1]
    row.cells[0].text = tid
    row.cells[1].text = desc
    row.cells[2].text = inp
    row.cells[3].text = exp
    row.cells[4].text = status
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

add_para("")

add_para("7.2.4 White Box Testing", bold=True, size=12)
add_para(
    'Tested all rule branches in ISL Grammar Engine (contraction expansion, article dropping, OSV reordering, '
    'WH repositioning). Verified threshold calibration logic and 258D feature projection from 390D MediaPipe output.'
)

add_para("7.2.5 Black Box Testing", bold=True, size=12)
add_para(
    '• All navigation links route correctly.\n'
    '• Camera/microphone permission requests appear appropriately.\n'
    '• Error messages display for invalid form submissions.\n'
    '• Responsive design renders on desktop, tablet, and mobile viewports.\n'
    '• Speech output plays in selected language.\n'
    '• 3D avatar gestures are smooth and correctly timed.'
)

add_page_break()

# ═══════════════════════════════════════════════════════════════
# 8. CONCLUSION
# ═══════════════════════════════════════════════════════════════
add_heading_centered("8. CONCLUSION", level=1, size=16)
add_para(
    'The Multi-Lingual Sign Language to Speech and Speech to Sign Translator (Samvak) successfully demonstrates a '
    'real-time, bidirectional, web-based translation system between Indian Sign Language and multiple spoken/written '
    'languages. The project achieves its core objectives:'
)
add_para(
    '1. Real-time sign recognition using LSTM-TCN trained on iSign, achieving 94.6% top-1 accuracy across 40 ISL phrases.\n'
    '2. Grammatically correct speech-to-sign translation using a local ISL grammar engine with OSV reordering.\n'
    '3. Multilingual support across 8 languages for both input and output.\n'
    '4. Accessible web deployment requiring only webcam and microphone, with in-browser TF.js inference.\n'
    '5. Interactive learning with ISL dictionary, gamified practice, and personalized dashboard.'
)
add_para(
    'The system integrates computer vision (MediaPipe), deep learning (TensorFlow/LSTM-TCN), NLP (spaCy, ISL grammar), '
    'speech processing (SpeechRecognition, gTTS), and 3D graphics (Three.js) into a cohesive, user-friendly application. '
    'The modular Flask blueprint architecture ensures maintainability and extensibility.'
)

add_page_break()

# ═══════════════════════════════════════════════════════════════
# 9. FUTURE ENHANCEMENT
# ═══════════════════════════════════════════════════════════════
add_heading_centered("9. FUTURE ENHANCEMENT", level=1, size=16)
add_para(
    '1. Continuous Sign Language Recognition (CSLR): Extend from isolated to sentence-level recognition using CTC.\n'
    '2. Expanded ISL Vocabulary: Increase beyond 40 phrases with additional training data and transfer learning.\n'
    '3. Facial Expression Integration: Incorporate eyebrow raises, mouth shapes as ISL grammatical markers.\n'
    '4. Regional ISL Dialect Support: Adapt to regional sign variations across Indian states.\n'
    '5. Photorealistic Avatar: Replace geometric avatar with neural-rendered or motion-captured signing.\n'
    '6. Offline Mode: Implement PWA with service workers for offline sign recognition.\n'
    '7. Video Conferencing Integration: Real-time translation during voice/video calls.\n'
    '8. Mobile Native App: Android/iOS apps with TensorFlow Lite for on-device inference.'
)

add_page_break()

# ═══════════════════════════════════════════════════════════════
# 10. REFERENCES
# ═══════════════════════════════════════════════════════════════
add_heading_centered("10. REFERENCES", level=1, size=16)

references = [
    '[1] Wadhawan, A. and Kumar, P. "Sign Language Recognition Systems: A Decade Systematic Literature Review." Archives of Computational Methods in Engineering, 28:785–813, 2021.',
    '[2] Lugaresi, C. et al. "MediaPipe: A Framework for Building Perception Pipelines." arXiv:1906.08172, 2019.',
    '[3] Hochreiter, S. and Schmidhuber, J. "Long Short-Term Memory." Neural Computation, 9(8):1735–1780, 1997.',
    '[4] Cui, R., Liu, H., and Zhang, C. "A Deep Neural Framework for Continuous Sign Language Recognition." IEEE Trans. Multimedia, 21(7):1880–1891, 2019.',
    '[5] Exploration Lab, IIT Kharagpur. "iSign: A Benchmark for Indian Sign Language Processing." ACL 2024.',
    '[6] Rastgoo, R., Kiani, K., and Escalera, S. "Sign Language Recognition: A Deep Survey." Expert Systems with Applications, 164:113794, 2021.',
    '[7] Saunders, B., Camgoz, N.C., and Bowden, R. "Progressive Transformers for End-to-End Sign Language Production." ECCV 2020.',
    '[8] Paudyal, P., Lee, J., and Banerjee, A. "A Survey on Sign Language Translation." ACM Computing Surveys, 55(3):1–37, 2023.',
    '[9] Zhang, F. et al. "MediaPipe Hands: On-device Real-time Hand Tracking." arXiv:2006.10214, 2020.',
    '[10] Abadi, M. et al. "TensorFlow: A System for Large-Scale Machine Learning." OSDI, pp. 265–283, 2016.',
]
for ref in references:
    add_para(ref, size=11)

# ── SAVE ──
output_path = os.path.join(os.path.dirname(__file__), "Samvak_ProjectReview_Final.docx")
doc.save(output_path)
print(f"SUCCESS: Document saved to {output_path}")
