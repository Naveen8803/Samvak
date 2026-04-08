import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document(r'E:\sam\Project documents\Samvak_Temp1.docx')
DIAGRAM_DIR = r"E:\sam\Project documents\diagrams"

def add_para(text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=13, space_before=None, space_after=None):
    p = doc.add_paragraph()
    p.alignment = align
    if space_before is not None: p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None: p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = 'Times New Roman'
    return p

def add_empty(n=1):
    for _ in range(n): doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# 3. SYSTEM ANALYSIS
# ═══════════════════════════════════════════════════════════════
add_para("3. SYSTEM ANALYSIS", True, WD_ALIGN_PARAGRAPH.LEFT, 16)
add_empty()

add_para("3.1 Existing System", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('Traditional existing implementations primarily center around heavy hardware setups including flex-sensor gloves and integrated Kinect depth sensors tracking spatial movement. Open source digital variants typically rely strictly on localized ASL patterns without comprehensive UI dashboards scaling to multilingual userbases.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("3.2 Disadvantages of Existing System", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('• High Hardware Cost: Reliance on specialized gloves reduces consumer accessibility.\n• Monolingual Bias: Output is natively routed into English without secondary language accommodations.\n• Grammatical Inconsistency: Naive text mappings ignore linguistic grammar transitions like ISL’s OSV format.\n• Static Availability: Systems lack portable web architectures compatible across generic browser structures.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("3.3 Proposed System", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('The system proposes a robust API-driven integration operating solely through regular webcams. Hand and multi-pose coordinates are aggregated directly via MediaPipe processes in JavaScript, feeding an optimal TCN model. Bidirectional conversion supports comprehensive accessibility through gamified interfaces and regional voice mappings.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("3.4 Advantages of Proposed System", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('• Platform Independence: Runs via browser engines utilizing TensorFlow.js without software installations.\n• Enhanced Accuracy: Deep Learning LSTM tracking accommodates temporal dynamic motions accurately natively over isolated static frames.\n• Syntactic Validity: Embeds spaCy linguistic processing mapping real-time syntax conversions.\n• Extended Portability: Deployed flexibly supporting immediate multilingual toggling targeting the diverse demographic of India.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("3.5 Workflow Explanation", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('The workflow delegates inference tasks between the UI and backend logic. Upon mode selection, video stream nodes map spatial vectors over 30-frame batches. Submitting this to the neural pipeline validates translations cascading back to the Google gTTS. Alternatively, dictation undergoes immediate parsing reconstructing vocabulary hierarchies before firing asynchronous events mapping dictionary libraries directly onto the skeletal bones of a 3D avatar rendering. This process is summarized visually herein.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

wf_img = os.path.join(DIAGRAM_DIR, "workflow_diagram.png")
if os.path.exists(wf_img):
    doc.add_picture(wf_img, width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_para("[Insert Diagram Screenshot Here]", True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 4. REQUIREMENT ANALYSIS
# ═══════════════════════════════════════════════════════════════
add_para("4. REQUIREMENT ANALYSIS", True, WD_ALIGN_PARAGRAPH.LEFT, 16)
add_empty()

add_para("4.1 Functional Requirements", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('1. The application must initialize local media tracking upon authorization without caching streams.\n2. The translation pipeline must classify static and temporal signs mapping outputs to multi-language bindings.\n3. User authentication tracking sessions tracking localized dictionary progress.\n4. Reordering voice captures through semantic dependencies eliminating stop words dynamically.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("4.2 Non-Functional Requirements", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('• Performance: Frame scaling must map inference checks under 200 milliseconds to preserve interactivity.\n• Reliability: Continuous operation must handle dropped packages adjusting to low bandwidth states visually gracefully.\n• Security: Data payloads (Passwords) hashed locally securing access via protected SQLite structures.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("4.3 Hardware Requirements", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('• Processor: Intel i3 / AMD Ryzen Class minimum\n• Memory: 8 GB Core RAM\n• Storage: 256 GB SSD allocation\n• Peripherals: 720p minimum Webcam capability, Internal/External Microphone.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("4.4 Software Requirements", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('• Environment: Python 3.10+, Node architectures enabled via HTML5 wrappers.\n• Frameworks: Flask 3.x, TensorFlow.js, Jinja2 template handling.\n• External Modules: MediaPipe Holistic, spaCy NLP binaries, gTTS routing, Three.js internal render engines.\n• Database Controller: SQLite 3 bindings scaling over SQLAlchemy ORM frameworks.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("4.5 System Architecture", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('The logical architectural configuration operates across three tiers. The client layer manages data capturing (Media devices) and WebGL visualizations (Three.js bounding boxes). The intermediate application logic deployed in Python handles mathematical routing mapping TCN predictions over Flask REST layers. The trailing persistence tier leverages rapid relational querying to log authenticated history operations maintaining high modular segregation ensuring code reusability.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 5. SYSTEM DESIGN
# ═══════════════════════════════════════════════════════════════
add_para("5. SYSTEM DESIGN", True, WD_ALIGN_PARAGRAPH.LEFT, 16)
add_empty()

diagrams = [
    ("5.1 Use Case Diagram", "use_case_diagram.png", "The Use Case diagram demarcates actor interaction boundaries encompassing the operational domain. The primary Actor initiates Login authentication, bridging access into dynamic endpoints like Sign-to-Text inference, Avatar Translation modules, or the Gamified ISL Dictionary logic structure."),
    ("5.2 Class Diagram", "class_diagram.png", "Detailed UML structure outlining precise object-oriented relationships between internal structures. Displays properties mapping User DB schemas, Tokenization handlers bridging Speech Recognition hierarchies, and the rendering logic mapping coordinate matrices scaling inside the Model Engine interfaces."),
    ("5.3 Sequence Diagram", "sequence_sign.png", "Highlights synchronous message broadcasting executing chronologically between User, Web UI wrapper, TensorFlow processor logic, and Backend translation formatting APIs validating payload returns."),
    ("5.4 Activity Diagram", "activity_diagram.png", "Plots the conditional decision matrix logic mapping branching operations like validating sign thresholds, adjusting spatial errors, and routing fallback actions depending upon visual coordinate occlusion parameters."),
    ("5.5 Component Diagram", "deployment_diagram.png", "Identifies macro logical building blocks compiling overarching dependencies. Outlines the Web Interface Node, Authentication Component, AI Extraction Library block, and Language Syntax logic processing interacting over HTTPS bindings. (Substituted with Deployment scope visual)."),
    ("5.6 Deployment Diagram", "deployment_diagram.png", "Models physical infrastructure topology. Summarizes client devices mapping to Web Application Firewalls cascading rendering tasks over Cloud hosting (Render) housing the Flask/SQlite ecosystem."),
    ("5.7 ER Diagram", "er_diagram.png", "The Entity-Relationship model illustrates the data persistence layer structures prioritizing associative mapping linking Student/User tracking metrics targeting individual session History keys and Dictionary Point allocations."),
    ("5.8 Data Flow Diagram", "dfd_level0.png", "The Level 0 scaling (Context Diagram) and subsequent Level 1 & Level 2 branches decompose process flows tracking user interactions passing raw multimedia packages across transformation functions standardizing final localized strings.")
]

for title, file, desc in diagrams:
    add_para(title, True, WD_ALIGN_PARAGRAPH.LEFT, 13)
    add_para(desc, False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
    
    img_path = os.path.join(DIAGRAM_DIR, file)
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para("[Insert Diagram Screenshot Here]", True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    add_empty()

doc.save(r'E:\sam\Project documents\Samvak_Temp2.docx')
