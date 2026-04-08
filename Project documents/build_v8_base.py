import os
from docx import Document
from docx.shared import Pt, Inches

ref_doc_path = r'E:\sam\Project documents\ProjectDoc.docx'
doc = Document(ref_doc_path)

def replace_paragraph_text(p, old_text, new_text, make_bold=False):
    if old_text in p.text:
        is_bold = make_bold
        font_size = None
        font_name = 'Times New Roman'
        
        for r in p.runs:
            if r.bold is not None: is_bold = r.bold
            if r.font.size is not None: font_size = r.font.size
            if r.font.name is not None: font_name = r.font.name
            
        full_text = p.text.replace(old_text, new_text)
        p.clear()
        
        new_run = p.add_run(full_text)
        new_run.bold = is_bold
        new_run.font.name = font_name
        if font_size is not None:
            new_run.font.size = font_size

ack_index = -1
for i, p in enumerate(doc.paragraphs):
    text = p.text
    if 'ATTRIBUTE-BASED DATA SHARING SCHEME REVISITED IN CLOUD COMPUTING' in text:
        replace_paragraph_text(p, 'ATTRIBUTE-BASED DATA SHARING SCHEME REVISITED IN CLOUD COMPUTING', 'MULTI-LINGUAL SIGN LANGUAGE TO SPEECH AND SPEECH TO SIGN TRANSLATOR', True)
    elif 'ATTRIBUTE' in text and 'DATA SHARING SCHEME' in text:
        replace_paragraph_text(p, text, text.replace('ATTRIBUTE-BASED DATA SHARING SCHEME REVISITED IN CLOUD COMPUTING', 'MULTI-LINGUAL SIGN LANGUAGE TO SPEECH AND SPEECH TO SIGN TRANSLATOR'), True)
    elif 'ATTRIBUTE' in text:
        if 'This is to certify that the Project' in text:
            p.text = 'This is to certify that the Project Report entitled "MULTI-LINGUAL SIGN LANGUAGE TO SPEECH AND SPEECH TO SIGN TRANSLATOR" is a bonafide record of the work done by Uppalapati Naveen Varma (2401600155) in partial fulfillment of the requirements for the award of the degree of Master of Computer Applications from the Department of Computer Science and Applications, Koneru Lakshmaiah Education Foundation, Vaddeswaram, during the academic year 2025–26.'
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(13)
        else:
            replace_paragraph_text(p, 'ATTRIBUTE', 'MULTI-LINGUAL SIGN LANGUAGE')
    if 'Shaik Salim' in text or 'shaik salim' in text.lower():
        # Avoid breaking if case is weird
        full_text = p.text.replace('Shaik Salim', 'Uppalapati Naveen Varma').replace('shaik salim', 'Uppalapati Naveen Varma')
        p.clear()
        r = p.add_run(full_text)
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(13)
    if '2201600096' in text:
        replace_paragraph_text(p, '2201600096', '2401600155', True)
        if 'ACKNOWLEDGEMENT' in [prev.text for prev in doc.paragraphs[max(0, i-20):i]]:
            ack_index = i
            
    if 'Dr.R D Sathiya' in text or 'Dr.R.D Sathiya' in text or 'Dr.R. D Sathiya' in text:
        p.text = p.text.replace('Dr.R D Sathiya', 'Mrs. Swathi Voddi').replace('Dr.R.D Sathiya', 'Mrs. Swathi Voddi').replace('Dr.R. D Sathiya', 'Mrs. Swathi Voddi')
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(13)
            run.font.name = 'Times New Roman'
            
    if '2023- 24' in text:
        replace_paragraph_text(p, '2023- 24', '2025- 26')
    if '2023 – 24' in text:
        replace_paragraph_text(p, '2023 – 24', '2025 – 26')

# Remove everything after Acknowledgement Student ID line
if ack_index > 0:
    for p in doc.paragraphs[ack_index+1:]:
        p._element.getparent().remove(p._element)
    for t in doc.tables:
        t._element.getparent().remove(t._element)

# Now doc has exactly the 4 front pages, beautifully formatted.
# Save it as base.
doc.save(r'E:\sam\Project documents\Samvak_Base.docx')
print("Base front pages generated successfully.")
