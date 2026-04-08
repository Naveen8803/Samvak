from docx import Document
path = r'E:\sam\Project documents\Samvak_Review_Final_V13b.docx'
doc = Document(path)

all_text = '\n'.join(p.text for p in doc.paragraphs)
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            all_text += '\n' + cell.text

checks = ['UT-01','IT-01','142 ms','PASS','Performance']
print("=== KEYWORD CHECK (paragraphs + table cells) ===")
for c in checks:
    print(f"  [{'OK' if c in all_text else 'MISSING'}] {c}")

print()
print("=== TABLE BREAKDOWN ===")
for i, t in enumerate(doc.tables):
    print(f"Table {i+1}: {len(t.rows)} rows x {len(t.columns)} cols")
    for j, row in enumerate(t.rows):
        cells = [row.cells[k].text[:25] for k in range(len(row.cells))]
        print(f"  Row {j}: {cells}")
    print()
