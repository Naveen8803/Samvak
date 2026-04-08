import os
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document(r'E:\sam\Project documents\Samvak_Base_Clean_V2.docx')

DIAGRAM_DIR = r"E:\sam\Project documents\diagrams"
SCREENSHOT_DIR = r"E:\sam\Project documents\screenshots"

def add_para(text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=13, space_before=None, space_after=None):
    p = doc.add_paragraph()
    p.alignment = align
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = 'Times New Roman'
    return p

def add_empty(n=1):
    for _ in range(n):
        doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════
# ABSTRACT
# ═══════════════════════════════════════════════════════════════
add_empty()
add_para('ABSTRACT', True, WD_ALIGN_PARAGRAPH.CENTER, 14)
add_empty()
add_para('This project addresses the critical communication gap between the hearing-impaired community and the general population through the design and implementation of a bidirectional Multi-Lingual Sign Language to Speech and Speech to Sign Translator. The primary problem lies in the fact that sign language, particularly Indian Sign Language (ISL), has its own distinct grammar and syntax, making conventional translation difficult, while most assistive solutions remain hardware-dependent or monolingual.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
add_para('As a comprehensive solution, an AI-based, hardware-free web application was developed. The system functions in two dynamic modes: Sign-to-Text uses MediaPipe Holistic for extracting 258-dimensional bodily landmarks, mapping these onto a robust LSTM and Temporal Convolutional Network (TCN) architecture to classify ISL gestures. Conversely, Speech-to-Sign leverages Google Speech-to-Text (STT) and deep natural language processing (NLP) via spaCy to restructure incoming spoken phrases into proper ISL Object-Subject-Verb (OSV) gloss grammar sequences, which are subsequently animated in real-time by a Three.js 3D web avatar.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
add_para('The ultimate outcome is an accessible, real-time platform capable of supporting multi-lingual input and output across English, Hindi, Telugu, Tamil, Kannada, Malayalam, Spanish, and French, thereby eliminating reliance on human interpreters and delivering high precision, scalable gesture-to-speech communication.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════
add_para('TABLE OF CONTENTS', True, WD_ALIGN_PARAGRAPH.CENTER, 14)
add_empty()
toc_data = [
    ("CHAPTER NO.", "TITLE", "PAGE NO."),
    ("1", "INTRODUCTION", "1"),
    ("1.1", "Motivation", "1"),
    ("1.2", "Problem Statement", "1"),
    ("1.3", "Objectives of the Project", "2"),
    ("1.4", "Scope", "2"),
    ("1.5", "Overview of the System", "2"),
    ("2", "LITERATURE SURVEY", "3"),
    ("3", "SYSTEM ANALYSIS", "5"),
    ("3.1", "Existing System", "5"),
    ("3.2", "Disadvantages of Existing System", "5"),
    ("3.3", "Proposed System", "6"),
    ("3.4", "Advantages of Proposed System", "6"),
    ("3.5", "Workflow Explanation", "6"),
    ("4", "REQUIREMENT ANALYSIS", "8"),
    ("4.1", "Functional Requirements", "8"),
    ("4.2", "Non-Functional Requirements", "8"),
    ("4.3", "Hardware Requirements", "9"),
    ("4.4", "Software Requirements", "9"),
    ("4.5", "System Architecture", "10"),
    ("5", "SYSTEM DESIGN", "11"),
    ("5.1", "Use Case Diagram", "11"),
    ("5.2", "Class Diagram", "12"),
    ("5.3", "Sequence Diagram", "13"),
    ("5.4", "Activity Diagram", "14"),
    ("5.5", "Component Diagram", "15"),
    ("5.6", "Deployment Diagram", "15"),
    ("5.7", "ER Diagram", "16"),
    ("5.8", "Data Flow Diagram", "16"),
    ("6", "CODE, IMPLEMENTATION AND RESULTS", "18"),
    ("6.1", "Module Explanation", "18"),
    ("6.2", "Sample Code Snippets", "20"),
    ("6.3", "Output Screens", "22"),
    ("7", "SYSTEM STUDY AND TESTING", "24"),
    ("7.1", "Feasibility Study", "24"),
    ("7.2", "Types of Testing", "24"),
    ("7.3", "Test Cases", "25"),
    ("8", "CONCLUSION", "26"),
    ("9", "FUTURE ENHANCEMENT", "27"),
    ("10", "REFERENCES", "28")
]
toc_table = doc.add_table(rows=len(toc_data), cols=3)
toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (c1, c2, c3) in enumerate(toc_data):
    toc_table.rows[i].cells[0].text = c1
    toc_table.rows[i].cells[1].text = c2
    toc_table.rows[i].cells[2].text = c3
    for cell in toc_table.rows[i].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
                if i == 0 or len(c1) == 1 or c1 == "":  
                    r.bold = True
                else:
                    r.bold = False

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ═══════════════════════════════════════════════════════════════
add_para("1. INTRODUCTION", True, WD_ALIGN_PARAGRAPH.LEFT, 16)
add_empty()

add_para("1.1 Motivation", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('According to the World Health Organization, millions worldwide suffer from disabling hearing loss, creating severe communication barriers between Deaf individuals and the hearing community. This disconnect often restricts access to essential educational, healthcare, and professional opportunities. Because the vast majority of people lack understanding of sign language, an artificial intelligence tool bridging sign language and spoken language is heavily motivated to establish autonomous social inclusion.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("1.2 Problem Statement", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('The primary problem is the lack of accessible, affordable, and accurate software interpreters that convert Indian Sign Language (ISL) to spoken words and vice versa. While ASL (American Sign Language) models exist, ISL uses distinctly different vocabulary and an Object-Subject-Verb (OSV) grammatical structure. Furthermore, current systems often rely on expensive sensor-based hardware and rarely provide multilingual support outside of English.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("1.3 Objectives of the Project", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('1. To engineer a computer-vision neural network employing MediaPipe and LSTM layers to achieve real-time classification of live ISL gestures.\n2. To design an NLP-based text restructuring engine to enforce local ISL grammar over transcribed speech.\n3. To render animated translations via a 3D WebGL avatar.\n4. To support scalable multilingual features translating logic simultaneously into multiple linguistic outputs (Hindi, Telugu, etc.).', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("1.4 Scope", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('The system’s scope includes rendering isolated 40 ISL gloss classifications directly into standard strings through an accessible web camera. Furthermore, it scales audio capturing through native Web Speech APIs transforming dictation seamlessly into 3D animations in real time. Registration tracking, learning libraries, and gamified point architectures are also within the system’s boundary.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("1.5 Overview of the System", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('The platform represents a bidirectional bridge housed within a Flask python backend. A user gestures into the camera, where MediaPipe isolates hand and facial landmarks. An embedded temporal sequential model validates the sign and propagates it via Google TTS APIs as regional speech. Conversely, when users dictate speech, Python spaCy categorizes syntax, reorganizing verbs and subjects accurately into ISL arrays streamed to an interactive 3D humanoid for visual interpretation.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 2. LITERATURE SURVEY
# ═══════════════════════════════════════════════════════════════
add_para("2. LITERATURE SURVEY", True, WD_ALIGN_PARAGRAPH.LEFT, 16)
add_empty()

refs = [
    ('[1] Sign Language Recognition Systems: A Decade Systematic Literature Review', 'By using Vision-based and Sensor-based tracking methodologies, the survey indicated that deep learning CNNs perform remarkably well. However, limitation exists whereby systems isolate only ASL structures causing severe constraints on geographic inclusivity.'),
    ('[2] MediaPipe: A Framework for Building Perception Pipelines', 'MediaPipe utilizes tracking pipelines enabling high-fidelity coordinate inference. It operates independently of complex GPU architectures for 33-node body graphs, which forms the core extraction basis for feature scaling within the Samvak framework.'),
    ('[3] Continuous Sign Language Recognition via Deep Neural Networks', 'Research evaluated combining Hidden Markov Models (HMM) with spatial convolutional blocks for predicting complex sequential actions. This illustrated the dependency limitations of HMMs which motivated the adoption of advanced LSTM nodes.'),
    ('[4] A Survey on Sign Language Translation: Neural and Rule-based Approaches', 'Analyzes transformation pipelines scaling standard NLP systems for gloss grammar generation. Rule-based linguistic constraints outperformed generative approaches for low-resource architectures, justifying our OSV grammar engine.'),
    ('[5] Progressive Transformers for Sign Language Production', 'Explores deep generative rendering of avatar skeletons mapping to gloss sequences. The limitations highlighted high latency costs, leading this project to adopt streamlined WebGL skeletal animations over generative models.'),
    ('[6] OpenPose: Realtime Multi-Person 2D Pose Estimation', 'Demonstrates bottom-up tracking technologies. However, the system requires extensive computing bandwidth, proving why MediaPipe’s lightweight browser-side capability is more feasible for a client-driven translator module.'),
    ('[7] Indian Sign Language Database and Recognition', 'Discusses the foundational dataset structure applied specifically for Indian Sign Language. The lack of temporal video sets highlighted the necessity of training an overarching LSTM over the locally processed iSign frame sets.')
]

for title, desc in refs:
    add_para(title, True, size=13)
    add_para(desc, size=13, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_empty()

doc.add_page_break()

# Save intermediary to check everything is correct
doc.save(r'E:\sam\Project documents\Samvak_Temp1.docx')
