import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Load the physically extracted and replaced reference Front Pages
doc = Document(r'E:\sam\Project documents\Samvak_FrontPages.docx')

DIAGRAM_DIR = r"E:\sam\Project documents\diagrams"
SCREENSHOT_DIR = r"E:\sam\Project documents\screenshots"

def add_para(text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=13, space_before=None, space_after=None):
    p = doc.add_paragraph()
    p.alignment = align
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = 'Times New Roman'
    return p

def add_empty(n=1):
    for _ in range(n):
        doc.add_paragraph()

doc.add_page_break()

# --- Page 5: Abstract ---
add_empty(2)
add_para('ABSTRACT', True, WD_ALIGN_PARAGRAPH.CENTER, 14)
add_empty(1)
add_para('This project presents the design and implementation of Samvak — a Multi-Lingual Sign Language to Speech and Speech to Sign Translator, a real-time, web-based application that bridges the communication gap between the hearing-impaired community and the general population. The system operates in two primary modes: (1) Sign-to-Text/Speech, where a webcam captures hand gestures performed in Indian Sign Language (ISL), processes them through a deep learning pipeline using MediaPipe for landmark extraction and an LSTM-based Temporal Convolutional Network (TCN) for classification, and produces text and spoken output in eight languages; and (2) Speech-to-Sign, where spoken or typed input is transcribed, translated into ISL gloss sequences using rule-based grammar transformation, and rendered as sign animations through a real-time 3D avatar in the browser.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
add_para('The application is built on a Flask backend with a responsive HTML/CSS/JavaScript frontend. MediaPipe Holistic provides 258-dimensional pose and hand feature vectors from each video frame, which are fed into a trained LSTM-TCN model capable of recognizing 40 ISL phrases with 94.6% top-1 accuracy. For speech-to-sign translation, the system employs the SpeechRecognition library for automatic speech recognition, spaCy for natural language tokenization, and a local ISL grammar engine that applies Object-Subject-Verb (OSV) reordering, article dropping, and WH-word repositioning to generate correct ISL gloss sequences. A 3D humanoid avatar, rendered using Three.js, performs the corresponding sign gestures in real time.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
add_para('The system supports multilingual output in English, Hindi, Telugu, Tamil, Kannada, Malayalam, Spanish, and French using the Google Translate API and gTTS for text-to-speech synthesis. Additional features include an interactive ISL dictionary, a gamified learning module with XP-based progress tracking, user authentication, a personalized dashboard, and translation history logging. The application is deployed on Render for cloud accessibility.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
add_empty(1)
p = add_para('Keywords: ', True, WD_ALIGN_PARAGRAPH.LEFT, 13)
p.add_run('Indian Sign Language, gesture recognition, speech recognition, MediaPipe, LSTM, real-time translation, 3D avatar, multilingual, accessibility').bold = False
doc.add_page_break()

# --- Page 6: TABLE OF CONTENTS ---
add_para('TABLE OF CONTENTS', True, WD_ALIGN_PARAGRAPH.CENTER, 14)
add_empty(1)
toc_data = [
    ("CHAPTER NO.", "TITLE", "PAGE NO."),
    ("1", "INTRODUCTION", "1"),
    ("1.1", "Motivation", "1"),
    ("1.2", "Problem Statement", "1"),
    ("1.3", "Objective of the Project", "2"),
    ("1.4", "Scope", "2"),
    ("1.5", "Project Introduction", "2"),
    ("2", "LITERATURE SURVEY", "3"),
    ("2.1", "Related Work", "3"),
    ("3", "SYSTEM ANALYSIS", "5"),
    ("3.1", "Existing System", "5"),
    ("3.2", "Disadvantages of Existing System", "5"),
    ("3.3", "Proposed System", "5"),
    ("3.4", "Advantages of Proposed System", "6"),
    ("3.5", "Workflow of Proposed System", "6"),
    ("4", "REQUIREMENT ANALYSIS", "8"),
    ("4.1", "Functional and Non-Functional Requirements", "8"),
    ("4.2", "Hardware Requirements", "8"),
    ("4.3", "Software Requirements", "9"),
    ("4.4", "Architecture", "9"),
    ("5", "SYSTEM DESIGN", "10"),
    ("5.1", "Input and Output Design", "10"),
    ("5.2", "UML Diagrams", "11"),
    ("5.3", "DFD Diagrams", "14"),
    ("6", "CODE, IMPLEMENTATION AND RESULT", "16"),
    ("7", "SYSTEM STUDY AND TESTING", "26"),
    ("7.1", "Feasibility Study", "26"),
    ("7.2", "System Testing", "27"),
    ("7.3", "Types of Tests", "27"),
    ("8", "CONCLUSION", "29"),
    ("9", "FUTURE ENHANCEMENT", "30"),
    ("", "REFERENCES", "31")
]
toc_table = doc.add_table(rows=len(toc_data), cols=3)
toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (c1, c2, c3) in enumerate(toc_data):
    toc_table.rows[i].cells[0].text = c1
    toc_table.rows[i].cells[1].text = c2
    toc_table.rows[i].cells[2].text = c3
    for cell in toc_table.rows[i].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
                if i == 0 or len(c1) == 1 or c1 == "":  
                    r.bold = True
                else:
                    r.bold = False

doc.add_page_break()
