from docx import Document

path = r'E:\sam\Project documents\Samvak_Review_Final_V13b.docx'
doc = Document(path)

# ── 1. All headings (14pt+ bold) ─────────────────────────────────────────────
print("=" * 60)
print("ALL HEADINGS (14pt+ bold)")
print("=" * 60)
for p in doc.paragraphs:
    t = p.text.strip()
    if not t or not p.runs:
        continue
    r = p.runs[0]
    sz = r.font.size.pt if r.font.size else None
    if sz and sz >= 14 and r.bold:
        print(f"  [{int(sz)}pt] {t[:80]}")

# ── 2. Code blocks (Courier New) ─────────────────────────────────────────────
print()
print("=" * 60)
print("CODE BLOCKS (Courier New paragraphs)")
print("=" * 60)
code_paras = [p for p in doc.paragraphs
              if p.runs and p.runs[0].font.name == 'Courier New']
print(f"Total code blocks: {len(code_paras)}")
for i, p in enumerate(code_paras):
    first_line = p.text.strip().split('\n')[0][:70]
    print(f"  CB{i+1}: {first_line}")

# ── 3. Sub-sections 6.1 – 6.8 ────────────────────────────────────────────────
print()
print("=" * 60)
print("CHAPTER 6 SUB-SECTIONS CHECK")
print("=" * 60)
required = ['6.1', '6.2', '6.3', '6.4', '6.5', '6.6', '6.7', '6.8']
found_subs = {k: False for k in required}
for p in doc.paragraphs:
    t = p.text.strip()
    for k in required:
        if t.startswith(k):
            found_subs[k] = True
for k, v in found_subs.items():
    print(f"  {k}: {'FOUND ✓' if v else 'MISSING ✗'}")

# ── 4. All tables ────────────────────────────────────────────────────────────
print()
print("=" * 60)
print("TABLES")
print("=" * 60)
for i, t in enumerate(doc.tables):
    hdr = [c.text[:25] for c in t.rows[0].cells] if t.rows else []
    print(f"  Table {i+1}: {len(t.rows)} rows × {len(t.columns)} cols")
    print(f"    Header: {hdr}")

# ── 5. Images embedded ───────────────────────────────────────────────────────
print()
print("=" * 60)
print(f"IMAGES: {len([r for r in doc.part.rels.values() if 'image' in r.reltype])}")
print("=" * 60)

# ── 6. Key content checks ────────────────────────────────────────────────────
print()
print("=" * 60)
print("KEY CONTENT CHECKS")
print("=" * 60)
full_text = "\n".join(p.text for p in doc.paragraphs)
checks = {
    "Abstract (94.6% accuracy)":           "94.6%" in full_text,
    "Abstract (258 features)":             "258" in full_text,
    "Abstract (30 frames)":                "30 frames" in full_text,
    "Literature Survey (8 papers)":        "[8]" in full_text,
    "iSign dataset mentioned":             "iSign" in full_text,
    "MediaPipe Holistic mentioned":        "MediaPipe" in full_text,
    "Gemini API mentioned":                "Gemini" in full_text,
    "TensorFlow/Keras LSTM":              "LSTM" in full_text,
    "Flask blueprint":                     "Blueprint" in full_text,
    "grammar_helper.py code":             "english_to_isl_glosses" in full_text,
    "sign.js code (extractFeatures)":     "extractFeatures" in full_text,
    "LSTM model (build_lstm_model)":      "build_lstm_model" in full_text,
    "isign_retrieval code":               "sequence_to_embedding" in full_text,
    "models.py code (UserProgress)":      "UserProgress" in full_text,
    "Feasibility study":                  "Feasibility" in full_text,
    "Test cases table":                   "UT-01" in full_text,
    "Integration test table":             "IT-01" in full_text,
    "Performance table":                  "142 ms" in full_text,
    "10 future enhancements":             "Federated Learning" in full_text,
    "14 references":                      "[14]" in full_text,
}
for desc, result in checks.items():
    print(f"  {'✓' if result else '✗'} {desc}")
