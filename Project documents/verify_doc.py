import os
from docx import Document
from lxml import etree

path = r'E:\sam\Project documents\Samvak_ProjectReview_Final.docx'
size = os.path.getsize(path)
d = Document(path)
imgs = list(d.inline_shapes)
tables = len(d.tables)
paras = len(d.paragraphs)

ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
borders = d.sections[0]._sectPr.findall('.//w:pgBorders', ns)
has_border = len(borders) > 0

print(f'File size: {size//1024} KB')
print(f'Paragraphs: {paras}')
print(f'Tables: {tables}')
print(f'Embedded images: {len(imgs)}')
print(f'Page border: {"YES" if has_border else "NO"}')
print(f'Sections: {len(d.sections)}')

for i, img in enumerate(imgs):
    w_in = img.width / 914400
    h_in = img.height / 914400
    print(f'  Image {i}: {w_in:.1f}" x {h_in:.1f}"')
