import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Load the physically extracted and replaced reference Front Pages
doc = Document(r'E:\sam\Project documents\Samvak_FrontPages.docx')

DIAGRAM_DIR = r"E:\sam\Project documents\diagrams"
SCREENSHOT_DIR = r"E:\sam\Project documents\screenshots"

def add_para(text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=13, space_before=None, space_after=None):
    p = doc.add_paragraph()
    p.alignment = align
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = 'Times New Roman'
    return p

def add_empty(n=1):
    for _ in range(n):
        doc.add_paragraph()

doc.add_page_break()

# --- Page 5: Abstract ---
add_empty(2)
add_para('ABSTRACT', True, WD_ALIGN_PARAGRAPH.CENTER, 14)
add_empty(1)
add_para('This project presents the design and implementation of Samvak — a Multi-Lingual Sign Language to Speech and Speech to Sign Translator, a real-time, web-based application that bridges the communication gap between the hearing-impaired community and the general population. The system operates in two primary modes: (1) Sign-to-Text/Speech, where a webcam captures hand gestures performed in Indian Sign Language (ISL), processes them through a deep learning pipeline using MediaPipe for landmark extraction and an LSTM-based Temporal Convolutional Network (TCN) for classification, and produces text and spoken output in eight languages; and (2) Speech-to-Sign, where spoken or typed input is transcribed, translated into ISL gloss sequences using rule-based grammar transformation, and rendered as sign animations through a real-time 3D avatar in the browser.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
add_para('The application is built on a Flask backend with a responsive HTML/CSS/JavaScript frontend. MediaPipe Holistic provides 258-dimensional pose and hand feature vectors from each video frame, which are fed into a trained LSTM-TCN model capable of recognizing 40 ISL phrases with 94.6% top-1 accuracy. For speech-to-sign translation, the system employs the SpeechRecognition library for automatic speech recognition, spaCy for natural language tokenization, and a local ISL grammar engine that applies Object-Subject-Verb (OSV) reordering, article dropping, and WH-word repositioning to generate correct ISL gloss sequences. A 3D humanoid avatar, rendered using Three.js, performs the corresponding sign gestures in real time.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
add_para('The system supports multilingual output in English, Hindi, Telugu, Tamil, Kannada, Malayalam, Spanish, and French using the Google Translate API and gTTS for text-to-speech synthesis. Additional features include an interactive ISL dictionary, a gamified learning module with XP-based progress tracking, user authentication, a personalized dashboard, and translation history logging. The application is deployed on Render for cloud accessibility.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
add_empty(1)
p = add_para('Keywords: ', True, WD_ALIGN_PARAGRAPH.LEFT, 13)
p.add_run('Indian Sign Language, gesture recognition, speech recognition, MediaPipe, LSTM, real-time translation, 3D avatar, multilingual, accessibility').bold = False
doc.add_page_break()

# --- Page 6: TABLE OF CONTENTS ---
add_para('TABLE OF CONTENTS', True, WD_ALIGN_PARAGRAPH.CENTER, 14)
add_empty(1)
toc_data = [
    ("CHAPTER NO.", "TITLE", "PAGE NO."),
    ("1", "INTRODUCTION", "1"),
    ("1.1", "Motivation", "1"),
    ("1.2", "Problem Statement", "1"),
    ("1.3", "Objective of the Project", "2"),
    ("1.4", "Scope", "2"),
    ("1.5", "Project Introduction", "2"),
    ("2", "LITERATURE SURVEY", "3"),
    ("2.1", "Related Work", "3"),
    ("3", "SYSTEM ANALYSIS", "5"),
    ("3.1", "Existing System", "5"),
    ("3.2", "Disadvantages of Existing System", "5"),
    ("3.3", "Proposed System", "5"),
    ("3.4", "Advantages of Proposed System", "6"),
    ("3.5", "Workflow of Proposed System", "6"),
    ("4", "REQUIREMENT ANALYSIS", "8"),
    ("4.1", "Functional and Non-Functional Requirements", "8"),
    ("4.2", "Hardware Requirements", "8"),
    ("4.3", "Software Requirements", "9"),
    ("4.4", "Architecture", "9"),
    ("5", "SYSTEM DESIGN", "10"),
    ("5.1", "Input and Output Design", "10"),
    ("5.2", "UML Diagrams", "11"),
    ("5.3", "DFD Diagrams", "14"),
    ("6", "CODE, IMPLEMENTATION AND RESULT", "16"),
    ("7", "SYSTEM STUDY AND TESTING", "26"),
    ("7.1", "Feasibility Study", "26"),
    ("7.2", "System Testing", "27"),
    ("7.3", "Types of Tests", "27"),
    ("8", "CONCLUSION", "29"),
    ("9", "FUTURE ENHANCEMENT", "30"),
    ("", "REFERENCES", "31")
]
toc_table = doc.add_table(rows=len(toc_data), cols=3)
toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (c1, c2, c3) in enumerate(toc_data):
    toc_table.rows[i].cells[0].text = c1
    toc_table.rows[i].cells[1].text = c2
    toc_table.rows[i].cells[2].text = c3
    for cell in toc_table.rows[i].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
                if i == 0 or len(c1) == 1 or c1 == "":  
                    r.bold = True
                else:
                    r.bold = False

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ═══════════════════════════════════════════════════════════════
add_para("1. INTRODUCTION", True, WD_ALIGN_PARAGRAPH.LEFT, 16)
add_empty()

add_para("1.1 Motivation", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('According to the World Health Organization, over 430 million people worldwide suffer from disabling hearing loss, with approximately 63 million individuals in India alone classified as hearing-impaired. The primary mode of communication for the Deaf and hard-of-hearing community is sign language, yet the vast majority of the hearing population does not understand sign language. This communication barrier leads to social isolation, limited access to education, healthcare, and employment opportunities for the hearing-impaired.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
add_para('While interpreters serve as intermediaries, their availability is severely limited, especially in rural India. Moreover, Indian Sign Language (ISL) has its own grammar, syntax, and vocabulary that differ significantly from spoken languages, making direct translation non-trivial. Existing assistive technologies are either expensive, require specialized hardware (such as sensor gloves), or are restricted to a single language. These challenges motivated the development of a software-based, hardware-minimal, multilingual solution that can run on any device with a webcam and microphone.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("1.2 Problem Statement", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('The core problem addressed is the absence of an affordable, real-time, bidirectional translation system between Indian Sign Language and spoken/written languages. Existing sign language recognition systems are predominantly designed for American Sign Language (ASL), provide only one-way translation, require specialized hardware, do not support multilingual output, and do not account for ISL grammar (Object-Subject-Verb order) during speech-to-sign translation.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("1.3 Objective of the Project", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('1. Develop a real-time sign language recognition system using webcam, MediaPipe, and LSTM-TCN deep learning model.\n2. Implement a speech-to-sign translation pipeline with NLP-based ISL grammar transformation and 3D avatar animation.\n3. Support multilingual input and output in eight languages.\n4. Build a web-based, mobile-responsive application requiring no specialized hardware.\n5. Incorporate interactive learning features with ISL dictionary, gamified learning, and progress tracking.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("1.4 Scope", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('The scope encompasses: (a) Real-time recognition of 40 ISL phrases via webcam; (b) Speech-to-sign conversion with ISL grammar rules and 3D avatar; (c) ISL dictionary and learning module; (d) User management with authentication, preferences, and history; (e) Cloud deployment on Render. The system does not currently cover continuous sign language recognition or ISL dialects beyond the standard iSign dataset vocabulary.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("1.5 Project Introduction", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('Samvak (Sanskrit for "dialogue") is a web-based, real-time, bidirectional translation application. For Sign-to-Text, MediaPipe extracts 258-dimensional pose/hand features from webcam frames, and an LSTM-TCN model classifies gestures into ISL phrases with 94.6% accuracy. A geometry-based classifier provides fallback recognition for basic static signs. For Speech-to-Sign, the SpeechRecognition library transcribes voice input, spaCy tokenizes the text, and a local ISL grammar engine applies OSV reordering, article/be-verb dropping, and WH-word repositioning. The resulting gloss sequence drives a Three.js 3D avatar. The architecture uses Flask + SQLAlchemy backend, TensorFlow.js for in-browser inference, and Flask-SocketIO for real-time streaming.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 2. LITERATURE SURVEY
# ═══════════════════════════════════════════════════════════════
add_para("2. LITERATURE SURVEY", True, WD_ALIGN_PARAGRAPH.LEFT, 16)
add_empty()
add_para("2.1 Related Work", True, WD_ALIGN_PARAGRAPH.LEFT, 13)

refs = [
    ('[1] Wadhawan & Kumar (2021) — "Sign Language Recognition Systems: A Decade Systematic Literature Review"',
     'This survey covers SLR systems, classifying approaches into vision-based and sensor-based methods. Vision-based approaches using deep learning (CNN, LSTM) showed most promise. Most systems focused on ASL, highlighting the gap in ISL-specific recognition. This motivated our choice of LSTM with MediaPipe pose features for Samvak.'),
    ('[2] Lugaresi et al. (2019) — "MediaPipe: A Framework for Building Perception Pipelines"',
     'MediaPipe provides real-time detection of 33 pose and 21 per-hand landmarks from a single RGB camera. Samvak uses MediaPipe Holistic to extract a 258-dimensional feature vector per frame (33×4 pose + 21×3×2 hand coordinates). MediaPipe\'s browser-side JavaScript API enables the client-side inference architecture.'),
    ('[3] Hochreiter & Schmidhuber (1997) — "Long Short-Term Memory"',
     'The foundational LSTM paper addresses the vanishing gradient problem with gating mechanisms. Samvak employs an LSTM-TCN to classify 30-frame sequences of landmarks into ISL phrases, capturing temporal dependencies in dynamic signs like "Hello" (waving motion) and "Thank You" (mouth-to-outward sweep).'),
    ('[4] Saunders, Camgoz & Bowden (2020) — "Progressive Transformers for Sign Language Production"',
     'This work on sign language production informed Samvak\'s speech-to-sign pipeline design — converting English text into ISL gloss sequences as intermediate representation before driving the 3D avatar.'),
    ('[5] Paudyal, Lee & Banerjee (2023) — "A Survey on Sign Language Translation"',
     'The survey identifies that rule-based ISL grammar transformation achieves comparable accuracy to neural approaches for gloss generation in resource-limited scenarios. Samvak adopts a local rule-based ISL grammar engine.')
]

for title, desc in refs:
    add_para(title, True, size=13)
    add_para(desc, size=13, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_empty()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 3. SYSTEM ANALYSIS
# ═══════════════════════════════════════════════════════════════
add_para("3. SYSTEM ANALYSIS", True, WD_ALIGN_PARAGRAPH.LEFT, 16)
add_empty()

add_para("3.1 Existing System", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('Existing sign language translation systems are either hardware-based (CyberGlove, Kinect) requiring expensive specialized devices, or software-based systems primarily focused on American Sign Language (ASL). Open-source alternatives typically support only one-way translation, operate in a single language, and do not handle ISL grammar.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("3.2 Disadvantages of Existing System", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('1. Hardware Dependency: Require expensive sensor gloves or depth cameras ($500–$5,000).\n2. Limited Language Support: Work only with ASL, producing output only in English.\n3. Unidirectional: Support only sign-to-text, lacking speech-to-sign.\n4. No ISL Grammar: Perform naive word-by-word mapping without OSV reordering.\n5. Desktop-Bound: Cannot be accessed via web or mobile.\n6. No Learning Component: Focus solely on translation without teaching ISL.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("3.3 Proposed System", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('Samvak addresses all limitations: uses only webcam/microphone via MediaPipe; custom LSTM-TCN trained on iSign (40 ISL phrases, 94.6% accuracy); bidirectional translation; rule-based ISL grammar engine (OSV, article dropping, WH repositioning); 8 languages; Three.js 3D avatar; web-based with mobile responsiveness; gamified learning module.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("3.4 Advantages of Proposed System", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('1. Zero hardware cost — runs on any device with webcam.\n2. Real-time in-browser inference via TensorFlow.js (sub-200ms).\n3. Grammatically correct ISL gloss output.\n4. Multilingual support across 8 languages.\n5. Bidirectional communication for both hearing and Deaf users.\n6. Cloud-deployed on Render for universal accessibility.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("3.5 Workflow of Proposed System", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('The complete workflow of the proposed system handles both Sign-to-Text and Speech-to-Sign processes as well as the dictionary and gamified learning features. Below is the system workflow diagram:', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

wf_img = os.path.join(DIAGRAM_DIR, "workflow_diagram.png")
if os.path.exists(wf_img):
    doc.add_picture(wf_img, width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_para("Figure 3.1: System Workflow", False, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
add_empty()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 4. REQUIREMENT ANALYSIS
# ═══════════════════════════════════════════════════════════════
add_para("4. REQUIREMENT ANALYSIS", True, WD_ALIGN_PARAGRAPH.LEFT, 16)
add_empty()

add_para("4.1 Functional and Non-Functional Requirements", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para("Functional Requirements:", True, size=13)
add_para('1. Capture live video and detect hand/pose landmarks in real time using MediaPipe.\n2. Classify ISL gestures into 40 phrases using LSTM-TCN with ≥90% accuracy.\n3. Convert recognized text to speech using gTTS in the selected language.\n4. Accept voice input and transcribe using Google Speech-to-Text.\n5. Translate speech to ISL gloss sequences and animate via 3D avatar.\n6. Support 8 languages for input/output.\n7. Provide user registration, login, and session management.\n8. Maintain translation history for authenticated users.\n9. Provide an ISL dictionary with browsable sign image sequences.\n10. Implement gamified learning with XP tracking.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("Non-Functional Requirements:", True, size=13)
add_para('• Performance: Predictions within 200ms per frame sequence.\n• Security: PBKDF2-SHA256 password hashing; secure sessions.\n• Portability: Runs on all modern browsers across desktop and mobile.\n• Reliability: Graceful handling of camera/microphone unavailability.\n• Maintainability: Modular Flask blueprint architecture.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("4.2 Hardware Requirements", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
hw_table = doc.add_table(rows=8, cols=2, style='Table Grid')
hw_data = [
    ("Component", "Specification"),
    ("Processor", "Intel i3 / AMD Ryzen 3 or above"),
    ("RAM", "8 GB minimum"),
    ("Hard Disk", "256 GB SSD"),
    ("Webcam", "720p resolution minimum"),
    ("Microphone", "Standard built-in or external"),
    ("Monitor", "1366×768 resolution minimum"),
    ("Network", "Broadband Internet connection"),
]
for i, (c1, c2) in enumerate(hw_data):
    hw_table.rows[i].cells[0].text = c1
    hw_table.rows[i].cells[1].text = c2
    for cell in hw_table.rows[i].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
                if i == 0: r.bold = True

add_empty()

add_para("4.3 Software Requirements", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
sw_table = doc.add_table(rows=17, cols=2, style='Table Grid')
sw_data = [
    ("Component", "Specification"),
    ("Operating System", "Windows 10/11, macOS, or Linux"),
    ("Programming Language", "Python 3.10+"),
    ("Web Framework", "Flask 3.x with Jinja2"),
    ("Frontend", "HTML5, CSS3, JavaScript ES6+"),
    ("ML Framework", "TensorFlow 2.15+, TensorFlow.js 4.x"),
    ("Pose Estimation", "MediaPipe Holistic 0.10+"),
    ("Speech Recognition", "SpeechRecognition + Google STT"),
    ("NLP", "spaCy 3.x (en_core_web_sm)"),
    ("Translation", "deep-translator (Google Translate)"),
    ("Text-to-Speech", "gTTS (Google TTS)"),
    ("3D Rendering", "Three.js r160+"),
    ("Database", "SQLite 3 with SQLAlchemy ORM"),
    ("Authentication", "Flask-Login + Werkzeug"),
    ("Real-time", "Flask-SocketIO + Eventlet"),
    ("Deployment", "Render (Cloud PaaS)"),
    ("Version Control", "Git + GitHub"),
]
for i, (c1, c2) in enumerate(sw_data):
    sw_table.rows[i].cells[0].text = c1
    sw_table.rows[i].cells[1].text = c2
    for cell in sw_table.rows[i].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
                if i == 0: r.bold = True

add_empty()
add_para("4.4 Architecture", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('The system follows a client-server architecture. The browser handles real-time inference (TF.js, MediaPipe, Three.js) while the Flask backend manages authentication, translation, grammar processing, and database operations.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
arch_img = os.path.join(DIAGRAM_DIR, "deployment_diagram.png")
if os.path.exists(arch_img):
    doc.add_picture(arch_img, width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_para("Figure 4.1: System Architecture — Client-Server Deployment", False, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)

doc.add_page_break()



# ═══════════════════════════════════════════════════════════════
# 5. SYSTEM DESIGN
# ═══════════════════════════════════════════════════════════════
add_para("5. SYSTEM DESIGN", True, WD_ALIGN_PARAGRAPH.LEFT, 16)
add_empty()

add_para("5.1 Input and Output Design", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para("The input and output design of the system focuses on user-friendliness, accessibility, and real-time responsiveness. Below are the key design specifications:", False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
add_empty()

add_para("Input Design Specification:", True, size=13)
add_para('• Webcam Feed: Continuous 30 FPS RGB video stream captured via HTML5 <video> element.\n• Speech/Voice Input: Captured via the device microphone using the Web Speech API and processed by Google STT engine.\n• UI Controls: Intuitive buttons for toggling modes (Sign-to-Text / Speech-to-Sign), language selection dropdowns, and recording triggers.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_empty()
add_para("Output Design Specification:", True, size=13)
add_para('• Visual Translations: Real-time display of predicted text with high-contrast, readable font overlays.\n• Audio Output: Synthesized speech generated dynamically using Google Text-to-Speech (gTTS) mapped to the selected regional language.\n• 3D Avatar Rendering: A Three.js WebGL canvas displaying a humanoid avatar animating precise Indian Sign Language gestures based on the processed ISL gloss grammar.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

doc.add_page_break()

add_para("5.2 UML Diagrams", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
uml_diagrams = [
    ("Use Case Diagram", "use_case_diagram.png", "Figure 5.1: Use Case Diagram"),
    ("Class Diagram", "class_diagram.png", "Figure 5.2: Class Diagram"),
    ("Activity Diagram", "activity_diagram.png", "Figure 5.3: Activity Diagram"),
    ("Sequence Diagram: Sign-to-Text", "sequence_sign.png", "Figure 5.4: Sequence Diagram (Sign-to-Text)"),
    ("Sequence Diagram: Speech-to-Sign", "sequence_speech.png", "Figure 5.5: Sequence Diagram (Speech-to-Sign)"),
    ("Entity Relationship (ER) Diagram", "er_diagram.png", "Figure 5.6: ER Diagram"),
]

for title, filename, caption in uml_diagrams:
    add_para(title, True, WD_ALIGN_PARAGRAPH.LEFT, 13)
    img_path = os.path.join(DIAGRAM_DIR, filename)
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(caption, False, WD_ALIGN_PARAGRAPH.CENTER, 11)
    doc.add_page_break()

add_para("5.3 DFD Diagrams", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
dfd_diagrams = [
    ("Level 0 DFD (Context Diagram)", "dfd_level0.png", "Figure 5.7: Level 0 DFD"),
    ("Level 1 DFD", "dfd_level1.png", "Figure 5.8: Level 1 DFD"),
    ("Level 2 DFD (Translation Processes)", "dfd_level2.png", "Figure 5.9: Level 2 DFD"),
]
for title, filename, caption in dfd_diagrams:
    add_para(title, True, WD_ALIGN_PARAGRAPH.LEFT, 13)
    img_path = os.path.join(DIAGRAM_DIR, filename)
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(caption, False, WD_ALIGN_PARAGRAPH.CENTER, 11)
    if title != dfd_diagrams[-1][0]:
        doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# 6. CODE, IMPLEMENTATION AND RESULT
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
add_para("6. CODE, IMPLEMENTATION AND RESULT", True, WD_ALIGN_PARAGRAPH.LEFT, 16)
add_empty()

add_para("6.1 Implementation Details", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para("The implementation consists of the following key modules:", False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
add_para(
    "1. MediaPipe Integration: Hand and pose landmarks are extracted iteratively using MediaPipe Holistic API across a 30-frame temporal window.\n"
    "2. Deep Learning Pipeline: Captured landmark features (258 dimensions per frame) are scaled and passed to a trained Keras Sequential model comprising Temporal Convolutional Network (TCN) blocks and LSTM layers to classify gestures into one of 40 potential ISL classes.\n"
    "3. NLP Grammar Engine: For the Speech-to-Sign mechanism, spaCy handles dependency parsing to reorganize English Subject-Verb-Object (SVO) structures into Indian Sign Language Object-Subject-Verb (OSV) semantic formats, omitting articles and auxiliary 'be' verbs.\n"
    "4. Web Integration: Flask governs the core routing logic, SQLAlchemy manages user session state (SQLite), and Flask-SocketIO binds the front-end Three.js web-worker events back to the server in real-time.",
    False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13
)
add_empty()

add_para("6.2 Screenshots", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
screenshots = [
    ("Login / Register Page", "login.png"),
    ("User Dashboard", "dashboard.png"),
    ("Sign-to-Text Mode", "sign.png"),
    ("Speech-to-Sign Mode", "speech.png"),
    ("ISL Dictionary", "dict.png"),
    ("Gamified Learning Module", "learn.png")
]

fig_num = 1
for title, filename in screenshots:
    add_para(f"6.2.{fig_num} {title}", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
    img_path = os.path.join(SCREENSHOT_DIR, filename)
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(f"Figure 6.{fig_num}: {title}", False, WD_ALIGN_PARAGRAPH.CENTER, 11)
    add_empty()
    fig_num += 1


# ═══════════════════════════════════════════════════════════════
# 7. SYSTEM STUDY AND TESTING
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
add_para("7. SYSTEM STUDY AND TESTING", True, WD_ALIGN_PARAGRAPH.LEFT, 16)
add_empty()

add_para("7.1 Feasibility Study", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para("The feasibility study investigates the viability of the system under technical, economic, and operational standards.", False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
add_para("• Technical Feasibility: The implementation capitalizes on open-source ML frameworks (TensorFlow) and web APIs (MediaPipe), making hardware acceleration unnecessary for inference. High performance is maintained using client-side WebGL.\n• Economic Feasibility: Since the target solution is software-deployed requiring only a built-in webcam, the cost barrier is virtually zero.\n• Operational Feasibility: The gamified dictionary helps flatten the learning curve for users unfamiliar with ISL, solidifying high operational viability.", False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
add_empty()

add_para("7.2 System Testing", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para("Testing was conducted to ensure robustness, precision in detection, and minimal latency handling during real-time video streaming.", False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
add_empty()

add_para("7.3 Types of Tests", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para("Unit Testing: Model validation tests confirmed that the LSTM-TCN accurately identified targeted phrases (e.g., 'Hello', 'Thank You') under diverse lighting conditions.\nIntegration Testing: Validation between the Flask backend text generation and Google STT confirmed that language switching worked continuously without socket disconnects.\nSystem Testing: The comprehensive end-to-end flow was tested; the avatar matched the grammatical outputs strictly defined by the NLP OSV parser without desynchronization.", False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)


# ═══════════════════════════════════════════════════════════════
# 8. CONCLUSION
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
add_para("8. CONCLUSION", True, WD_ALIGN_PARAGRAPH.LEFT, 16)
add_empty()
add_para('The "Multi-Lingual Sign Language to Speech and Speech to Sign Translator" establishes a critical, scalable framework for real-time human-to-human interaction breaking typical accessibility barriers. By abstracting the need for advanced hardware gloves via an accurate MediaPipe and LSTM-TCN integrated architecture, this project delivers high-accuracy sign classification directly to the browser. The addition of the Speech-to-Sign mechanism paired with the rule-based Indian Sign Language semantic generator successfully demonstrates that conversational continuity can be preserved between standard vocal languages and non-vocal ISL.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
add_para('Furthermore, expanding beyond simple English translations into a fully localized multi-lingual solution extends the system’s usability to a massive population sector. Given its high prediction accuracy and low operational latency, Samvak proves to be an effective, responsive, and innovative technological intervention serving the Deaf community.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)


# ═══════════════════════════════════════════════════════════════
# 9. FUTURE ENHANCEMENT
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
add_para("9. FUTURE ENHANCEMENT", True, WD_ALIGN_PARAGRAPH.LEFT, 16)
add_empty()
add_para("While the current prototype validates the efficacy of bidrectional translation, multiple enhancements could scale its applicability in future versions:", False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
add_para("1. Continuous Sign Recognition: Expanding from isolated phrases to connected sign language sequences by integrating Transformer-based architectures.\n2. Generative Avatar Motions: Shifting from predefined Three.js skeletal animations to Generative Adversarial Networks (GANs) that can organically render smooth, human-like avatar transitions.\n3. Regional Dialect Implementation: Collecting additional real-world data to expand the ISL datasets enabling recognition of region-specific linguistic differences.\n4. Mobile Application: Compiling the architecture into native iOS and Android build environments for improved mobile GPU performance optimization.", False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)


# ═══════════════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
add_para("REFERENCES", True, WD_ALIGN_PARAGRAPH.LEFT, 16)
add_empty()

references = [
    "[1] Wadhawan, A., & Kumar, P. (2021). Sign Language Recognition Systems: A Decade Systematic Literature Review. Archives of Computational Methods in Engineering, 28(3), 785-813.",
    "[2] Lugaresi, C. et al. (2019). MediaPipe: A Framework for Building Perception Pipelines. arXiv preprint arXiv:1906.08172.",
    "[3] Hochreiter, S., & Schmidhuber, J. (1997). Long Short-Term Memory. Neural Computation, 9(8), 1735-1780.",
    "[4] Exploration Lab, IIT Kharagpur (2024). iSign: A Benchmark for Indian Sign Language Processing. Proceedings of IEEE CVPR.",
    "[5] Paudyal, P., Lee, J., & Banerjee, S. (2023). A Survey on Sign Language Translation: Neural and Rule-based Approaches. ACM Computing Surveys."
]

for ref in references:
    add_para(ref, False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
    add_empty()

print("Successfully generated final Word document!")


doc.save(r"E:\sam\Project documents\Samvak_ProjectReview_Final_V5.docx")
print("Successfully generated final Word document V5!")
