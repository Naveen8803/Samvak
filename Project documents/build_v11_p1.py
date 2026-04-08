import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document(r'E:\sam\Project documents\Samvak_Base_Clean_V2.docx')
DIAGRAM_DIR = r"E:\sam\Project documents\diagrams"

def add_para(text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=13, space_before=None, space_after=None):
    p = doc.add_paragraph()
    p.alignment = align
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = 'Times New Roman'
    return p

def add_empty(n=1):
    for _ in range(n):
        doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# ABSTRACT
# ═══════════════════════════════════════════════════════════════
add_empty()
add_para('ABSTRACT', True, WD_ALIGN_PARAGRAPH.CENTER, 14)
add_empty()
add_para('This project aims to solve the severe communication gap between the hearing-impaired community and the general population through an AI-based bidirectional translation system. Traditional communication relies heavily on human sign language interpreters, limiting accessibility and independence for Deaf individuals.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
add_para('The proposed Multi-Lingual Sign Language to Speech and Speech to Sign Translator provides a software-based solution leveraging Machine Learning, Natural Language Processing (NLP), and Computer Vision (CV). Operating via a standard webcam and microphone, the system utilizes MediaPipe Hand and Pose tracking to extract gesture landmarks, which are classified into accurate text using an advanced deep learning model. For bidirectional dialogue, the system captures spoken audio, employs Google Speech-to-Text for transcription, processes it through an NLP engine to map spoken grammar into Indian Sign Language (ISL) syntactical structure, and visually renders the translation using a 3D avatar animations.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
add_para('The outcome is a highly accessible, multi-lingual web platform capable of recognizing complex gestures and converting them to regional speech while simultaneously translating audio input directly into real-time ISL gestures. This abolishes the reliance on costly sensor hardware, drastically improving independence and inclusivity.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════
add_para('TABLE OF CONTENTS', True, WD_ALIGN_PARAGRAPH.CENTER, 14)
add_empty()
toc_data = [
    ("CHAPTER NO.", "TITLE", "PAGE NO."),
    ("1", "INTRODUCTION", "1"),
    ("1.1", "Motivation", "1"),
    ("1.2", "Problem Statement", "1"),
    ("1.3", "Objectives", "2"),
    ("1.4", "Scope", "2"),
    ("1.5", "Overview of the System", "2"),
    ("2", "LITERATURE SURVEY", "3"),
    ("3", "SYSTEM ANALYSIS", "5"),
    ("3.1", "Existing System", "5"),
    ("3.2", "Disadvantages of Existing System", "5"),
    ("3.3", "Proposed System", "6"),
    ("3.4", "Advantages of Proposed System", "6"),
    ("3.5", "Workflow Explanation", "6"),
    ("4", "REQUIREMENT ANALYSIS", "8"),
    ("4.1", "Functional Requirements", "8"),
    ("4.2", "Non-Functional Requirements", "8"),
    ("4.3", "Hardware Requirements", "9"),
    ("4.4", "Software Requirements", "9"),
    ("4.5", "System Architecture", "10"),
    ("5", "SYSTEM DESIGN", "11"),
    ("5.1", "Use Case Diagram", "11"),
    ("5.2", "Class Diagram", "12"),
    ("5.3", "Sequence Diagram", "13"),
    ("5.4", "Activity Diagram", "14"),
    ("5.5", "Component Diagram", "15"),
    ("5.6", "Deployment Diagram", "15"),
    ("5.7", "ER Diagram", "16"),
    ("5.8", "Data Flow Diagram", "16"),
    ("6", "CODE, IMPLEMENTATION AND RESULTS", "18"),
    ("6.1", "Modules Overview", "18"),
    ("6.2", "Sample Implementation Code", "20"),
    ("6.3", "Output Screens", "22"),
    ("7", "SYSTEM STUDY AND TESTING", "24"),
    ("7.1", "Feasibility Study", "24"),
    ("7.2", "Types of Testing", "24"),
    ("7.3", "Test Cases", "25"),
    ("8", "CONCLUSION", "26"),
    ("9", "FUTURE ENHANCEMENT", "27"),
    ("10", "REFERENCES", "28")
]
toc_table = doc.add_table(rows=len(toc_data), cols=3)
toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (c1, c2, c3) in enumerate(toc_data):
    toc_table.rows[i].cells[0].text = c1
    toc_table.rows[i].cells[1].text = c2
    toc_table.rows[i].cells[2].text = c3
    for cell in toc_table.rows[i].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
                if i == 0 or len(c1) == 1 or c1 == "":  
                    r.bold = True
                else:
                    r.bold = False

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ═══════════════════════════════════════════════════════════════
add_para("1. INTRODUCTION", True, WD_ALIGN_PARAGRAPH.LEFT, 16)
add_empty()

add_para("1.1 Motivation", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('According to global health records, millions of individuals navigate society primarily through sign language. Because the general public is unversed in sign vocabulary, the Deaf community faces profound societal isolation. Developing an autonomous, AI-driven software translating sign language to speech breaks this barrier natively without demanding users to wear intrusive sensors.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("1.2 Problem Statement", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('Current interaction solutions between hearing and non-hearing individuals heavily depend on manual interpretation. Digital alternatives usually concentrate uniquely on American Sign Language while relying on high-cost hardware gloves. A lack of bidirectional functionality—meaning they translate signs to words but fail to translate spoken words back into sign patterns—renders daily conversational integration highly inefficient.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("1.3 Objectives", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('1. Develop a bidirectional Sign-to-Speech and Speech-to-Sign translation system using Computer Vision.\n2. Utilize MediaPipe arrays to dynamically recognize multi-hand structures substituting hardware layers.\n3. Integrate Google Speech-to-Text establishing robust microphone transcription.\n4. Design scalable multi-lingual pipelines permitting real-time vocalized outputs natively in regional languages.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("1.4 Scope", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('The software bounds encompass isolated sign classification over a webcam matrix parsing direct NLP manipulations through Python pipelines. It enables users to switch localized vernacular outputs through a web interface, rendering audio and a dynamically mapped 3D avatar executing precise gloss motions accurately modeling Indian Sign Language mechanics.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

add_para("1.5 Overview of the System", True, WD_ALIGN_PARAGRAPH.LEFT, 13)
add_para('The system encapsulates both detection and rendering modules natively within web protocols. Camera streams feed isolated landmark features (hands, shoulders, elbows) matching complex gestures through integrated deep neural weights generating accurate strings. Subsequent text arrays trigger real-time Text-to-Speech functions mapping output sounds. Inversely, user audio feeds invoke Natural Language processing isolating the Subject/Object logic mapping resulting tokens rendering 3D avatar skeletal nodes executing translation sequences comprehensively.', False, WD_ALIGN_PARAGRAPH.JUSTIFY, 13)

doc.add_page_break()

# Save intermediary script 1
doc.save(r'E:\sam\Project documents\Samvak_Temp1.docx')
